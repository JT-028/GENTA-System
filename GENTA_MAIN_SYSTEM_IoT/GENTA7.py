import google.generativeai as genai
import os
import sys
import logging
import time
import docx

# Force UTF-8 encoding for stdout to prevent UnicodeEncodeError
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')
import random
import re
import html
import csv
import mysql.connector
from mysql.connector import Error
from datetime import datetime
from os import environ

# Suppress pygame and pkg_resources deprecation warnings
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="pygame.pkgdata")
warnings.filterwarnings("ignore", category=DeprecationWarning, module="pkg_resources")

environ['PYGAME_HIDE_SUPPORT_PROMPT'] = '1'
# Setup simple logging for the application. Level can be controlled via GENTA_LOG_LEVEL env var.
LOG_LEVEL = os.environ.get('GENTA_LOG_LEVEL', 'INFO').upper()
logging.basicConfig(level=getattr(logging, LOG_LEVEL, logging.INFO), format='%(asctime)s %(levelname)s: %(message)s')
import config
import pygame
import requests
from pydub import AudioSegment
from google.cloud import speech, translate_v2, texttospeech_v1
import threading
import sys
import os.path
import shutil
import subprocess
import io
import urllib.parse
from concurrent.futures import ThreadPoolExecutor
from functools import lru_cache
import wave
import struct
import socket
import ast

# ============================================================================
# ESP32 AUTO-DISCOVERY: Fetch current IPs from Flask UDP discovery system
# ============================================================================
def discover_esp32_devices(flask_host='localhost', flask_port=5000, timeout=2, retry_interval=5, block_until_found=True):
    """
    Query Flask server for discovered ESP32 devices via UDP announcements.
    
    Args:
        flask_host: Flask server hostname
        flask_port: Flask server port
        timeout: HTTP request timeout
        retry_interval: Seconds to wait between retries
        block_until_found: If True, will retry indefinitely until devices found
    
    Returns:
        dict: {'recorder': 'IP', 'player': 'IP'} or None if discovery fails (when block_until_found=False)
    """
    # Allow overriding the Flask discovery endpoint via environment variables so
    # users can point discovery at an ngrok/public URL. Supported env vars:
    # - GENTA_FLASK_URL (full URL, e.g. https://<subdomain>.ngrok-free.dev)
    # - GENTA_FLASK_HOST (host[:port], e.g. localhost:5000)
    gflask_url = os.environ.get('GENTA_FLASK_URL') or os.environ.get('GENTA_FLASK_HOST')
    use_full_url = False
    base_url = None
    if gflask_url:
        # Normalize: if it looks like a full URL (starts with http), use directly
        if gflask_url.startswith('http://') or gflask_url.startswith('https://'):
            base_url = gflask_url.rstrip('/')
            use_full_url = True
        else:
            # treat as host[:port]
            base_url = f'http://{gflask_url.rstrip('/')}'
            use_full_url = True

    attempt = 0
    while True:
        attempt += 1
        try:
            # Choose endpoint: prefer environment-configured base_url (ngrok/public tunnel)
            if base_url:
                query_url = f"{base_url}/api/discovered_devices"
            else:
                query_url = f'http://{flask_host}:{flask_port}/api/discovered_devices'
            print(f"[Discovery Attempt #{attempt}] Querying Flask at {query_url} ...")
            response = requests.get(query_url, timeout=timeout)
            if response.status_code == 200:
                devices = response.json()
                print(f"[Discovery] Flask responded with {len(devices)} device(s)")
                
                result = {}
                for mac, info in devices.items():
                    role = info.get('role', '').lower()
                    ip = info.get('ip')
                    last_seen = info.get('last_seen', 0)
                    if ip:
                        print(f"[Discovery] Found {role}: {ip} (MAC: {mac})")
                        if 'record' in role:
                            result['recorder'] = ip
                        elif 'play' in role:
                            result['player'] = ip
                
                if result:
                    print(f"✓ Auto-discovered ESP32 devices: {result}")
                    return result
                else:
                    print("⚠ No recorder or player devices found in response")
                    if not block_until_found:
                        return None
            else:
                print(f"⚠ Flask API returned status {response.status_code}")
                if not block_until_found:
                    return None
        except requests.exceptions.ConnectionError as e:
            print(f"⚠ Flask server not reachable (attempt #{attempt})")
            if not block_until_found:
                print(f"  Make sure Flask is running: python GENTA_Flask.py")
                return None
        except requests.exceptions.Timeout as e:
            print(f"⚠ Flask server timeout (attempt #{attempt})")
            if not block_until_found:
                return None
        except Exception as e:
            print(f"⚠ Discovery error (attempt #{attempt}): {e}")
            if not block_until_found:
                return None
        
        # If we get here and block_until_found is True, wait and retry
        print(f"⏳ Waiting {retry_interval} seconds before retry...")
        print("   Make sure:")
        print("   1. Flask is running: python GENTA_Flask.py")
        print("   2. ESP32 devices are powered on and connected to WiFi")
        print("   3. GENTA.ino announces role='recorder'")
        print("   4. GENTA2.ino announces role='player'")
        print("   5. All devices are on the same network")
        import time
        time.sleep(retry_interval)

print("\n" + "="*70)
print("ESP32 DEVICE DISCOVERY")
print("="*70)

# Try auto-discovery first, fall back to environment variables or hardcoded defaults
_discovered = discover_esp32_devices()

# Determine recorder IP: prefer auto-discovery, but fall back to env/config if discovery missed it
recorder_ip = None
if _discovered and 'recorder' in _discovered:
    recorder_ip = _discovered['recorder']
else:
    # Try environment variable or config fallback
    recorder_ip = os.environ.get('GENTA_ESP_RECORD_IP') or getattr(config, 'ESP_RECORD_IP', None)

# Send wakeup signal to ESP32 recorder to trigger eye animation (use fallback if needed)
if recorder_ip:
    print("\n" + "="*70)
    print("SENDING WAKEUP SIGNAL TO ESP32 (recorder) - (non-blocking)")
    print("="*70)

    def _send_wakeup_async(ip, timeout=5):
        # Attempt a fire-and-forget raw-socket GET first so we don't block
        # waiting for a full HTTP response when the ESP performs the action
        # but doesn't send headers/body promptly (causes requests ReadTimeout).
        try:
            print(f"Sending wakeup signal to {ip} (fire-and-forget socket)...")
            try:
                s = socket.create_connection((ip, 80), timeout=2)
                req = f"GET /wakeup HTTP/1.1\r\nHost: {ip}\r\nConnection: close\r\n\r\n"
                s.sendall(req.encode('utf-8'))
                # Close immediately without waiting for response
                s.close()
                print("✓ ESP32 wakeup request sent (not waiting for response).")
                print("  OLED eyes should now be showing the startup animation")
                return
            except Exception:
                # If raw socket approach fails (e.g., windows/network rules), fall back
                # to the original requests-based call but with a slightly larger
                # read timeout to allow the device to respond.
                wakeup_url = f'http://{ip}/wakeup'
                print(f"Socket method failed; falling back to HTTP GET {wakeup_url} (timeout={timeout}s)...")
                try:
                    resp = requests.get(wakeup_url, timeout=(2, max(timeout, 6)))
                    if getattr(resp, 'status_code', None) == 200:
                        print("✓ ESP32 wakeup sequence initiated successfully (HTTP)!")
                        print("  OLED eyes should now be showing the startup animation")
                        return
                    else:
                        print(f"⚠ Wakeup signal sent but received status {getattr(resp, 'status_code', 'N/A')}")
                        return
                except Exception as e2:
                    print(f"⚠ Could not send wakeup signal (HTTP fallback): {e2}")
                    print("  OLED will remain in booting state")
                    return
        except Exception as e:
            print(f"⚠ Could not send wakeup signal: {e}")
            print("  OLED will remain in booting state")

    # Launch wakeup in background so slow/unresponsive devices won't hang startup
    try:
        threading.Thread(target=_send_wakeup_async, args=(recorder_ip, 5), daemon=True).start()
    except Exception:
        # Fallback: try synchronous (best-effort)
        try:
            _send_wakeup_async(recorder_ip, timeout=5)
        except Exception:
            pass

    print("="*70 + "\n")
else:
    print("⚠ No recorder IP available (discovery failed and no GENTA_ESP_RECORD_IP/config.ESP_RECORD_IP). Skipping wakeup.")

print("\n" + "="*70)
print("ESP32 CONFIGURATION")
print("="*70)

# ESP32 Player (GENTA2.ino - Speaker/Playback device)
if _discovered and 'player' in _discovered:
    esp_playback_host = _discovered['player']
    print(f"✓ ESP32 Player: {esp_playback_host} (auto-discovered)")
else:
    esp_playback_host = os.environ.get('GENTA_ESP_PLAYBACK_IP', getattr(config, 'ESP_PLAYBACK_HOST', '192.168.50.62'))
    if 'GENTA_ESP_PLAYBACK_IP' in os.environ:
        print(f"ESP32 Player: {esp_playback_host} (from environment variable)")
    else:
        print(f"⚠ ESP32 Player: {esp_playback_host} (default - auto-discovery failed)")
        print("  Set GENTA_ESP_PLAYBACK_IP environment variable or start Flask server")

print("="*70 + "\n")

ARDUINO_DATA_DIR = getattr(config, 'ARDUINO_DATA_DIR', r"c:\Users\vonti\OneDrive\Desktop\GENTA SYS\ARDUINO\GENTA2\data")

# Module-level TTS helpers (global) - used by multiple functions across the module.
def synthesize_speech(text, out_path='response.wav', sample_rate_hz=24000):
    """Module-level synthesize wrapper that writes linear16 WAV to out_path.
    Default writes to 'response.wav' at 24kHz to match ESP32 player expectations.
    Returns out_path on success.
    """
    try:
        # Determine voice settings from environment or centralized config
        # Priority: env vars > config.py > built-in defaults
        voice_name = os.environ.get('GENTA_TTS_VOICE') or getattr(config, 'GENTA_TTS_VOICE', None) or 'en-US-Neural2-C'
        language_code = os.environ.get('GENTA_TTS_LANGUAGE') or getattr(config, 'GENTA_TTS_LANGUAGE', None) or 'en-US'
        # pitch may be provided as float-like string; default 0.0 for neutral voice
        try:
            pitch = float(os.environ.get('GENTA_TTS_PITCH', getattr(config, 'GENTA_TTS_PITCH', 0.0)))
        except Exception:
            pitch = 0.0

        client = texttospeech_v1.TextToSpeechClient()
        voice = texttospeech_v1.VoiceSelectionParams(name=voice_name, language_code=language_code)
        audio_cfg = texttospeech_v1.AudioConfig(
            audio_encoding=texttospeech_v1.AudioEncoding.LINEAR16,
            pitch=float(pitch),
            sample_rate_hertz=int(sample_rate_hz)
        )
        response = client.synthesize_speech(input=texttospeech_v1.SynthesisInput(text=text), voice=voice, audio_config=audio_cfg)
        with open(out_path, 'wb') as fh:
            fh.write(response.audio_content)
        return out_path
    except Exception as e:
        print(f"⚠ synthesize_speech (module-level) failed: {e}")
        raise

def speak_and_play_text(text, max_chunk_chars=240):
    """Chunked TTS + playback helper.

        Splits `text` into sentence-like chunks (or by roughly `max_chunk_chars`) and
        synthesizes/uploads/plays each chunk immediately so audio starts earlier.
        Preference order for playback/upload:
            1. Upload to `/upload` (LittleFS) which autoplays then removes file
            2. Fallback to local `play_audio`
    Returns True if at least one chunk played/uploaded successfully.
    """
    try:
        if not text or not text.strip():
            return False

        # Simple sentence splitter using punctuation; fallback to fixed-size slices
        import re
        sentences = re.split(r'(?<=[\.\!\?])\s+', text.strip())
        # Merge small sentences to avoid too-short chunks
        chunks = []
        cur = ''
        for s in sentences:
            if len(cur) + len(s) + 1 <= max_chunk_chars:
                cur = (cur + ' ' + s).strip() if cur else s
            else:
                if cur:
                    chunks.append(cur)
                cur = s
        if cur:
            chunks.append(cur)

        # If splitting produced no meaningful sentences (very long without punctuation), slice by chars
        if not chunks:
            chunks = [text[i:i+max_chunk_chars].strip() for i in range(0, len(text), max_chunk_chars)]

        played_any = False
        for idx, chunk in enumerate(chunks):
            try:
                fname = f'response_chunk_{idx}.wav'
                synthesize_speech(chunk, out_path=fname, sample_rate_hz=24000)

                # Use the default upload endpoint (LittleFS) on the device
                try_endpoints = ['/upload']
                uploaded = False
                for ep in try_endpoints:
                    try:
                        ok = esp_upload_file(fname, endpoint=ep, max_retries=2, auto_delete_after_playback=(ep == '/upload'))
                        if ok:
                            uploaded = True
                            played_any = True
                            break
                    except Exception:
                        continue

                if not uploaded:
                    # Fallback to local playback
                    try:
                        pl = globals().get('play_audio')
                        if callable(pl):
                            try:
                                play_ok = pl(fname)
                                if play_ok:
                                    played_any = True
                            except Exception:
                                pass
                    except Exception:
                        pass

                # Tiny gap so device has time to start playback; keeps flow snappy
                time.sleep(0.18)
            except Exception as e:
                print(f"⚠ chunk playback/upload failed for chunk {idx}: {e}")
                continue

        return played_any
    except Exception as e:
        print(f"⚠ speak_and_play_text exception: {e}")
        return False


# --- Canonical global helpers (single implementations) -------------------
def TranslateToFil(text):
    """Translate `text` to Filipino using Google Translate API and return text.
    Writes a copy to `output.txt` under `UPLOAD_DIR` when possible.
    """
    try:
        translate_client = translate_v2.Client()
        translated_response = translate_client.translate(text, target_language="fil")
        translated_text = translated_response.get('translatedText', translated_response)
        try:
            outp = U('output.txt')
            os.makedirs(os.path.dirname(outp), exist_ok=True)
        except Exception:
            outp = 'output.txt'
        try:
            with open(outp, 'w', encoding='utf-8') as f:
                f.write(translated_text)
        except Exception:
            pass
        return translated_text
    except Exception as e:
        print(f"⚠ TranslateToFil failed: {e}")
        return text

def TranslateToEng(text):
    """Translate `text` to English using Google Translate API and return text."""
    try:
        translate_client = translate_v2.Client()
        translated_response = translate_client.translate(text, target_language="en")
        translated_text = translated_response.get('translatedText', translated_response)
        return translated_text
    except Exception as e:
        print(f"⚠ TranslateToEng failed: {e}")
        return text

def play_audio(file):
    """Module-level play_audio: upload to ESP32 player (preferred) or fallback to local playback.
    Returns True if ESP playback was used, False otherwise.
    """
    try:
        destination_file = U('GENTA_response.mp3')
        os.makedirs(os.path.dirname(destination_file), exist_ok=True)
    except Exception:
        destination_file = 'GENTA_response.mp3'

    if os.path.exists(destination_file):
        try:
            os.remove(destination_file)
        except Exception:
            pass

    def _upload_and_play(wav_path):
        try:
            try:
                os.makedirs(ARDUINO_DATA_DIR, exist_ok=True)
                dest = os.path.join(ARDUINO_DATA_DIR, 'response.wav')
                shutil.copy2(wav_path, dest)
            except Exception:
                pass
            ok = esp_upload_file(wav_path, endpoint='/upload', max_retries=3)
            return ok
        except Exception:
            return False

    esp_success = False
    try:
        import socket
        try:
            s = socket.create_connection((esp_playback_host, 80), timeout=1.0)
            s.close()
            esp_success = _upload_and_play(file)
        except Exception:
            esp_success = False
    except Exception:
        esp_success = False

    if not esp_success:
        # Local fallback for development/test environments
        try:
            pygame.mixer.init()
            pygame.mixer.music.load(file)
            pygame.mixer.music.play()
            while pygame.mixer.music.get_busy():
                pygame.time.Clock().tick(10)
            pygame.mixer.quit()
        except Exception:
            pass

    try:
        os.rename(file, destination_file)
    except Exception:
        pass

    return esp_success

# -------------------------------------------------------------------------


# OPTIMIZATION: Global HTTP session for connection pooling (reuse connections)
_http_session = requests.Session()
_http_session.headers.update({
    'ngrok-skip-browser-warning': 'true',
    'User-Agent': 'GENTA-System/1.0'
})


def http_get_with_retries(url, params=None, timeout=1.0, retries=3, backoff=0.5):
    """GET with simple retries/backoff and detailed logging. Returns Response or None."""
    last_exc = None
    for attempt in range(1, retries + 1):
        try:
            resp = _http_session.get(url, params=params, timeout=timeout)
            # Log non-2xx/3xx responses for debugging
            if resp is not None and getattr(resp, 'status_code', None) and resp.status_code >= 400:
                try:
                    body = resp.text[:400]
                except Exception:
                    body = '<unreadable body>'
                print(f"[HTTP GET] {url} returned status {resp.status_code} (attempt {attempt}/{retries}) - body: {body}")
            return resp
        except Exception as e:
            last_exc = e
            print(f"[HTTP GET] Exception for {url} (attempt {attempt}/{retries}): {e}")
            if attempt < retries:
                time.sleep(backoff * attempt)
    print(f"[HTTP GET] All {retries} attempts failed for {url}: {last_exc}")
    return None


def try_direct_oled_status(recorder_ip, timeout=4.0):
    """Attempt to call the recorder device's /oled_status directly. Returns parsed JSON dict or None."""
    if not recorder_ip:
        return None
    try:
        url = f'http://{recorder_ip}/oled_status'
        r = _http_session.get(url, timeout=timeout)
        if r is None or getattr(r, 'status_code', None) != 200:
            return None
        try:
            return r.json()
        except Exception:
            # If device returns non-JSON, attempt to interpret text
            text = getattr(r, 'text', None)
            return {'completion_played': False, 'raw': text}
    except Exception:
        return None


def try_direct_oled_set(recorder_ip, params=None, timeout=4.0):
    """Attempt to call the recorder device's /oled endpoint directly. Returns True on 200."""
    if not recorder_ip:
        return False
    try:
        url = f'http://{recorder_ip}/oled'
        r = _http_session.get(url, params=params, timeout=timeout)
        return r is not None and getattr(r, 'status_code', None) == 200
    except Exception:
        return False


def http_post_with_retries(url, data=None, files=None, timeout=3.0, retries=3, backoff=0.5):
    """POST with retries/backoff. Returns Response or None."""
    last_exc = None
    for attempt in range(1, retries + 1):
        try:
            resp = _http_session.post(url, data=data, files=files, timeout=timeout)
            if resp is not None and getattr(resp, 'status_code', None) and resp.status_code >= 400:
                try:
                    body = resp.text[:400]
                except Exception:
                    body = '<unreadable body>'
                print(f"[HTTP POST] {url} returned status {resp.status_code} (attempt {attempt}/{retries}) - body: {body}")
            return resp
        except Exception as e:
            last_exc = e
            print(f"[HTTP POST] Exception for {url} (attempt {attempt}/{retries}): {e}")
            if attempt < retries:
                time.sleep(backoff * attempt)
    print(f"[HTTP POST] All {retries} attempts failed for {url}: {last_exc}")
    return None

# OPTIMIZATION: Thread pool for parallel operations
_thread_pool = ThreadPoolExecutor(max_workers=3)

# OPTIMIZATION: Cache ffmpeg path lookup
_cached_ffmpeg_path = None

def _get_ffmpeg_path():
    """Cached ffmpeg path lookup to avoid repeated filesystem searches."""
    global _cached_ffmpeg_path
    if _cached_ffmpeg_path is None:
        _cached_ffmpeg_path = shutil.which('ffmpeg') or shutil.which('ffmpeg.exe') or ''
    return _cached_ffmpeg_path if _cached_ffmpeg_path else None

# Canonical list of supported OLED expressions. Other helpers reference this.
OLED_VALID_EXPRESSIONS = [
    'idle', 'listening', 'thinking', 'quiz', 'assist',
    'correct', 'happy', 'excited', 'glee', 'awe', 'surprised',
    'incorrect', 'sad', 'worried', 'frustrated', 'annoyed', 'angry', 'furious', 'scared',
    'curious', 'skeptical', 'suspicious', 'focused', 'squint', 'unimpressed', 'sleepy',
    'error'
]

def set_oled_expression(expression: str):
    """
    Control the ESP32 OLED display to show different eye expressions.
    NON-BLOCKING: Runs in background thread to avoid delays.
    
    Args:
        expression: Eye expression name (see available expressions below)
    
    Available Cozmo-style expressions:
        Core:
        - idle: Default neutral eyes
        - listening: Eyes squint (attentive, focused look)
        - thinking: Inverted arcs (pondering look, like Cozmo)
        
        Positive emotions:
            - happy: Happy eyes with upward curve
        - excited: Rapid shimmy animation (celebration!)
        - glee: Happy with bounce (playful joy)
        - awe: Very wide amazed eyes
        - surprised: Sudden wide open eyes
        
        Negative emotions:
        - incorrect/sad: Looking down, droopy eyes
        - worried: Raised inner corners, anxious
        - frustrated: Eyes drift to side, droopy
        - annoyed: Asymmetric squint
        - angry: Sharp inward tilt, narrowed
        - furious: Intense anger with shake
        - scared: Wide eyes looking up
        
        Neutral/Curious:
        - curious: Asymmetric eyes with scanning motion
            - skeptical: One eye smaller, questioning
            - suspicious: Side-eye with squint
        - focused: Narrow intense eyes
        - squint: Both eyes narrow
        - unimpressed: Flat half-closed eyes
        - sleepy: Very narrow, slow blink
        
        System:
        - error: Sleep animation (closed eyes)
    """
    # Use the module-level canonical list so other helpers can reference it
    try:
        valid_expressions = OLED_VALID_EXPRESSIONS
    except NameError:
        valid_expressions = [
            'idle', 'listening', 'thinking', 'quiz', 'assist',
            'correct', 'happy', 'excited', 'glee', 'awe', 'surprised',
            'incorrect', 'sad', 'worried', 'frustrated', 'annoyed', 'angry', 'furious', 'scared',
            'curious', 'skeptical', 'suspicious', 'focused', 'squint', 'unimpressed', 'sleepy',
            'error'
        ]
    if expression not in valid_expressions:
        print(f"[OLED] ⚠ Invalid expression '{expression}'. Valid: {valid_expressions}")
        return
    # Special-case: 'assist' is a one-shot transition; avoid re-sending repeatedly
    try:
        global _last_oled_expression
        if expression == 'assist' and _last_oled_expression == 'assist':
            # Skip redundant assist requests to avoid replaying animation on every loop
            return
        # Remember last expression immediately so rapid repeats don't trigger multiple HTTP calls
        _last_oled_expression = expression
    except Exception:
        pass

    def _set_expression_async():
        try:
            # Call Flask proxy endpoint which routes to ESP32 OLED
            oled_url = f"http://127.0.0.1:5000/oled?value={expression}"
            response = _http_session.get(oled_url, timeout=2)
            # Silently succeed or fail - don't spam console with OLED errors
            # OLED expressions are non-critical visual feedback
        except Exception:
            # Silently ignore OLED failures - they're not critical to functionality
            pass
    
    # Run in background thread so it doesn't block main execution
    threading.Thread(target=_set_expression_async, daemon=True).start()


def show_oled_text(line1: str, line2: str = '', hold_ms: int = 2000):
    """Ask the Flask proxy / OLED to show two-line custom text on the recorder device.
    Non-blocking (spawns a background thread). Uses the same proxy used by set_oled_expression.

    Args:
        line1: First line of text (left-aligned)
        line2: Second line of text (optional)
        hold_ms: How long the device should hold the completed text (milliseconds)
    """
    def _do_show():
        try:
            params = {'value': 'text', 'line1': line1}
            if line2:
                params['line2'] = line2
            params['hold'] = str(int(hold_ms))
            _http_session.get('http://127.0.0.1:5000/oled', params=params, timeout=2)
        except Exception:
            # Silent fail - OLED is non-critical
            pass

    threading.Thread(target=_do_show, daemon=True).start()


def _split_into_two_lines(text: str, max_per_line: int = 20):
    """Split a long text into two roughly balanced lines at a space boundary.
    Keeps words intact. Returns (line1, line2).
    """
    if not text:
        return '', ''
    t = text.strip()
    if len(t) <= max_per_line:
        return t, ''
    # Prefer splitting near the middle
    mid = len(t) // 2
    # Find nearest space to the left of mid, otherwise to the right
    left = t.rfind(' ', 0, mid)
    if left == -1:
        right = t.find(' ', mid)
        if right == -1:
            return t[:max_per_line], t[max_per_line:]
        else:
            left = right
    line1 = t[:left].strip()
    line2 = t[left+1:].strip()
    # If line2 still too long, truncate gracefully
    if len(line2) > max_per_line:
        line2 = line2[:max_per_line]
    return line1, line2

def exercise_all_oled_expressions(delay: float = 0.6):
    """Non-blocking: Cycle through all known OLED expressions for testing.

    - Runs in a background thread.
    - Respects `assist` one-shot behavior by clearing `_last_oled_expression` before sending it.
    - Leaves the display in `idle` when finished.
    """
    try:
        import time

        def _runner():
            try:
                print("[OLED TEST] Exercising OLED expressions...")
                for expr in OLED_VALID_EXPRESSIONS:
                    try:
                        # Allow assist to play by clearing last-expression sentinel
                        if expr == 'assist':
                            try:
                                globals()['_last_oled_expression'] = None
                            except Exception:
                                pass
                        set_oled_expression(expr)
                        time.sleep(float(delay))
                    except Exception:
                        # continue through list even if one fails
                        pass
                # Return to idle at end
                try:
                    set_oled_expression('idle')
                except Exception:
                    pass
                print("[OLED TEST] Done exercising OLED expressions")
            except Exception:
                pass

            threading.Thread(target=_runner, daemon=True).start()
    except Exception:
        pass
def _start_state_processing_animation():
    """Start a simple processing/dots animation on the OLED until stopped.

    Uses module-level `_state_change_stop_event` to control the loop.
    """
    global _state_change_stop_event
    try:
        _state_change_stop_event = threading.Event()
    except Exception:
        return

    def _loop():
        idx = 0
        # Replace three-dot 'processing' animation with repeated 'thinking'
        seq = ['thinking', 'thinking', 'thinking']
        try:
            while not _state_change_stop_event.is_set():
                try:
                    set_oled_expression(seq[idx % len(seq)])
                except Exception:
                    pass
                idx += 1
                time.sleep(0.5)
        finally:
            try:
                set_oled_expression('idle')
            except Exception:
                pass

    threading.Thread(target=_loop, daemon=True).start()

def _start_report_creation_animation():
    """Start an OLED animation indicating report creation (runs until stopped).

    Uses module-level `_report_creation_stop_event` to control the loop.
    The animation cycles through 'processing' and 'excited' with a slow rhythm
    to indicate a longer background task (report generation).
    """
    global _report_creation_stop_event
    try:
        _report_creation_stop_event = threading.Event()
    except Exception:
        return

    def _loop():
        # This loop will read the shared `_report_progress` variable and
        # send progress updates to the ESP OLED via the local Flask proxy.
        # If no explicit progress is provided (value < 0), an indeterminate
        # animation is shown on the device.
        idx = 0
        try:
            while not _report_creation_stop_event.is_set():
                try:
                    # Read the shared progress variable (0-100)
                    prog = int(_report_progress) if isinstance(_report_progress, (int, float)) else 0
                    # If progress is between 0 and 100, report it, otherwise show indeterminate animation
                    if prog >= 0 and prog <= 100:
                        # Deduplicate identical progress updates to avoid flooding the proxy/ESP.
                        try:
                            global _last_report_progress_sent, _last_report_progress_ts
                            nowt = time.time()
                            # Only send if progress changed OR at least 1.2s has passed since last send
                            if _last_report_progress_sent is None or _last_report_progress_sent != prog or (nowt - (_last_report_progress_ts or 0.0)) > 1.2:
                                _last_report_progress_sent = prog
                                _last_report_progress_ts = nowt

                                # Try local Flask proxy first (faster and usually available).
                                url_local = f'http://127.0.0.1:5000/oled?value=report&progress={prog}'
                                try:
                                    print(f"[ReportAnim] Sending progress update to OLED: {prog}% -> {url_local}")
                                    # Use a slightly larger timeout so slow proxies have a chance to respond
                                    resp = _http_session.get(url_local, timeout=(2, 6))
                                    if resp is not None and getattr(resp, 'status_code', None) == 200:
                                        # Successful - reset failure counter
                                        globals()['_report_anim_failures'] = 0
                                    else:
                                        # Non-200 or None: try device endpoints directly (recorder then playback)
                                        tried_dev = False
                                        for host_ip in (globals().get('esp_record_host'), globals().get('esp_playback_host')):
                                            try:
                                                if host_ip:
                                                    url_dev = f'http://{host_ip}/oled?value=report&progress={prog}'
                                                    print(f"[ReportAnim] Attempting direct device oled: {url_dev}")
                                                    rdev = _http_session.get(url_dev, timeout=(2, 4))
                                                    if rdev is not None and getattr(rdev, 'status_code', None) == 200:
                                                        tried_dev = True
                                                        break
                                            except Exception:
                                                continue
                                        if not tried_dev:
                                            # As last resort show fallback animation (no 'processing')
                                            set_oled_expression(['thinking', 'idle', 'excited'][idx % 3])
                                except Exception as e:
                                    # On repeated failures, back off to avoid spamming logs and network
                                    try:
                                        _report_anim_failures = (globals().get('_report_anim_failures') or 0) + 1
                                        globals()['_report_anim_failures'] = _report_anim_failures
                                    except Exception:
                                        globals()['_report_anim_failures'] = 1
                                    if globals().get('_report_anim_failures', 0) <= 3:
                                        print(f"[ReportAnim] Error sending progress update: {e}")
                                    # Provide a visual fallback so OLED isn't blank
                                    set_oled_expression(['thinking', 'idle', 'excited'][idx % 3])
                                    # If failures are frequent, sleep a bit longer to reduce load
                                    if globals().get('_report_anim_failures', 0) > 5:
                                        time.sleep(1.0)
                            else:
                                # Skip sending duplicate progress update
                                pass
                        except Exception:
                            # Keep looping even if dedupe bookkeeping fails
                            pass
                    else:
                        set_oled_expression(['thinking', 'idle', 'excited'][idx % 3])
                except Exception:
                    pass
                idx += 1
                # slower cadence for longer task visibility
                time.sleep(0.6)
        finally:
            try:
                # Ensure progress is shown as complete before we clear
                try:
                    print("[ReportAnim] Finalizing: sending 100% progress to OLED (guaranteed resend sequence)")
                    # Try immediate, +0.5s, +1.5s attempts to increase chance of delivery
                    try:
                        resp = _http_session.get('http://127.0.0.1:5000/oled?value=report&progress=100', timeout=(2, 8))
                        print(f"[ReportAnim] Final OLED proxy response (attempt 1): {getattr(resp, 'status_code', 'N/A')}")
                    except Exception as e1:
                        print(f"[ReportAnim] Final attempt 1 error: {e1}")

                    time.sleep(0.5)
                    try:
                        resp = _http_session.get('http://127.0.0.1:5000/oled?value=report&progress=100', timeout=(2, 10))
                        print(f"[ReportAnim] Final OLED proxy response (attempt 2): {getattr(resp, 'status_code', 'N/A')}")
                    except Exception as e2:
                        print(f"[ReportAnim] Final attempt 2 error: {e2}")

                    time.sleep(1.0)
                    try:
                        resp = _http_session.get('http://127.0.0.1:5000/oled?value=report&progress=100', timeout=(2, 12))
                        print(f"[ReportAnim] Final OLED proxy response (attempt 3): {getattr(resp, 'status_code', 'N/A')}")
                    except Exception as e3:
                        print(f"[ReportAnim] Final attempt 3 error: {e3}")
                except Exception:
                    pass
                set_oled_expression('idle')
            except Exception:
                pass

    threading.Thread(target=_loop, daemon=True).start()

# ============================================================================
# PLEASE WAIT HELPER: Notify user during long processing delays
# ============================================================================
def start_please_wait_timer(delay_seconds: float = 5.0, message: str = "Please wait, I'm still thinking..."):
    """Start a timer that plays a 'please wait' message if processing takes too long.
    
    Returns a threading.Event that can be set() to cancel the timer.
    Use this before slow operations (model calls, database queries, etc.)
    
    Example:
        stop_event = start_please_wait_timer(5.0)
        # ... do slow operation ...
        stop_event.set()  # Cancel timer when done
    """
    stop_event = threading.Event()
    
    def _wait_and_notify():
        # Wait for the delay period (or until cancelled)
        if stop_event.wait(timeout=delay_seconds):
            return  # Cancelled before timeout
        
        # Timeout reached - notify user
        try:
            print(f"[Please Wait] Processing taking longer than {delay_seconds}s - notifying user")
            set_oled_expression('thinking')  # Show thinking animation
            
            # Quick TTS + playback
            try:
                synthesize_speech(message, out_path='wait.wav')
                play_audio('wait.wav')
            except Exception as e:
                print(f"[Please Wait] Could not play message: {e}")
        except Exception as e:
            print(f"[Please Wait] Error in notification: {e}")
    
    threading.Thread(target=_wait_and_notify, daemon=True).start()
    return stop_event

# Directory where repeated/auto-play audio files may be dropped. We'll move processed files
# into a `processed` subfolder after we answer so they don't retrigger.


def esp_upload_file(wav_path: str, endpoint: str = '/upload', max_retries: int = 3, auto_delete_after_playback: bool = True) -> bool:
    """Upload a file to the ESP32 with retry logic and adaptive timeouts.

    Returns True if upload returned HTTP 200. Centralized helper to handle large TTS payloads
    and reduce intermittent failures on the ESP32 by applying retries and longer timeouts for
    large files.
    """
    try:
        filesize = os.path.getsize(wav_path) if os.path.exists(wav_path) else 0
    except Exception:
        filesize = 0

    # Adaptive timeout: larger files need more time to upload and for ESP to write LittleFS
    if filesize > 300_000:
        timeout = (10, 90)  # connect timeout, read timeout
    elif filesize > 100_000:
        timeout = (8, 45)
    else:
        timeout = (5, 30)

    url = f'http://{esp_playback_host}{endpoint}'

    # Pre-upload: ensure remote device has enough free space. If not, try to reclaim by deleting
    # non-critical files via the device /delete endpoint (or /list to choose candidates).
    def _ensure_remote_space(required_bytes, safety_margin=20*1024):
        """Return True if device likely has enough free space, or after attempting cleanup.

        Behavior:
        - Use the device `/list` (or the local Flask proxy `/list`) to estimate used space.
        - Estimate free using a conservative default partition size if exact info is unavailable.
        - Attempt to delete largest non-critical files (via device `/delete` then proxy `/delete`) until enough
          space is reclaimed or candidates exhausted.
        """
        try:
            total_needed = int(required_bytes) + int(safety_margin)

            # Use /list to get files and attempt deletion if needed
            list_urls = [f'http://{esp_playback_host}/list', 'http://127.0.0.1:5000/list']
            files = None
            for lu in list_urls:
                try:
                    lr = requests.get(lu, timeout=(2, 6))
                    if lr.status_code == 200:
                        try:
                            files = lr.json()
                            break
                        except Exception:
                            pass
                except Exception:
                    pass

            # If we couldn't get /list, abort cleanup attempt (can't reason about space)
            if not files:
                print("⚠ Could not retrieve remote file list (/list). Cannot perform cleanup.")
                return False

            # Files expected as list of {name, size}. Compute used and estimate free using a default partition size.
            used = 0
            for fi in files:
                try:
                    used += int(fi.get('size', 0))
                except Exception:
                    pass

            # Conservative default partition size (in bytes). Many ESP32 LittleFS partitions are ~1-2MB.
            DEFAULT_PARTITION_BYTES = 1_500_000
            free_est = max(0, DEFAULT_PARTITION_BYTES - used)

            if free_est >= total_needed:
                print(f"✓ Estimated free ({free_est}) appears sufficient for upload")
                return True

            print(f"⚠ Estimated free={free_est}, need {total_needed}. Attempting to delete large non-critical files...")

            # Prepare candidates: exclude protected files
            protected_prefixes = ['/ca.pem', '/state.txt', '/response.wav']
            candidates = [f for f in files if not any(f.get('name','').startswith(p) for p in protected_prefixes)]
            candidates.sort(key=lambda x: x.get('size', 0), reverse=True)

            # Try deleting until we've reclaimed enough estimated space
            reclaimed = 0
            for fobj in candidates:
                name = fobj.get('name')
                size = int(fobj.get('size', 0)) if fobj.get('size') else 0
                if not name:
                    continue
                print(f"Trying to delete remote file {name} ({size} bytes)")

                deleted_ok = False
                # Try device direct delete then proxy delete
                try_urls = [f'http://{esp_playback_host}/delete', 'http://127.0.0.1:5000/delete']
                for du in try_urls:
                    try:
                        resp = requests.delete(du, params={'name': name}, timeout=(3, 6))
                        if resp.status_code in (200, 204, 202, 404):
                            print(f"✓ Deleted via {du} -> {resp.status_code}")
                            deleted_ok = True
                            break
                        else:
                            # try GET fallback
                            resp2 = requests.get(du, params={'name': name}, timeout=(3, 6))
                            if resp2.status_code in (200, 204, 202, 404):
                                print(f"✓ Deleted (GET) via {du} -> {resp2.status_code}")
                                deleted_ok = True
                                break
                    except Exception:
                        pass

                if deleted_ok:
                    reclaimed += size
                    free_est += size
                    if free_est >= total_needed:
                        print(f"✓ Reclaimed enough space (free_est={free_est})")
                        return True
                else:
                    print(f"⚠ Could not delete {name} via known delete endpoints")

            # If we get here, deletions exhausted and still insufficient
            print(f"⚠ Cleanup attempted but insufficient space remains (free_est={free_est}, need={total_needed})")
            return False
        except Exception as e:
            print(f"⚠ Pre-upload space check failed: {e}")
            return False

    # Attempt remote space check/cleanup once before upload attempts
    try:
        if os.path.exists(wav_path) and os.path.getsize(wav_path) > 0:
            need = os.path.getsize(wav_path)
            if not _ensure_remote_space(need):
                print(f"⚠ Aborting upload: insufficient remote space for {wav_path} ({need} bytes)")
                # Try a compressed MP3 fallback before giving up (smaller bitrate)
                try:
                    if os.path.exists(wav_path):
                        from pydub import AudioSegment as _AudioSegment
                        tmp_mp3 = wav_path + '.mp3'
                        try:
                            audio_seg = _AudioSegment.from_file(wav_path)
                            # Downsample and mono to reduce size
                            audio_seg = audio_seg.set_frame_rate(22050).set_channels(1)
                            audio_seg.export(tmp_mp3, format='mp3', bitrate='96k')
                            mp3_size = os.path.getsize(tmp_mp3)
                            print(f"⚠ Trying compressed MP3 fallback ({mp3_size} bytes)")
                            # Attempt upload of compressed MP3 to device using standard /upload endpoint
                            sd_url = f'http://{esp_playback_host}/upload'
                            try:
                                with open(tmp_mp3, 'rb') as fh:
                                    files = {'file': (os.path.basename(tmp_mp3), fh, 'audio/mpeg')}
                                    r = requests.post(sd_url, files=files, timeout=timeout)
                                if r.status_code == 200:
                                    print(f"✓ ESP upload successful (compressed MP3, endpoint=/upload, size={mp3_size} bytes)")
                                    try:
                                        os.remove(tmp_mp3)
                                    except Exception:
                                        pass
                                    return True
                                else:
                                    print(f"⚠ Compressed MP3 upload returned {r.status_code}")
                            except Exception as _e:
                                print(f"⚠ Compressed MP3 upload attempt failed: {_e}")
                        finally:
                            try:
                                if os.path.exists(tmp_mp3):
                                    os.remove(tmp_mp3)
                            except Exception:
                                pass
                except Exception as _ce:
                    print(f"⚠ Compressed fallback creation failed: {_ce}")

                return False
    except Exception:
        # If pre-check fails for any reason, continue to attempt upload — don't be overly restrictive
        pass

    for attempt in range(1, max_retries + 1):
        try:
            with open(wav_path, 'rb') as fh:
                files = {'file': (os.path.basename(wav_path), fh, 'audio/wav')}
                r = requests.post(url, files=files, timeout=timeout)
            if r.status_code == 200:
                # Give ESP a short moment to kick off playback (non-blocking for most flows)
                time.sleep(0.35)
                print(f"✓ ESP upload successful (endpoint={endpoint}, size={filesize} bytes, attempt={attempt})")

                # Optionally schedule a background delete AFTER estimated playback finishes.
                # We estimate playback duration for 16kHz/16bit/mono WAV as ~32000 bytes/sec.
                if auto_delete_after_playback and filesize and filesize > 0:
                    try:
                        def _schedule_delete(local_file_path, basename, size_bytes):
                            try:
                                # Attempt to compute accurate playback duration from WAV header
                                estimated_sec = None
                                try:
                                    import wave as _wave
                                    if local_file_path and os.path.exists(local_file_path):
                                        try:
                                            with _wave.open(local_file_path, 'rb') as wf:
                                                fr = wf.getframerate()
                                                nch = wf.getnchannels()
                                                sampw = wf.getsampwidth()
                                                bytes_per_sec = fr * nch * sampw
                                                if bytes_per_sec > 0:
                                                    estimated_sec = float(size_bytes) / float(bytes_per_sec)
                                        except Exception:
                                            estimated_sec = None
                                except Exception:
                                    estimated_sec = None

                                # Fallback conservative estimate if WAV header unavailable
                                if estimated_sec is None:
                                    estimated_sec = max(0.5, float(size_bytes) / 32000.0)
                                # Add small buffer to allow playback and filesystem write completion
                                wait = estimated_sec + 0.8
                                print(f"⏳ Scheduling remote delete for '{basename}' in {wait:.1f}s (estimated duration {estimated_sec:.2f}s)")
                                time.sleep(wait)

                                # Candidate delete endpoints to try (GET or DELETE)
                                candidates = [
                                    f'http://{esp_playback_host}/delete?name={basename}',
                                    f'http://{esp_playback_host}/remove?name={basename}',
                                    f'http://{esp_playback_host}/rm?name={basename}',
                                    f'http://{esp_playback_host}/delete_file?name={basename}',
                                    f'http://{esp_playback_host}/remove_file?name={basename}',
                                ]

                                # Also try Flask proxy delete route (if running locally)
                                try:
                                    candidates.append(f'http://127.0.0.1:5000/delete?name={basename}')
                                    candidates.append(f'http://127.0.0.1:5000/remove?name={basename}')
                                except Exception:
                                    pass

                                deleted = False
                                for c in candidates:
                                    try:
                                        # Try HTTP DELETE first
                                        resp = requests.delete(c, timeout=(3, 8))
                                        code = getattr(resp, 'status_code', None)
                                        if code in (200, 204, 202, 404):
                                            # If response.wav was used, 404 likely means the firmware already removed
                                            if basename.lower() == 'response.wav' and code == 404:
                                                print(f"✓ Attempted delete via DELETE {c} -> {code} (already removed on device - expected)")
                                            else:
                                                print(f"✓ Attempted delete via DELETE {c} -> {code}")
                                            deleted = True
                                            break
                                    except Exception:
                                        # Try GET/POST fallbacks
                                        try:
                                            resp2 = requests.get(c, timeout=(3, 8))
                                            code2 = getattr(resp2, 'status_code', None)
                                            if code2 in (200, 204, 202, 404):
                                                if basename.lower() == 'response.wav' and code2 == 404:
                                                    print(f"✓ Attempted delete via GET {c} -> {code2} (already removed on device - expected)")
                                                else:
                                                    print(f"✓ Attempted delete via GET {c} -> {code2}")
                                                deleted = True
                                                break
                                        except Exception:
                                            pass

                                if not deleted:
                                    print(f"⚠ Could not delete remote file '{basename}' - device may not expose delete endpoint")
                            except Exception as exc:
                                print(f"⚠ Background delete thread error: {exc}")

                        # Launch background thread (non-blocking)
                        try:
                            basename = os.path.basename(wav_path)
                            threading.Thread(target=_schedule_delete, args=(wav_path, basename, filesize), daemon=True).start()
                        except Exception:
                            pass
                    except Exception:
                        pass

                return True
            else:
                print(f"⚠ ESP upload returned {r.status_code} (endpoint={endpoint}, attempt={attempt})")
        except requests.exceptions.Timeout as e:
            print(f"⚠ ESP upload timeout (attempt {attempt}): {e}")
        except requests.exceptions.ConnectionError as e:
            print(f"⚠ ESP upload connection error (attempt {attempt}): {e}")
        except Exception as e:
            print(f"⚠ ESP upload failed (attempt {attempt}): {e}")

        # Backoff before retrying
        if attempt < max_retries:
            time.sleep(0.6 * attempt)

    return False


# Directory where repeated/auto-play audio files may be dropped. We'll move processed files
# into a `processed` subfolder after we answer so they don't retrigger.
REPEAT_AUDIO_DIR = getattr(config, 'REPEAT_AUDIO_DIR', r"C:\Users\vonti\OneDrive\Desktop\GENTA SYS\RepeatAudio")

# ESP32 Recorder (GENTA.ino - Recording device)
# Select whether to fetch recordings via the local proxy (Flask/ngrok) or directly from the ESP.
# If GENTA_USE_PROXY is set to 1/true/yes we will download from `audio_raw_url` (default: http://localhost:5000/recording.wav)
# Otherwise we will contact the ESP directly at GENTA_ESP_RECORD_IP.
USE_PROXY = os.environ.get('GENTA_USE_PROXY', '1').lower() in ('1', 'true', 'yes')
if USE_PROXY:
    esp_record_host = None
    print("ESP32 Recorder: Using Flask proxy for recording downloads")
    # When using the local Flask proxy, ensure the canonical download URL
    # points at the local proxy endpoint so record_and_transcribe downloads
    # from `http://127.0.0.1:5000/recording.wav` by default.
    try:
        audio_raw_url = getattr(config, 'AUDIO_RAW_URL', None) or 'http://127.0.0.1:5000/recording.wav'
        # Also ensure audio_raw_url is local if config explicitly requested proxy usage
        if not str(audio_raw_url).startswith('http://127.0.0.1') and not str(audio_raw_url).startswith('http://localhost'):
            audio_raw_url = 'http://127.0.0.1:5000/recording.wav'
        print(f"[Config] Using proxy audio URL: {audio_raw_url}")
        # If discovery found a recorder IP, still remember it so we can prefer
        # direct-device downloads when available while keeping proxy as default.
        try:
            if _discovered and isinstance(_discovered, dict) and 'recorder' in _discovered:
                esp_record_host = _discovered.get('recorder')
                print(f"[Config] Also discovered recorder at {esp_record_host}; will prefer direct device when reachable")
        except Exception:
            pass
    except Exception:
        audio_raw_url = 'http://127.0.0.1:5000/recording.wav'
else:
    # Try auto-discovery for recorder
    if _discovered and 'recorder' in _discovered:
        esp_record_host = _discovered['recorder']
        print(f"✓ ESP32 Recorder: {esp_record_host} (auto-discovered)")
    else:
        esp_record_host = os.environ.get('GENTA_ESP_RECORD_IP', getattr(config, 'ESP_RECORD_IP', '192.168.50.62'))
        if 'GENTA_ESP_RECORD_IP' in os.environ:
            print(f"ESP32 Recorder: {esp_record_host} (from environment variable)")
        else:
            print(f"⚠ ESP32 Recorder: {esp_record_host} (default - auto-discovery failed)")
            print("  Set GENTA_ESP_RECORD_IP environment variable or start Flask server")

print("="*70 + "\n")

# Centralized config-driven endpoints (fall back to existing hardcoded test URLs if config missing)
BAKurl_state = getattr(config, 'BAKURL_STATE', "https://nonbasic-bob-inimical.ngrok-free.dev/download/state.txt")
BAKaudio_raw_url = getattr(config, 'BAK_AUDIO_RAW_URL', "https://nonbasic-bob-inimical.ngrok-free.dev/download/recording.wav")
url_state = getattr(config, 'URL_STATE', "https://nonbasic-bob-inimical.ngrok-free.dev/state.txt")
audio_raw_url = getattr(config, 'AUDIO_RAW_URL', "https://nonbasic-bob-inimical.ngrok-free.dev/download_recording")
student_id_url = getattr(config, 'STUDENT_ID_URL', "https://nonbasic-bob-inimical.ngrok-free.dev/student_id.txt")

# Ensure GOOGLE_APPLICATION_CREDENTIALS from centralized config is applied if available
if getattr(config, 'GOOGLE_APPLICATION_CREDENTIALS', None):
    if not os.environ.get('GOOGLE_APPLICATION_CREDENTIALS'):
        logging.info("Setting GOOGLE_APPLICATION_CREDENTIALS from config: %s", config.GOOGLE_APPLICATION_CREDENTIALS)
        os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = config.GOOGLE_APPLICATION_CREDENTIALS
#    else:
#        logging.info("GOOGLE_APPLICATION_CREDENTIALS already set in environment")
# Configure genai from environment variable or centralized config (do not hardcode keys)
# Priority: environment `GENAI_API_KEY` > `config.GENAI_API_KEY`
_genai_key = os.environ.get('GENAI_API_KEY') or getattr(config, 'GENAI_API_KEY', None)
if _genai_key:
    try:
        genai.configure(api_key=_genai_key)
        if os.environ.get('GENAI_API_KEY'):
            logging.info("genai configured from environment variable GENAI_API_KEY")
        else:
            logging.info("genai configured from config.GENAI_API_KEY")
    except Exception as e:
        logging.error("Failed to configure genai client: %s", e)
else:
    logging.warning("GENAI_API_KEY not set in env or config; genai.configure skipped. Set GENAI_API_KEY env var or update config.py to enable model access.")
# Resilient genai.chat caller: normalizes model name, message shape, and maps/removes
# legacy kwargs so this code works across multiple installed google.generativeai versions.
def _call_genai_chat(model_name, messages, kwargs=None):
    import inspect
    if kwargs is None:
        kwargs = {}

    # Inspect supported parameters for genai.chat
    try:
        sig = inspect.signature(genai.chat)
        supported = set(sig.parameters.keys())
    except Exception:
        supported = set()

    # Normalize model name to expected prefix
    mdl = model_name
    try:
        if isinstance(mdl, str) and not (mdl.startswith('models/') or mdl.startswith('tunedModels/')):
            mdl = f"models/{mdl}"
    except Exception:
        pass

    # Prepare joined text from messages if needed
    try:
        if isinstance(messages, list):
            if all(isinstance(x, str) for x in messages):
                joined_text = '\n'.join(messages)
            else:
                # try to stringify elements
                joined_text = '\n'.join(str(x) for x in messages)
        else:
            joined_text = str(messages)
    except Exception:
        joined_text = ''

    # Decide whether to pass 'messages' or 'prompt' based on signature
    call_kwargs = {}
    if 'model' in supported:
        call_kwargs['model'] = mdl
    else:
        call_kwargs['model'] = mdl

    if 'messages' in supported:
        # Build a simple structured message if caller passed strings
        if isinstance(messages, list) and messages and all(isinstance(x, dict) for x in messages):
            call_kwargs['messages'] = messages
        else:
            call_kwargs['messages'] = [{'author': 'user', 'content': [{'type': 'text', 'text': joined_text}]}]
    else:
        # Use prompt/input fallbacks
        if 'prompt' in supported:
            call_kwargs['prompt'] = joined_text
        elif 'input' in supported:
            call_kwargs['input'] = joined_text
        else:
            call_kwargs['prompt'] = joined_text

    # Map legacy generation kwargs to supported names where possible
    # Prefer modern parameter names first. Do not inject unknown params as a fallback.
    legacy_map = {
        'max_output_tokens': ['max_output_tokens', 'max_new_tokens', 'max_tokens'],
        'candidate_count': ['candidate_count', 'candidateCount', 'n'],
        'temperature': ['temperature'],
        'top_p': ['top_p', 'topP'],
        'top_k': ['top_k', 'topK']
    }

    for src, targets in legacy_map.items():
        if src in kwargs:
            val = kwargs[src]
            placed = False
            for t in targets:
                if t in supported:
                    call_kwargs[t] = val
                    placed = True
                    logging.debug("Mapping legacy kw '%s' -> '%s'", src, t)
                    break
            # If none of the candidate target names are supported, do not add an unsupported
            # keyword to call_kwargs. We'll let the call proceed without this parameter.
            if not placed:
                logging.debug("Legacy kw '%s' could not be mapped to supported chat params; skipping", src)

    # pass through simple safe kwargs
    for k in ('temperature',):
        if k in kwargs and k not in call_kwargs:
            call_kwargs[k] = kwargs[k]

    # Final attempt: try function-based chat API if present
    try:
        if hasattr(genai, 'chat') and callable(getattr(genai, 'chat')):
            logging.debug("Calling genai.chat with keys: %s", list(call_kwargs.keys()))
            # Attempt call; if TypeError arises due to unexpected kwarg, strip it and retry once.
            try:
                return genai.chat(**call_kwargs)
            except TypeError as te:
                msg = str(te)
                logging.debug("genai.chat TypeError: %s", msg)
                m = None
                try:
                    # Parse unexpected keyword name from message like "chat() got an unexpected keyword argument 'max_tokens'"
                    import re as _re
                    m = _re.search(r"unexpected keyword argument '?([a-zA-Z0-9_]+)'?", msg)
                except Exception:
                    m = None
                if m:
                    bad_key = m.group(1)
                    if bad_key in call_kwargs:
                        logging.info("genai.chat rejects kw '%s' — retrying without it", bad_key)
                        call_kwargs.pop(bad_key, None)
                        try:
                            return genai.chat(**call_kwargs)
                        except Exception:
                            logging.exception("genai.chat retry after removing '%s' failed", bad_key)
                # If we couldn't parse or retry failed, log and continue to other fallbacks
                logging.exception("genai.chat final call failed for model '%s': %s", mdl, te)
    except Exception as e:
        logging.exception("genai.chat final call failed for model '%s': %s", mdl, e)

    # If module does not expose chat/generate_text functions (older/newer SDK),
    # try class-based API provided by the installed module (genai.GenerativeModel)
    try:
        if hasattr(genai, 'GenerativeModel') and callable(getattr(genai, 'GenerativeModel')):
            try:
                # Instantiate provider's GenerativeModel and use its methods
                prov_model_name = mdl if isinstance(mdl, str) else mdl
                try:
                    prov = genai.GenerativeModel(prov_model_name)
                except Exception:
                    # try short name
                    short = prov_model_name.replace('models/', '', 1) if isinstance(prov_model_name, str) else prov_model_name
                    prov = genai.GenerativeModel(short)

                # If generate_content exists, use it (synchronous text generation)
                joined = joined_text if 'joined_text' in locals() else ''
                if hasattr(prov, 'generate_content'):
                    logging.debug("Delegating to prov.generate_content for model '%s'", prov_model_name)
                    # Provider expects a 'generation_config' parameter rather than flat kwargs
                    gen_cfg = {}
                    # Map common legacy kwargs into provider generation_config fields
                    if isinstance(kwargs, dict):
                        if 'max_output_tokens' in kwargs:
                            gen_cfg['max_output_tokens'] = kwargs.get('max_output_tokens')
                        if 'candidate_count' in kwargs:
                            gen_cfg['candidate_count'] = kwargs.get('candidate_count')
                        if 'temperature' in kwargs:
                            gen_cfg['temperature'] = kwargs.get('temperature')
                        if 'top_p' in kwargs:
                            gen_cfg['top_p'] = kwargs.get('top_p')
                        if 'top_k' in kwargs:
                            gen_cfg['top_k'] = kwargs.get('top_k')
                    # Call provider with a generation_config mapping
                    try:
                        return prov.generate_content(joined, generation_config=gen_cfg if gen_cfg else None)
                    except TypeError as te:
                        logging.exception("Provider generate_content TypeError for model '%s': %s", prov_model_name, te)
                        # If provider rejects the mapping, try call without generation_config
                        try:
                            return prov.generate_content(joined)
                        except Exception:
                            raise

                # If start_chat exists, start a session and send a message
                if hasattr(prov, 'start_chat'):
                    logging.debug("Delegating to prov.start_chat/send_message for model '%s'", prov_model_name)
                    sess = prov.start_chat(history=None)
                    if hasattr(sess, 'send_message'):
                        return sess.send_message(joined)
            except Exception as e:
                logging.exception("Provider GenerativeModel usage failed for '%s': %s", mdl, e)

    except Exception:
        pass

    # As last resort, try minimal function-based calls if generate_text exists
    try:
        if hasattr(genai, 'generate_text') and callable(getattr(genai, 'generate_text')):
            logging.debug("Calling genai.generate_text as fallback")
            return genai.generate_text(model=mdl, prompt=joined_text if 'joined_text' in locals() else '', **{k: v for k, v in kwargs.items() if isinstance(v, (str, int, float))})
    except Exception as e:
        logging.exception("generate_text fallback failed for model '%s': %s", mdl, e)

    # If we get here, no provider API worked
    raise RuntimeError("No supported genai API available on this environment (neither chat/generate_text nor GenerativeModel)")

# Compatibility shim: provide a GenerativeModel class that adapts the
# installed function-based `google.generativeai` API to the older
# class-based usage expected by this code (generate_content, start_chat).
class GenerativeModel:
    def __init__(self, model_name="models/gemini-2.5-flash"):
        # Normalize model name to expected prefix used by installed genai versions
        try:
            if isinstance(model_name, str) and not (model_name.startswith('models/') or model_name.startswith('tunedModels/')):
                self.model_name = f"models/{model_name}"
            else:
                self.model_name = model_name
        except Exception:
            self.model_name = model_name

        # If the installed google.generativeai module provides its own GenerativeModel
        # implementation, prefer delegating to it to maintain compatibility across
        # differing SDK versions (some expose a class-based API rather than module-level
        # chat/generate_text functions).
        self._inner = None
        try:
            if hasattr(genai, 'GenerativeModel') and callable(getattr(genai, 'GenerativeModel')):
                try:
                    # Try to instantiate the provider's wrapper using the normalized name
                    self._inner = genai.GenerativeModel(self.model_name)
                    logging.debug("Using installed genai.GenerativeModel for model '%s'", self.model_name)
                except Exception:
                    # Try without the 'models/' prefix if instantiation fails
                    try:
                        short = self.model_name.replace('models/', '', 1) if isinstance(self.model_name, str) else self.model_name
                        self._inner = genai.GenerativeModel(short)
                        logging.debug("Using installed genai.GenerativeModel (short name) for model '%s'", short)
                    except Exception:
                        self._inner = None
        except Exception:
            self._inner = None

    def generate_content(self, prompt, max_output_tokens=None, temperature=None, candidate_count=None, **kwargs):
        """
        Map legacy generate_content(...) calls to genai.generate_text(...).
        Returns the Completion-like object produced by genai.generate_text.
        """
        # If we have a provider-supplied inner model, attempt to delegate to its generate_content.
        # Newer provider implementations expect a `generation_config` dict rather than flat kwargs.
        generation_config = {}
        if max_output_tokens is not None:
            generation_config['max_output_tokens'] = max_output_tokens
        if temperature is not None:
            generation_config['temperature'] = temperature
        if candidate_count is not None:
            generation_config['candidate_count'] = candidate_count
        # Merge any known compatible keys from kwargs
        for k in ('top_p', 'top_k'):
            if k in kwargs:
                generation_config[k] = kwargs.pop(k)

        # Try provider inner instance first (class-based API)
        if self._inner is not None and hasattr(self._inner, 'generate_content'):
            try:
                # Preferred calling style for provider: contents + generation_config keyword
                try:
                    return self._inner.generate_content(prompt, generation_config=generation_config or None, **kwargs)
                except TypeError:
                    # Some older provider wrappers accepted flat kwargs; try that as a fallback
                    return self._inner.generate_content(prompt, max_output_tokens=max_output_tokens, temperature=temperature, candidate_count=candidate_count, **kwargs)
            except Exception:
                # If provider call fails, fall through to function-based fallback below
                pass

        # Fallback: try to call function-based API if available on module
        try:
            return genai.generate_text(
                model=self.model_name,
                prompt=prompt,
                max_output_tokens=max_output_tokens,
                temperature=temperature,
                candidate_count=candidate_count,
                **kwargs,
            )
        except Exception:
            # Re-raise to let callers handle errors; higher-level code will attempt other fallbacks
            raise

    def start_chat(self, history=None):
        """
        Create a simple chat session object that exposes send_message(...)
        which maps to genai.chat(...). The history argument is expected to be
        a list of dicts with 'role' and 'parts' (the existing code passes that
        form). We'll flatten parts -> strings and maintain an internal
        message list for the session.
        """
        messages = []
        if history:
            for item in history:
                txt = None
                try:
                    parts = item.get('parts') if isinstance(item, dict) else None
                    if parts:
                        # join part texts
                        txt = ''.join(
                            (p.get('text') if isinstance(p, dict) else str(p)) for p in parts
                        )
                    else:
                        if isinstance(item, dict) and 'text' in item:
                            txt = item['text']
                        else:
                            txt = str(item)
                except Exception:
                    try:
                        txt = str(item)
                    except Exception:
                        txt = ''
                messages.append(txt)

        # If provider has a native GenerativeModel and it supports start_chat, delegate to it
        try:
            if self._inner is not None and hasattr(self._inner, 'start_chat'):
                try:
                    # Many provider start_chat implementations return a session-like object.
                    sess = self._inner.start_chat(history=history)
                    return sess
                except Exception:
                    pass
        except Exception:
            pass

        return _ChatSession(self.model_name, messages)


class _ChatSession:
    def __init__(self, model_name, messages=None):
        self.model_name = model_name
        self.history = list(messages) if messages else []

    def send_message(self, prompt_text, generation_config=None):
        """
        Send a message in the session. Maps to genai.chat and returns the
        ChatResponse from the installed library. generation_config is a dict
        that may include keys like max_output_tokens, temperature, candidate_count, top_p, top_k.
        """
        msgs = list(self.history)
        msgs.append(prompt_text)

        # Map generation_config keys into genai.chat kwargs
        kwargs = {}
        if generation_config:
            if isinstance(generation_config, dict):
                if 'max_output_tokens' in generation_config:
                    kwargs['max_output_tokens'] = generation_config['max_output_tokens']
                if 'temperature' in generation_config:
                    kwargs['temperature'] = generation_config['temperature']
                if 'candidate_count' in generation_config:
                    kwargs['candidate_count'] = generation_config['candidate_count']
                if 'top_p' in generation_config:
                    kwargs['top_p'] = generation_config['top_p']
                if 'top_k' in generation_config:
                    kwargs['top_k'] = generation_config['top_k']

        # Use tolerant wrapper to handle signature differences across genai versions
        try:
            resp = _call_genai_chat(self.model_name, msgs, kwargs)
        except Exception as e:
            logging.debug("_call_genai_chat failed: %s", e)
            # Final fallback: try calling genai.chat directly without extra kwargs
            try:
                resp = genai.chat(model=self.model_name, messages=msgs)
            except Exception:
                raise

        # Update history with prompt and last reply (if any)
        try:
            self.history.append(prompt_text)
            last = resp.last if hasattr(resp, 'last') else None
            if last:
                self.history.append(last)
        except Exception:
            pass

        return resp


# Instantiate wrapper so existing code using `model = genai.GenerativeModel(...)`
# continues to work. Allow override via environment or config for flexibility.
_default_model = os.environ.get('GENAI_MODEL') or getattr(config, 'GENAI_MODEL', 'gemini-2.5-flash')
try:
    model = GenerativeModel(model_name=_default_model)
    logging.info("Generative model initialized as: %s", _default_model)
except Exception as e:
    logging.error("Failed to initialize GenerativeModel wrapper: %s", e)
    # Fallback to a raw model name (will be normalized by GenerativeModel)
    model = GenerativeModel(model_name='gemini-2.5-flash')

# Flag indicating whether GenAI model is reachable and usable
_GENAI_AVAILABLE = True

def verify_genai_model(timeout: float = 6.0) -> bool:
    """Quick check to ensure the configured genai model is reachable.

    Returns True if a lightweight call to the GenAI API succeeds, False otherwise.
    This helps avoid starting long report generation when the model is unavailable.
    """
    global _GENAI_AVAILABLE, _default_model
    tried = []
    try:
        logging.info("Probing GenAI model '%s'...", _default_model)
        # Try a few model-name variants to handle differences between genai versions
        variants = []
        variants.append(_default_model)
        if isinstance(_default_model, str):
            if not _default_model.startswith('models/'):
                variants.append(f"models/{_default_model}")
            else:
                # also try without the models/ prefix
                variants.append(_default_model.replace('models/', '', 1))

        # Deduplicate while preserving order
        seen = set()
        probes = []
        for v in variants:
            if v and v not in seen:
                seen.add(v)
                probes.append(v)

        for probe in probes:
            try:
                logging.info("Trying probe via _call_genai_chat with model='%s'", probe)
                resp = _call_genai_chat(probe, ['ping'], kwargs={'temperature': 0.0, 'max_output_tokens': 1})
                logging.info("GenAI probe succeeded for model '%s' (via chat). Response: %s", probe, getattr(resp, 'last', resp))
                _GENAI_AVAILABLE = True
                # If wrapper exists, prefer storing normalized name
                try:
                    if hasattr(model, 'model_name'):
                        model.model_name = probe if probe.startswith('models/') else f"models/{probe}"
                    # If the installed provider exposes a GenerativeModel class, instantiate and attach it
                    try:
                        if hasattr(genai, 'GenerativeModel') and getattr(genai, 'GenerativeModel'):
                            try:
                                prov_inst = genai.GenerativeModel(probe if probe.startswith('models/') else f"models/{probe}")
                                # attach to our wrapper instance for future delegated calls
                                if hasattr(model, '_inner'):
                                    model._inner = prov_inst
                                    logging.info("Attached provider GenerativeModel instance to wrapper for model '%s'", probe)
                            except Exception:
                                pass
                    except Exception:
                        pass
                except Exception:
                    pass
                return True
            except Exception as e:
                logging.exception("Probe via chat failed for model '%s': %s", probe, e)
                tried.append((probe, str(e)))

        # Try generate_text fallback with same variants
        for probe in probes:
            try:
                probe_name = probe if probe.startswith('models/') else f"models/{probe}"
                logging.info("Trying generate_text fallback for model '%s'", probe_name)
                gresp = genai.generate_text(model=probe_name, prompt='ping', max_output_tokens=1)
                logging.info("generate_text fallback succeeded for model '%s'", probe_name)
                _GENAI_AVAILABLE = True
                return True
            except Exception as e:
                logging.exception("generate_text fallback failed for '%s': %s", probe, e)
                tried.append((probe, str(e)))

        logging.warning("All GenAI probes failed. Tried variants: %s", tried)
    except Exception as final_e:
        logging.exception("verify_genai_model unexpected error: %s", final_e)

    _GENAI_AVAILABLE = False
    logging.warning("GenAI model not reachable: %s", _default_model)
    return False

# Immediately verify model availability on startup so callers can skip work early
try:
    verify_genai_model()
except Exception:
    pass


ctime = datetime.now()
ftime = str(ctime.strftime("%Y-%m-%d %H:%M:%S"))
# Centralized upload directory from config (config already ensures existence)
UPLOAD_DIR = getattr(config, 'UPLOAD_DIR', os.path.abspath(r"C:\Users\vonti\OneDrive\Desktop\GENTA SYS\MAIN_SYSTEM\uploads"))

try:
    os.makedirs(UPLOAD_DIR, exist_ok=True)
except Exception:
    pass

def U(path_tail: str) -> str:
    """Helper: path inside canonical upload directory."""
    return os.path.join(UPLOAD_DIR, path_tail) if path_tail else UPLOAD_DIR

file_path = U('Math1.txt')
conversation_file_path = U('conversation_log.txt')
output_docx_path = U('analysis_result.docx')
output_docx_tailoredmodule_path = U('tailored_module.docx')
audio_raw_path = U('Recording.wav')
audio_converted_path = U('Recording_Converted.wav')
audio_mono_path = U('Recording_Mono.wav')

# Current session student/teacher context (set at GENTA startup after LRN prompt)
CURRENT_STUDENT_ID = None
CURRENT_TEACHER_ID = None
CURRENT_TEACHER_NAME = None
CURRENT_STUDENT_NAME = None

# Animation stop events (module-level) - controlled by background threads
_state_change_stop_event = None
_quiz_loading_stop_event = None
_report_creation_stop_event = None
# Last recording diagnostics: bytes seen in most recent download and whether audio was present
LAST_RECORDING_BYTES = 0
LAST_AUDIO_PRESENT = False
# Threshold (bytes) to consider a recording "present" (avoid tiny placeholders)
AUDIO_PRESENT_THRESHOLD = 2000
# Flag to indicate report creation is in progress. When True, suppress state-change
# confirmations (so the system doesn't interrupt report generation with prompts).
_REPORT_CREATION_ACTIVE = False
# Flag to indicate the whole QUIZ flow is active (after student confirmation).
# When True, state monitoring should avoid fetching state so it doesn't interfere.
_QUIZ_FLOW_ACTIVE = False
# Progress for report creation (0-100)
_report_progress = 0
# Last-sent report progress (to avoid spamming identical updates)
_last_report_progress_sent = None
# Timestamp when last report progress was sent
_last_report_progress_ts = 0.0
# Last expression sent to OLED proxy (used to avoid replaying one-shot animations)
_last_oled_expression = None

# Global flag: True after startup cleanup completes (prevents re-clearing during session)
_STARTUP_CLEANUP_DONE = False

# State monitoring flags
_STATE_CHANGE_REQUESTED = False
_STATE_MONITOR_ACTIVE = False
_CURRENT_STATE = "0"

def record_and_transcribe(esp_host: str = None, audio_raw_path_local: str = None, audio_mono_local: str = None,
                          poll_for_recording: bool = True, max_poll_seconds: int = 30, use_english: bool = False,
                          use_event_driven: bool = True) -> tuple:
    """Event-driven recording detection and transcription.
    
    Uses Flask proxy /wait_for_recording endpoint for fast, efficient recording detection.
    No polling - just waits for notification that recording is ready, then downloads and transcribes.
    
    Returns: (transcript_text, forced_timeout)
    """
    if audio_raw_path_local is None:
        audio_raw_path_local = audio_raw_path
    if audio_mono_local is None:
        audio_mono_local = audio_mono_path

    forced_timeout = False

    # If a state change was requested (e.g., user pressed quiz button) we should
    # abort any ongoing recording waits so the main loop can proceed to the
    # confirmation flow. This is a fast pre-check to avoid initiating network
    # calls unnecessarily when the system is transitioning state.
    try:
        if globals().get('_STATE_CHANGE_REQUESTED'):
            print('[record_and_transcribe] Aborting due to pending state change request')
            return "", False
    except Exception:
        pass

    # Determine download URL
    if esp_host:
        download_url = f'http://{esp_host}/recording.wav'
    else:
        download_url = audio_raw_url

    print("record_and_transcribe: Waiting for recording...")

    # ===== SMART CHECK: If recording already exists, use it immediately =====
    # This handles cases where Flask monitor missed the notification or old recording exists
    try:
        # Quick check if file exists and has content
        check_start = time.time()
        size_check_url = f'http://{esp_host}/size' if esp_host else None
        
        if size_check_url:
            try:
                size_resp = _http_session.get(size_check_url, timeout=2)
                if size_resp.status_code == 200:
                    current_size = int(size_resp.text.strip())
                    if current_size > 1024:  # Recording exists
                        print(f"[Smart Check] Recording already exists ({current_size} bytes) - downloading immediately")
                        # Skip event-driven wait, go directly to download
                        use_event_driven = False
            except Exception:
                pass
    except Exception:
        pass

    # ===== EVENT-DRIVEN RECORDING DETECTION (if needed) =====
    if use_event_driven and poll_for_recording:
        # FAST PRE-CHECK: Ask the local proxy if a recording already appears ready
        # This avoids waiting the full long-poll timeout when the recorder is already
        # finished writing a file (helps when monitor missed a notification).
        try:
            probe = _http_session.get('http://127.0.0.1:5000/poll_recording_status', timeout=1.2)
            if probe is not None and getattr(probe, 'status_code', None) == 200:
                try:
                    pj = probe.json()
                    if pj.get('status') == 'ready' and int(pj.get('size', 0)) > 1024:
                        print("[Event-Driven] Fast poll: recording already present (skipping long-poll)")
                        # proceed to download section without waiting for notification
                        pass
                    else:
                        # No ready recording yet; continue to long-poll
                        print("[Event-Driven] No ready recording from fast poll; waiting for proxy notification...")
                except Exception:
                    print("[Event-Driven] Fast poll returned unexpected payload; continuing to long-poll")
            else:
                print("[Event-Driven] Fast poll failed or non-200; waiting for proxy notification...")
        except Exception:
            # If fast poll fails (proxy down/unreachable), fall back to long-poll behavior
            print("[Event-Driven] Fast poll to proxy failed; falling back to long-poll")

        # Wait for Flask proxy notification - it monitors ESP32 and notifies when recording is ready
        print("[Event-Driven] Waiting for recording notification from Flask proxy...")
        proxy_url = 'http://127.0.0.1:5000/wait_for_recording'
        wait_start = time.time()
        
        try:
            # Long-polling request to Flask - it will block until recording is ready
            response = requests.get(
                proxy_url,
                params={'timeout': max_poll_seconds},
                timeout=max_poll_seconds + 5  # Client timeout slightly longer than server
            )
            
            if response.status_code == 200:
                data = response.json()
                # Support cancellation notification from Flask proxy: abort the
                # recording flow and return so the main loop can proceed to
                # the state-change confirmation flow.
                try:
                    if isinstance(data, dict) and data.get('status') == 'cancelled':
                        print('[Event-Driven] Recording wait cancelled by state change')
                        forced_timeout = False
                        return "", forced_timeout
                except Exception:
                    pass
                elapsed = time.time() - wait_start
                print(f"[Event-Driven] ✓ Recording ready after {elapsed:.1f}s!")
                # Recording is ready - proceed to download
            elif response.status_code == 408:
                # Timeout from Flask - BUT recording might still exist! Check before giving up
                print(f"[Event-Driven] Timeout from Flask, checking if recording exists anyway...")
                try:
                    if esp_host:
                        size_resp = _http_session.get(f'http://{esp_host}/size', timeout=2)
                        if size_resp.status_code == 200:
                            current_size = int(size_resp.text.strip())
                            if current_size > 1024:
                                print(f"[Event-Driven] Recording EXISTS despite timeout ({current_size} bytes) - continuing to download")
                                # Continue to download section
                            else:
                                print(f"[Event-Driven] No recording found (size={current_size})")
                                forced_timeout = True
                                return "", forced_timeout
                    else:
                        # Try downloading from tunnel anyway
                        print(f"[Event-Driven] Trying direct download from tunnel despite timeout...")
                        # Continue to download section
                except Exception as e:
                    print(f"[Event-Driven] Size check failed: {e}")
                    forced_timeout = True
                    return "", forced_timeout
            else:
                print(f"[Event-Driven] Unexpected response: {response.status_code}")
                forced_timeout = True
                return "", forced_timeout
        except requests.Timeout:
            print(f"[Event-Driven] Connection timeout after {max_poll_seconds}s")
            # BUT recording might still exist! Try to download anyway
            print("[Event-Driven] Attempting download despite connection timeout...")
        except Exception as e:
            print(f"[Event-Driven] Error: {e}")
            # Try to download anyway
            print("[Event-Driven] Attempting download despite error...")
    
    # ===== DOWNLOAD RECORDING =====
    # Recording is ready, now download it. Try multiple candidate URLs in order
    print(f"Attempting to download recording (candidates will be tried)")
    ffmpeg_path = _get_ffmpeg_path()
    converted_bytes = None
    try:
        # Show thinking animation
        set_oled_expression('thinking')

        print(f"⏱ Starting download at {time.strftime('%H:%M:%S')}...")
        download_start = time.time()

        # Prepare candidate URLs: prefer direct esp_host, then local proxy, then configured download_url
        candidates = []
        try:
            if esp_host:
                candidates.append(f'http://{esp_host}/recording.wav')
        except Exception:
            pass
        # Local proxy
        candidates.append('http://127.0.0.1:5000/recording.wav')
        # Configured download_url (may be tunnel/ngrok)
        try:
            candidates.append(download_url)
        except Exception:
            pass

        r = None
        cl = 'unknown'
        written = 0
        for cand in candidates:
            try:
                print(f"[Download] Trying {cand} ...")
                r = _http_session.get(cand, stream=True, timeout=(3, 20))
                if r is not None and getattr(r, 'status_code', None) == 200:
                    cl = r.headers.get('content-length', 'unknown')
                    print(f"[Download] Connected to {cand}, downloading {cl} bytes...")
                    # Save to local file
                    try:
                        os.makedirs(os.path.dirname(audio_raw_path_local), exist_ok=True)
                    except Exception:
                        pass
                    with open(audio_raw_path_local, 'wb') as out_f:
                        for chunk in r.iter_content(chunk_size=32768):
                            if chunk:
                                out_f.write(chunk)
                                written += len(chunk)
                    download_finished = time.time()
                    download_duration = download_finished - download_start
                    print(f"✓ Downloaded {written:,} bytes from {cand} in {download_duration:.2f}s")

                    # If we received a very small file, it might be a placeholder or transient.
                    # Retry the same candidate a few times before moving on to the next.
                    SMALL_BYTES_THRESHOLD = 1400
                    if written > 0 and written < SMALL_BYTES_THRESHOLD:
                        retry_small_attempts = 0
                        max_small_retries = 3
                        while retry_small_attempts < max_small_retries and written < SMALL_BYTES_THRESHOLD:
                            retry_small_attempts += 1
                            print(f"[Download] Received small file ({written} bytes) from {cand}, retrying ({retry_small_attempts}/{max_small_retries})...")
                            try:
                                time.sleep(0.45)
                                # re-request
                                r2 = _http_session.get(cand, stream=True, timeout=(3, 20))
                                if r2 is not None and getattr(r2, 'status_code', None) == 200:
                                    written = 0
                                    with open(audio_raw_path_local, 'wb') as out_f2:
                                        for chunk in r2.iter_content(chunk_size=32768):
                                            if chunk:
                                                out_f2.write(chunk)
                                                written += len(chunk)
                                    print(f"[Download] After retry #{retry_small_attempts} downloaded {written} bytes from {cand}")
                                    if written >= SMALL_BYTES_THRESHOLD:
                                        break
                            except Exception as rexc:
                                print(f"[Download] Retry error from {cand}: {rexc}")
                                continue

                    # If after retries we still have too-small file, treat as failed and try next candidate
                    if written == 0:
                        print(f"[Download] {cand} delivered no content; trying next candidate")
                        continue
                    if written < SMALL_BYTES_THRESHOLD:
                        print(f"[Download] {cand} delivered too-small file ({written} bytes); trying next candidate")
                        # reset written so next candidate can accumulate
                        written = 0
                        continue

                    # Successful download; stop trying candidates
                    break
                else:
                    print(f"[Download] {cand} returned status {getattr(r, 'status_code', None)}")
            except Exception as de:
                print(f"[Download] Error from {cand}: {de}")
                continue

        # Record diagnostics about the downloaded bytes so callers can decide
        # whether a real audio was present (avoid spurious "didn't hear" prompts)
        try:
            globals()['LAST_RECORDING_BYTES'] = int(written or 0)
            globals()['LAST_AUDIO_PRESENT'] = (int(written or 0) >= globals().get('AUDIO_PRESENT_THRESHOLD', 1400))
        except Exception:
            globals()['LAST_RECORDING_BYTES'] = 0
            globals()['LAST_AUDIO_PRESENT'] = False

        if written == 0:
            print("⚠ Download failed from all candidates")
            return "", False

        # Convert to 16kHz mono WAV for Google STT
        if written > 100:
            if ffmpeg_path:
                try:
                    print(f"⏱ Converting to 16kHz mono...")
                    ffmpeg_start = time.time()

                    cmd = [ffmpeg_path, '-threads', '0', '-y', '-i', audio_raw_path_local,
                           '-vn', '-sn', '-f', 'wav', '-ac', '1', '-ar', '16000',
                           '-acodec', 'pcm_s16le', 'pipe:1', '-hide_banner', '-loglevel', 'error']

                    proc = subprocess.Popen(cmd, stdin=subprocess.DEVNULL,
                                          stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                    out, err = proc.communicate(timeout=30)

                    if proc.returncode == 0:
                        converted_bytes = out
                        ffmpeg_duration = time.time() - ffmpeg_start
                        print(f"✓ Converted in {ffmpeg_duration:.2f}s ({len(out):,} bytes)")
                    else:
                        print(f"⚠ FFmpeg failed: {err.decode('utf-8', errors='ignore')[:200]}")
                        converted_bytes = None
                except Exception as e:
                    print(f"⚠ FFmpeg error: {e}")
                    converted_bytes = None
            else:
                # Fallback to pydub when ffmpeg not available
                try:
                    audio = AudioSegment.from_file(audio_raw_path_local)
                    audio = audio.set_channels(1).set_frame_rate(16000)
                    buf = io.BytesIO()
                    audio.export(buf, format='wav')
                    converted_bytes = buf.getvalue()
                except Exception as e:
                    print(f"⚠ pydub failed: {e}")
                    converted_bytes = None
            
    except Exception as e:
        print(f"⚠ Download error: {e}")
        return "", False
    
    # If download failed, return empty
    if not converted_bytes:
        print("⚠ No audio data to transcribe")
        return "", False
    
    # ===== TRANSCRIBE WITH GOOGLE SPEECH-TO-TEXT =====
    try:
        # OPTIMIZATION: Use cached client to avoid repeated initialization
        if not hasattr(record_and_transcribe, '_stt_client'):
            # Use default client which respects GOOGLE_APPLICATION_CREDENTIALS env var
            record_and_transcribe._stt_client = speech.SpeechClient()
            
            # Per user request: use English-only recognition and English keyword lists
            record_and_transcribe._stt_config_english = speech.RecognitionConfig(
                encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
                sample_rate_hertz=16000,
                language_code='en-US',
                enable_automatic_punctuation=False,
                use_enhanced=True,
                model='default',
                speech_contexts=[
                    speech.SpeechContext(
                        phrases=[
                            # English digits and common responses
                            "0", "1", "2", "3", "4", "5", "6", "7", "8", "9",
                            "zero", "one", "two", "three", "four", "five",
                            "six", "seven", "eight", "nine", "oh", "double zero",
                            "yes", "no", "correct", "wrong", "okay", "answer", "question"
                        ],
                        boost=20.0
                    )
                ],
                enable_word_time_offsets=False,
                enable_word_confidence=True,
                profanity_filter=False,
                max_alternatives=5
            )

            # Conversation-focused config (also English)
            record_and_transcribe._stt_config_filipino = speech.RecognitionConfig(
                encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
                sample_rate_hertz=16000,
                language_code='en-US',
                enable_automatic_punctuation=True,
                speech_contexts=[
                    speech.SpeechContext(
                        phrases=[
                            "yes", "no", "correct", "wrong", "okay", "answer", "question",
                            "one", "two", "three", "four", "five", "ten"
                        ],
                        boost=12.0
                    )
                ],
                enable_word_time_offsets=False,
                enable_word_confidence=False,
                profanity_filter=False,
                max_alternatives=3
            )

            # Default to English recognition
            record_and_transcribe._stt_config = record_and_transcribe._stt_config_english
        
        print(f"⏱ Starting transcription at {time.strftime('%H:%M:%S')}...")
        transcription_start = time.time()
        
        # Choose config based on use_english flag
        config_to_use = record_and_transcribe._stt_config_english if use_english else record_and_transcribe._stt_config_filipino
        
        if use_english:
            print("🔢 Using BILINGUAL model (English primary + Filipino secondary)")
            print(f"   Primary: en-US | Secondary: fil-PH | Enhanced: True | Alternatives: 5")
            print(f"   Optimized for: Digits, numbers, and mixed English/Filipino speech")
        else:
            print("🗣️ Using BILINGUAL model (Filipino primary + English secondary)")
            print(f"   Primary: fil-PH | Secondary: en-US | Alternatives: 3")
            print(f"   Optimized for: Conversation, yes/no, and mixed Filipino/English speech")
        
        audio_file = speech.RecognitionAudio(content=converted_bytes)
        response = record_and_transcribe._stt_client.recognize(
            config=config_to_use, 
            audio=audio_file
        )
        
        transcription_finished = time.time()
        transcription_duration = transcription_finished - transcription_start
        print(f"✓ Transcription completed in {transcription_duration:.2f}s")
        
    except Exception as exc:
        print('Transcription error:', exc)
        return "", forced_timeout

    complete_text = ""
    best_confidence = 0.0

    # Collect all alternatives across all results and pick a single best hypothesis
    if response and hasattr(response, 'results') and len(response.results) > 0:
        alternatives_info = []
        for ridx, result in enumerate(response.results):
            if not getattr(result, 'alternatives', None):
                continue
            for aidx, alt in enumerate(result.alternatives):
                text = alt.transcript if hasattr(alt, 'transcript') else ''
                confidence = alt.confidence if hasattr(alt, 'confidence') else 0.0
                digit_count = sum(c.isdigit() for c in text)
                word_count = len(text.split())
                text_len = len(text.replace(' ', '')) or 1
                digit_ratio = digit_count / text_len
                alternatives_info.append({
                    'text': text.strip(),
                    'confidence': confidence,
                    'digit_count': digit_count,
                    'word_count': word_count,
                    'digit_ratio': digit_ratio,
                    'result_idx': ridx,
                    'alt_idx': aidx
                })

        # If we collected any alternatives, score them and pick the single best
        if alternatives_info:
            # Scoring rules: if use_english (digit-focused) prefer more digits then confidence,
            # otherwise prefer higher confidence and penalize digit-heavy alternatives for conversational answers.
            def score_alt(x):
                if use_english:
                    return (x['digit_count'], x['confidence'], -x['digit_ratio'])
                else:
                    # For Filipino/conversational mode: prefer higher confidence, fewer digits
                    return (x['confidence'], -x['digit_ratio'], x['word_count'])

            alternatives_info.sort(key=score_alt, reverse=True)
            best = alternatives_info[0]
            complete_text = best['text']
            best_confidence = best['confidence']
            try:
                print(f"\n📊 Selected single best transcription: '{complete_text}' (conf: {best_confidence:.1%}, digits: {best['digit_count']}, ratio: {best['digit_ratio']:.2f})")
            except Exception:
                pass
    
    # Print total processing time
    if 'download_start' in locals():
        total_time = time.time() - download_start
        print(f"⏱ Total processing time: {total_time:.2f}s (download + convert + transcribe)")
    
    print(f"Transcribed text: {complete_text}")

    # Save transcription to file
    try:
        with open(U('transcribed_text.txt'), 'w', encoding='utf-8') as f:
            f.write(complete_text)
    except Exception:
        pass

    # Clean up both local cached file AND remote recording after transcription
    try:
        if os.path.exists(audio_raw_path_local):
            os.remove(audio_raw_path_local)
            print(f"✓ Cleaned up local cache: {audio_raw_path_local}")
    except Exception:
        pass
    
    # Clear remote recording after successful transcription
    print("Clearing remote recording after transcription...")
    
    clear_success = False
    
    # METHOD 1: Direct ESP32 recorder (use discovered IP or esp_host parameter)
    clear_target_ip = esp_host if esp_host else (_discovered.get('recorder') if _discovered else None)
    
    if clear_target_ip:
        try:
            # Primary: ESP /clear endpoint with discovered IP
            clear_resp = _http_session.get(f'http://{clear_target_ip}/clear', timeout=3.0)
            if clear_resp.status_code == 200:
                print(f"✓ ESP /clear successful (IP: {clear_target_ip})")
                clear_success = True
            
            # Also call /stop to ensure recording stopped
            _http_session.get(f'http://{clear_target_ip}/stop', timeout=2.0)
        except Exception as e:
            print(f"⚠ Could not reach ESP at {clear_target_ip}: {e}")
    
    # METHOD 2: Flask proxy (reliable fallback)
    if not clear_success:
        try:
            proxy_clear = _http_session.get('http://localhost:5000/clear', timeout=2.0)
            if proxy_clear.status_code == 200:
                print(f"✓ Flask proxy /clear successful")
                clear_success = True
        except Exception as e:
            print(f"⚠ Flask proxy clear failed: {e}")
    
    # METHOD 3: ngrok tunnel endpoints
    if not clear_success:
        try:
            parsed = urllib.parse.urlparse(audio_raw_url)
            if parsed.netloc:
                base = f'{parsed.scheme}://{parsed.netloc}'
                for path in ['/clear', '/clear_recording']:
                    try:
                        resp = _http_session.get(f'{base}{path}', timeout=2.0)
                        if resp.status_code == 200:
                            print(f"✓ Tunnel {path} successful")
                            clear_success = True
                            break
                    except Exception:
                        pass
        except Exception:
            pass
    
    # METHOD 4: DELETE method on recording URL (last resort)
    if not clear_success:
        try:
            delete_resp = _http_session.delete(audio_raw_url, timeout=2.0)
            if delete_resp.status_code in [200, 204, 404]:
                print(f"✓ DELETE successful")
                clear_success = True
        except Exception:
            pass
    
    if clear_success:
        print("✓ Remote recording cleared successfully")
        
        # CRITICAL: Reset Flask notification queue to prevent stale notifications
        # This ensures next recording detection waits for a NEW recording, not old notification
        try:
            reset_resp = _http_session.get('http://localhost:5000/reset_recording_notification', timeout=1.0)
            if reset_resp.status_code == 200:
                print("✓ Flask notification queue reset")
        except Exception:
            pass  # Not critical if this fails
    else:
        print("⚠ Warning: Could not verify remote recording cleared")

    return complete_text, forced_timeout


def is_profane_by_ai(text_to_check):
    """Uses the AI's safety filters to check for profanity."""
    if not text_to_check or not text_to_check.strip():
        return False
    try:
        # We send the user's text to the model. If the model's safety filter
        # blocks the prompt, we consider it to contain profanity.
        try:
            response = model.generate_content(text_to_check)
        except Exception as e:
            # If the AI check fails, be permissive and don't block
            print(f"Warning: AI profanity check failed: {e}")
            return False

        # Defensive: response may be None or shaped differently across SDKs
        if response is None:
            return False

        candidates = getattr(response, 'candidates', None)
        prompt_feedback = getattr(response, 'prompt_feedback', None)

        if (not candidates) and prompt_feedback and getattr(prompt_feedback, 'block_reason', None) == 'SAFETY':
            return True

        return False
    except Exception as e:
        print(f"Warning: Could not check for profanity via AI due to an error: {e}")
        return False # Fail safe: don't block if the check fails.

def is_profane_input(text: str) -> bool:
    """Fast, local profanity detector for input text.
    Uses a small curated list of common Filipino/English swear words and a
    word-boundary regex. This is intentionally conservative and meant only
    as a first-line filter to avoid sending explicit profanity to the model.
    """
    try:
        if not text or not isinstance(text, str):
            return False
        t = text.lower()
        # Small curated list (kept intentionally short and conservative)
        swear_words = [
            'putangina', 'putang ina', 'p*tang ina', 'pakshet', 'gago', 'tanga', 'ulol',
            'sh*t', 'shit', 'fuck', 'f**k', 'bastard', 'bobo', 'buwisit'
        ]
        # Build a simple regex that matches whole words or common obfuscated variants
        for w in swear_words:
            # escape and use word boundaries where sensible
            pattern = r"\b" + re.escape(w) + r"\b"
            if re.search(pattern, t):
                return True
        # Also flag repeated non-word sequences that often represent profanity
        if re.search(r"[!@#\$%\^&\*]{2,}", text):
            return True
    except Exception:
        # On error, be permissive (do not block) to avoid false positives
        return False
    return False

def read_state_text_file(url_state):
    # Try several sensible sources in order: local Flask proxy, discovered player IP, then provided url_state (ngrok)
    candidates = []
    # local Flask proxy first
    candidates.append('http://127.0.0.1:5000/state.txt')
    # discovered player IP (if auto-discovery populated it)
    try:
        if _discovered and isinstance(_discovered, dict) and 'player' in _discovered:
            candidates.append(f"http://{_discovered['player']}/state.txt")
    except Exception:
        pass
    # finally the configured url_state (may be ngrok)
    if url_state:
        candidates.append(url_state)

    for candidate in candidates:
        try:
            response = requests.get(candidate, timeout=(2, 4))
            if response.status_code == 200:
                return response.text.strip().lower()
            else:
                # Silently skip common/expected failures: 404 (not found), 5xx (server errors)
                # Only log unusual status codes
                if response.status_code not in [404, 500, 502, 503, 504]:
                    print(f"⚠ State check: {candidate} returned {response.status_code}")
        except requests.exceptions.Timeout:
            # timeout is common if the ESP is busy; try next candidate
            continue
        except requests.exceptions.RequestException as e:
            # Log only if it's an unexpected error (not simple connection refused)
            if "Connection" not in str(e) and "Timeout" not in str(e):
                print(f"⚠ State check error for {candidate}: {e}")
            continue

    # Nothing succeeded
    return None

def set_state_to_assisting_mode():
    """Force state to 0 (Assisting Mode) via ngrok tunnel"""
    try:
        # Prefer local Flask proxy /set_state, then discovered player IP, then ngrok as last resort
        candidates = []
        candidates.append('http://127.0.0.1:5000/set_state')
        try:
            if _discovered and isinstance(_discovered, dict) and 'player' in _discovered:
                candidates.append(f"http://{_discovered['player']}/set_state")
        except Exception:
            pass
        # ngrok fallback (if configured)
        candidates.append('https://nonbasic-bob-inimical.ngrok-free.dev/set_state')

        print("\n" + "="*60)
        print("INITIALIZING GENTA SYSTEM - Setting state to Assisting Mode...")
        print("="*60)

        for set_state_url in candidates:
            try:
                response = requests.get(f"{set_state_url}?value=0", timeout=(3, 10))
                if response.status_code == 200:
                    print(f"✓ State set to 0 (Assisting Mode) via {set_state_url}")
                    print("✓ State button DISABLED until LRN is entered")
                    return True
                else:
                    print(f"⚠ Warning: Could not set state via {set_state_url}. Status: {response.status_code}")
                    # Continue to next candidate
            except Exception as e:
                # Try next candidate; log first few failures for diagnostics
                print(f"[set_state] Could not reach {set_state_url}: {e}")

        print("⚠ Warning: Could not set state via any candidate endpoints.")
        return False
            
    except Exception as e:
        print(f"⚠ Warning: Could not set initial state: {e}")
        print("  System will use current ESP32 state")
        return False

def check_for_state_change():
    """Quick state check - returns True if state has changed"""
    global _STATE_CHANGE_REQUESTED, _CURRENT_STATE
    try:
        new_state = read_state_text_file(url_state)
        if new_state and new_state.strip() and new_state.strip() != _CURRENT_STATE:
            set_oled_expression('idle')
            print(f"\n[State Change Detected] {_CURRENT_STATE} → {new_state.strip()}")
            # Start a lightweight processing animation to indicate we're handling the state change
            try:
                _start_state_processing_animation()
                print("[State Monitor] Started processing animation due to state change")
            except Exception:
                pass
            _STATE_CHANGE_REQUESTED = True
            return True
    except Exception:
        pass
    return False

def load_melcs_knowledge():
    """Load MELCs (Most Essential Learning Competencies) from database to use as knowledge base.
    Returns a formatted string containing all MELCs content for the AI model.
    """
    melcs_content = ""
    try:
        connection = mysql.connector.connect(
            host=getattr(config, 'MYSQL_HOST', 'localhost'),
            port=getattr(config, 'MYSQL_PORT', 3306),
            database=getattr(config, 'MYSQL_DB', ''),
            user=getattr(config, 'MYSQL_USER', ''),
            password=getattr(config, 'MYSQL_PASS', '')
        )
        
        if connection and connection.is_connected():
            cursor = connection.cursor()
            # Fetch all MELCs for Mathematics (subject_id = 1)
            # If CURRENT_TEACHER_ID is set, filter by teacher_id to show only that teacher's MELCs
            if CURRENT_TEACHER_ID:
                cursor.execute("""
                    SELECT id, description 
                    FROM melcs 
                    WHERE subject_id = 1 AND teacher_id = %s
                    ORDER BY id
                """, (CURRENT_TEACHER_ID,))
            else:
                # Load all Mathematics MELCs if no teacher context
                cursor.execute("""
                    SELECT id, description 
                    FROM melcs 
                    WHERE subject_id = 1
                    ORDER BY id
                """)
            
            melcs_rows = cursor.fetchall()
            
            if melcs_rows:
                melcs_content = "\n\nMATHEMATICS KNOWLEDGE BASE (Grade 3 MELCs):\n"
                melcs_content += "="*60 + "\n"
                
                for row in melcs_rows:
                    melc_id = row[0] if row[0] else ""
                    description = row[1] if row[1] else ""
                    
                    if description:
                        melcs_content += f"\n[MELC-{melc_id}] {description}\n"
                
                melcs_content += "\n" + "="*60 + "\n"
                print(f"✓ Loaded {len(melcs_rows)} MELCs from database for AI knowledge base")
            else:
                print("⚠ No MELCs found in database - AI will use general knowledge only")
            
            cursor.close()
            connection.close()
    
    except Exception as e:
        print(f"⚠ Could not load MELCs from database: {e}")
        print("  AI will use general mathematics knowledge")
    
    return melcs_content

def GENTA():
    # === STARTUP CLEANUP: Clear any old recordings before beginning ===
    print("\n" + "="*60)
    print("GENTA STARTUP: Clearing old recordings from previous sessions...")
    print("="*60)
    
    # Delete local recording file if it exists
    try:
        if os.path.exists(audio_raw_path):
            os.remove(audio_raw_path)
            print(f"✓ Deleted local file: {audio_raw_path}")
        else:
            print(f"✓ No local recording file found")
    except Exception as e:
        print(f"⚠ Could not delete local file: {e}")
    
    # Clear remote recordings - Use DISCOVERED ESP32 recorder IP (not hardcoded)
    # This is critical - we must use the actual recorder IP from discovery
    esp_recorder_ip = None
    if _discovered and 'recorder' in _discovered:
        esp_recorder_ip = _discovered['recorder']
        print(f"✓ Using discovered recorder IP: {esp_recorder_ip}")
    else:
        # Fallback: try Flask proxy method
        print("⚠ No recorder IP discovered, will use Flask proxy to clear")
    
    clear_success = False
    
    # METHOD 1: Direct ESP32 recorder (BEST - most reliable)
    if esp_recorder_ip:
        print(f"Attempting to clear ESP32 directly at {esp_recorder_ip}...")
        try:
            # Try ESP /clear endpoint directly (most reliable method)
            clear_resp = _http_session.get(f'http://{esp_recorder_ip}/clear', timeout=3.0)
            if clear_resp.status_code == 200:
                print(f"✓ ESP /clear successful (HTTP {clear_resp.status_code})")
                clear_success = True
            
            # Also try /stop to ensure recording stopped
            stop_resp = _http_session.get(f'http://{esp_recorder_ip}/stop', timeout=2.0)
            if stop_resp.status_code == 200:
                print(f"✓ ESP /stop successful (HTTP {stop_resp.status_code})")
        except Exception as e:
            print(f"⚠ Could not reach ESP recorder directly: {e}")
    
    # METHOD 2: Flask proxy endpoints (backup)
    if not clear_success:
        print("Trying Flask proxy clear endpoints...")
        try:
            # Try Flask proxy /clear (routes to discovered recorder)
            proxy_resp = _http_session.get('http://localhost:5000/clear', timeout=3.0)
            if proxy_resp.status_code == 200:
                print(f"✓ Flask proxy /clear successful")
                clear_success = True
        except Exception as e:
            print(f"⚠ Flask proxy /clear failed: {e}")
    
    # METHOD 3: ngrok tunnel endpoints (if proxy failed)
    if not clear_success:
        try:
            parsed = urllib.parse.urlparse(audio_raw_url)
            if parsed.netloc:
                base = f'{parsed.scheme}://{parsed.netloc}'
                
                # Try various proxy clear endpoints
                for path in ['/clear', '/clear_recording', '/reset']:
                    try:
                        resp = _http_session.get(f'{base}{path}', timeout=3.0)
                        if resp.status_code == 200:
                            print(f"✓ Tunnel {path} successful")
                            clear_success = True
                            break
                    except Exception:
                        pass
        except Exception:
            pass
    
    # Wait for deletions to complete
    if clear_success:
        print("Waiting for deletions to propagate...")
        time.sleep(1.0)
        
        # Verify recording is actually gone (use discovered IP)
        verification_attempts = 3
        recording_still_exists = False
        
        for attempt in range(verification_attempts):
            try:
                # Check via discovered recorder IP
                check_url = f'http://{esp_recorder_ip}/recording.wav' if esp_recorder_ip else 'http://localhost:5000/recording.wav'
                verify_resp = _http_session.head(check_url, timeout=2.0)
                
                if verify_resp.status_code == 200:
                    size = verify_resp.headers.get('content-length', '0')
                    print(f"  Attempt {attempt + 1}: Recording exists (size={size})")
                    if attempt < verification_attempts - 1:
                        # Try clearing again with discovered IP
                        if esp_recorder_ip:
                            _http_session.get(f'http://{esp_recorder_ip}/clear', timeout=2.0)
                        else:
                            _http_session.get('http://localhost:5000/clear', timeout=2.0)
                        time.sleep(0.5)
                    else:
                        recording_still_exists = True
                else:
                    print(f"✓ Verified: Recording cleared")
                    break
            except Exception:
                print(f"✓ Verified: Recording cleared or unreachable")
                break
        
        if not recording_still_exists:
            print("✓ Startup cleanup completed successfully")
        else:
            print("⚠ WARNING: Recording may still exist - will use size detection")
    else:
        print("⚠ WARNING: Could not clear recordings - system will detect and skip old files")
    
    # Helper to try toggling state button on speaker via multiple possible routes
    def _try_toggle_state_button(enable: bool):
        """Try multiple endpoints (local Flask proxy, configured ESP_SPEAKER, playback host, legacy IP)
        Returns (success: bool, url_used: str|None)
        """
        action = 'enable_state_button' if enable else 'disable_state_button'
        candidates = []

        # 1) Prefer local Flask proxy (if running) - this will proxy to the correct ESP
        candidates.append(f'http://127.0.0.1:5000/{action}')

        # 2) If environment variable ESP_SPEAKER is set (e.g. http://192.168.0.109), try it
        esp_speaker_env = os.getenv('ESP_SPEAKER')
        if esp_speaker_env:
            candidates.append(f"{esp_speaker_env.rstrip('/')}/{action}")

        # 3) Try configured playback host (if present)
        try:
            if esp_playback_host:
                candidates.append(f'http://{esp_playback_host}/{action}')
        except Exception:
            pass

        # 4) Legacy hard-coded address (keeps backward compatibility)
        candidates.append(f'http://192.168.50.70/{action}')

        for url in candidates:
            try:
                resp = _http_session.get(url, timeout=5.0)
                if resp.status_code == 200:
                    return True, url
            except Exception:
                # ignore and continue to next candidate
                pass
        return False, None

    # Disable state button at startup (will be enabled after LRN entry)
    # SKIP button disable if returning from quiz mode (student already logged in)
    global CURRENT_STUDENT_ID, CURRENT_TEACHER_ID, CURRENT_TEACHER_NAME, CURRENT_STUDENT_NAME
    
    if CURRENT_STUDENT_ID and CURRENT_STUDENT_NAME:
        # Returning from quiz - button was already re-enabled, keep it enabled
        print("✓ Returning from quiz - keeping state button ENABLED")
    else:
        # First time startup - disable button until LRN is collected
        try:
            ok, used = _try_toggle_state_button(enable=False)
            if ok:
                print(f"✓ State button DISABLED (via {used})")
            else:
                print("⚠ Could not disable state button on any known endpoint")
        except Exception as e:
            print(f"⚠ Could not disable state button: {e}")
    
    print("="*60)
    print("GENTA READY: Starting LRN collection...\n")
    
    # Mark startup cleanup as complete
    global _STARTUP_CLEANUP_DONE
    _STARTUP_CLEANUP_DONE = True
    
    # SKIP LRN collection if student already logged in (returning from Quiz Mode)
    # Note: globals already declared above for button check
    if CURRENT_STUDENT_ID and CURRENT_STUDENT_NAME:
        print("\n" + "="*60)
        print(f"✓ Welcome back, {CURRENT_STUDENT_NAME}!")
        print(f"✓ LRN: {CURRENT_STUDENT_ID}")
        print(f"✓ Teacher: {CURRENT_TEACHER_NAME}")
        print("Skipping LRN collection - continuing session...")
        print("="*60 + "\n")
        # Skip to main loop (don't collect LRN again)
        pass
    else:
        # First time - need to collect LRN
        print("New session - collecting student LRN...")
    
    def retrieve_and_store_remarks(host, database, user, password, student_id, port=3306):
        """Return a dict with student and teacher info for the given student_id.
        Keys returned: remark, student_name, teacher_id, teacher_name
        """
        connection = None
        cursor = None
        try:
            # Connect to the database
            connection = mysql.connector.connect(
                host=host,
                database=database,
                user=user,
                password=password,
                port=port
            )

            remark = None
            student_name = None
            teacher_id = None
            teacher_name = None

            if connection and connection.is_connected():
                cursor = connection.cursor()
                # Try some common lookup strategies: id (numeric), lrn/student_number (string)
                tried = []
                # sanitize candidate
                cand = str(student_id).strip()
                # Prefer explicit LRN column lookup first (user specified column name 'lrn')
                if cand:
                    try:
                        cursor.execute("SELECT remarks, name, teacher_id FROM students WHERE lrn = %s LIMIT 1", (cand,))
                        row = cursor.fetchone()
                        tried.append(('lrn', cand))
                        if row:
                            remark = row[0] if len(row) > 0 else None
                            student_name = row[1] if len(row) > 1 else None
                            teacher_id = row[2] if len(row) > 2 else None
                    except Exception:
                        pass

                # If not found via lrn, attempt numeric id lookup (common path)
                if not student_name and cand.isdigit():
                    try:
                        cursor.execute("SELECT remarks, name, teacher_id FROM students WHERE id = %s", (int(cand),))
                        row = cursor.fetchone()
                        tried.append(('id', cand))
                        if row:
                            remark = row[0] if len(row) > 0 else None
                            student_name = row[1] if len(row) > 1 else None
                            teacher_id = row[2] if len(row) > 2 else None
                    except Exception:
                        pass

                # If still not found, try matching common alternate columns (student_number or id)
                if not student_name:
                    try:
                        cursor.execute("SELECT remarks, name, teacher_id FROM students WHERE student_number = %s OR id = %s", (cand, cand))
                        row = cursor.fetchone()
                        tried.append(('student_number/id', cand))
                        if row:
                            remark = row[0] if len(row) > 0 else None
                            student_name = row[1] if len(row) > 1 else None
                            teacher_id = row[2] if len(row) > 2 else None
                    except Exception:
                        pass

                # If we have teacher_id, fetch teacher's full name (firstname + lastname)
                if teacher_id:
                    try:
                        cursor.execute("SELECT firstname, lastname FROM teachers WHERE id = %s", (teacher_id,))
                        trow = cursor.fetchone()
                        if trow and len(trow) >= 2:
                            # Concatenate firstname and lastname to get full name
                            firstname = trow[0] if trow[0] else ""
                            lastname = trow[1] if trow[1] else ""
                            teacher_name = f"{firstname} {lastname}".strip()
                    except Exception:
                        teacher_name = None

                final_remark = None
                if remark:
                    final_remark = str(remark) + " This is the remark on the student that you need to reinforce:"

                return {
                    'remark': final_remark,
                    'student_name': student_name,
                    'teacher_id': teacher_id,
                    'teacher_name': teacher_name,
                }

        except mysql.connector.Error as e:
            print("Error while connecting to MySQL", e)
            return None
        finally:
            try:
                if cursor:
                    cursor.close()
            except Exception:
                pass
            try:
                if connection and connection.is_connected():
                    connection.close()
            except Exception:
                pass
    # Voice-driven mandatory LRN collection: keep prompting until student confirms their LRN
    def _speak_and_play(text: str):
        """Small early TTS+play helper used before full synthesize_speech/play_audio are defined.
        Generates response.wav via Google TTS and plays on ESP32 (primary) with PC fallback.
        """
        if not text or not text.strip():
            print("⚠ _speak_and_play: Empty text provided")
            return
            
        # Delegate to the canonical module-level speak_and_play_text helper so
        # all TTS/upload behavior is centralized. This prevents inconsistent
        # behaviors caused by duplicate local helpers.
        try:
            top = globals().get('speak_and_play_text')
            if callable(top):
                try:
                    top(text)
                    return
                except Exception:
                    pass
        except Exception:
            pass

        # Last-resort fallback: synthesize and attempt the legacy upload behavior
        try:
            print(f"🔊 Fallback TTS generation: '{text[:50]}...'")
            synthesize_speech(text, out_path='response.wav', sample_rate_hz=24000)
            try:
                esp_upload_file('response.wav', endpoint='/upload', max_retries=2)
            except Exception:
                print('⚠ Fallback ESP upload failed')
        except Exception as e:
            print(f'⚠ Early TTS/play fallback failed: {e}')

    def _play_response_wav():
        """Play response.wav on ESP32 (primary) with PC fallback. ESP32 autoplays after upload."""
        # Try ESP32 first
        esp_success = False
        try:
            import socket
            s = socket.create_connection((esp_playback_host, 80), timeout=1.0)
            s.close()
            url = f'http://{esp_playback_host}/upload'
            try:
                ok = esp_upload_file('response.wav', endpoint='/upload', max_retries=3)
                if ok:
                    esp_success = True
                else:
                    print('⚠ ESP32 upload failed (response.wav)')
            except Exception:
                pass
        except Exception:
            pass
        
        # Fallback to local
        if not esp_success:
            # Enforce ESP-only playback. Do not fallback to local speaker.
            print("⚠ ESP32 playback failed and local fallback is disabled for response.wav.")

    def ask_for_lrn_via_voice():
        attempts = 0
        failed_speech_attempts = 0

        while True:
            attempts += 1
            
            # Only say full prompt on FIRST attempt, not on retries
            if attempts == 1:
                # Verbose console prompt so user sees what's expected
                print("Please provide your LRN (voice).")
                try:
                    # Ask student to say their LRN (full prompt only on first attempt)
                    prompt = "Please say your full L-R-N number now. Say all digits clearly."
                    try:
                        _speak_and_play(prompt)
                    except Exception as _e:
                        print('LRN prompt playback failed:', _e)
                except Exception:
                    pass
            
            # Quick HEAD to the download URL to see if the proxy is up
            try:
                # OPTIMIZATION: Use session and reduced timeout
                _ = _http_session.head(audio_raw_url, timeout=0.8)
            except Exception:
                if attempts == 1:  # Only log this once
                    print('Recording proxy not reachable; continuing with voice-only LRN flow.')

            # Wait briefly then record/transcribe
            time.sleep(0.3)
            lrn_text = ''
            forced = False
            print('Attempting to capture LRN via recording (short timeout)...')
            try:
                # Use English model for MUCH BETTER digit recognition
                lrn_text, forced = record_and_transcribe(
                    esp_host=esp_record_host, 
                    poll_for_recording=True, 
                    max_poll_seconds=30,  # INCREASED: 30 seconds to give user more time for LRN entry
                    use_english=False  # Accept Filipino number words (isa, dalawa, etc.)
                )
                print(f'LRN recording attempt returned: {lrn_text!r}, forced={forced}')
            except Exception as e:
                print('LRN recording failed (exception):', e)
                lrn_text = ''
                forced = False

            if not lrn_text or not lrn_text.strip():
                failed_speech_attempts += 1
                
                # Single clear message for retry (no redundant double messages)
                try:
                    if failed_speech_attempts >= 3:
                        # After multiple failures, give encouraging message and reset counter
                        _speak_and_play("Let's try again. Please say your L-R-N number.")
                        failed_speech_attempts = 0
                    else:
                        # Simple retry message
                        _speak_and_play('I did not hear you. Please say your L-R-N again.')
                except Exception:
                    pass
                
                time.sleep(0.5)  # Brief pause before next attempt
                continue

            lrn_text = lrn_text.strip()
            
            # Show what was transcribed
            print(f'DEBUG: Captured LRN text (raw): "{lrn_text}"')
            
            # Convert word numbers to digits (handles mixed transcriptions like "one 2 three")
            def words_to_digits(text):
                """Convert number words to digits and extract only digits.
                Handles English and Filipino number words.
                Example: "one two 3 four" -> "1234"
                """
                # Mapping of word numbers to digits
                word_to_digit = {
                    # English
                    'zero': '0', 'one': '1', 'two': '2', 'three': '3', 'four': '4',
                    'five': '5', 'six': '6', 'seven': '7', 'eight': '8', 'nine': '9',
                    # Filipino
                    'sero': '0', 'isa': '1', 'dalawa': '2', 'tatlo': '3', 'apat': '4',
                    'lima': '5', 'anim': '6', 'pito': '7', 'walo': '8', 'siyam': '9',
                    'sampu': '10',  # Special case: "sampu" = 10 (two digits)
                }
                
                # Lowercase and split into words
                words = text.lower().split()
                result = []
                
                for word in words:
                    # Remove punctuation from word
                    clean_word = re.sub(r'[^\w]', '', word)
                    
                    if clean_word in word_to_digit:
                        # Convert word to digit
                        result.append(word_to_digit[clean_word])
                    elif clean_word.isdigit():
                        # Already a digit, keep it
                        result.append(clean_word)
                    # Ignore non-numeric words
                
                # Join all digits (no spaces)
                return ''.join(result)
            
            # Convert words to digits and extract
            digits_only = words_to_digits(lrn_text)
            print(f'DEBUG: Converted to digits: "{digits_only}" (length: {len(digits_only)})')
            
            # Also try pure digit extraction as fallback
            pure_digits = re.sub(r"\D", "", lrn_text)
            if len(pure_digits) == 12 and len(digits_only) != 12:
                # If pure extraction gives exactly 12 but conversion doesn't, use pure
                print(f'DEBUG: Using pure digit extraction: "{pure_digits}"')
                digits_only = pure_digits
            elif len(digits_only) != 12 and len(pure_digits) != 12:
                # Neither worked perfectly, prefer the conversion result
                print(f'DEBUG: Both methods failed, using word conversion result')
            
            print(f'DEBUG: Final digits: "{digits_only}" (length: {len(digits_only)})')

            # STRICT VALIDATION #1: Must be EXACTLY 12 digits
            if len(digits_only) != 12:
                print(f'DEBUG: LRN rejected - not 12 digits (got {len(digits_only)} digits)')
                try:
                    set_oled_expression('idle')
                    if len(digits_only) < 12:
                        _speak_and_play(f'The number is too short. I only heard {len(digits_only)} digits. The L-R-N must be exactly 12 digits long. Please try again.')
                    else:
                        _speak_and_play(f'The number is too long. I heard {len(digits_only)} digits. The L-R-N must be exactly 12 digits long. Please try again.')
                except Exception:
                    pass
                time.sleep(0.3)
                continue

            # STRICT VALIDATION #2: Must contain ONLY digits (no letters/special chars in original)
            if not digits_only.isdigit():
                print('DEBUG: LRN rejected - contains non-numeric characters')
                try:
                    set_oled_expression('idle')
                    _speak_and_play('The L-R-N must contain digits only. Please repeat it clearly.')
                except Exception:
                    pass
                time.sleep(0.3)
                continue

            # STRICT VALIDATION #3: Must exist in database
            print(f'DEBUG: Checking database for LRN: {digits_only}')
            temp_student_info = retrieve_and_store_remarks(getattr(config, 'MYSQL_HOST', 'localhost'), getattr(config, 'MYSQL_DB', ''), getattr(config, 'MYSQL_USER', ''), getattr(config, 'MYSQL_PASS', ''), digits_only, port=getattr(config, 'MYSQL_PORT', 3306))
            
            if not temp_student_info or not isinstance(temp_student_info, dict) or not temp_student_info.get('student_name'):
                print('DEBUG: LRN rejected - not found in database')
                try:
                    # Privacy: do not speak the numeric LRN aloud. Keep message generic.
                    _speak_and_play('The L-R-N you gave is not registered in our system. Please check your L-R-N and try again.')
                except Exception:
                    pass
                time.sleep(0.3)
                continue

            # ALL VALIDATIONS PASSED - Now ask for confirmation
            student_name = temp_student_info.get('student_name', 'Estudyante')
            print(f'DEBUG: LRN valid! Found student: {student_name}')
            
            try:
                set_oled_expression('idle')
                # Read back the student name for confirmation only (avoid repeating the LRN aloud)
                confirm_prompt = f"Are you {student_name}? Please answer 'yes' or 'no'."
                try:
                    _speak_and_play(confirm_prompt)
                except Exception as _e:
                    print('Confirm prompt playback failed:', _e)
            except Exception:
                pass

            # CRITICAL: Wait for prompt to finish playing AND give student time to respond
            # This prevents premature detection of old recordings
            print("Waiting for student to respond to confirmation...")
            time.sleep(2.0)  # Wait 2 seconds: 1s for audio to finish, 1s for student to prepare

            # We'll allow up to 2 confirmation attempts: the first unclear answer triggers
            # a simple clarifying yes/no prompt; if still unclear, we will NOT go back to
            # re-collecting the full LRN (avoids looping the long LRN capture).
            confirmation_attempts = 0
            while True:
                conf_text = ''
                try:
                    # Accept both Filipino (oo/hindi) and English (yes/no) confirmations
                    print("Ready to capture confirmation response...")
                    conf_text, _ = record_and_transcribe(esp_host=esp_record_host, poll_for_recording=True, max_poll_seconds=30, use_english=False)
                    print(f'DEBUG: Confirmation transcription result: "{conf_text}"')
                except Exception as e:
                    print('LRN confirmation recording failed (exception):', e)
                    conf_text = ''

                if conf_text and conf_text.strip():
                    ct = conf_text.lower().strip()
                    print(f'DEBUG: Checking confirmation text (lowercased): "{ct}"')
                    # Check for affirmative responses
                    if any(w in ct for w in ['yes', 'y', 'yeah', 'yep', 'correct', 'right']):
                        print('DEBUG: User confirmed identity! Returning validated LRN.')
                        return digits_only
                    elif any(w in ct for w in ['no', 'not', 'wrong', 'nope']):
                        print('DEBUG: User denied identity. Will ask for LRN again.')
                        try:
                            _speak_and_play("Okay, please say the correct L-R-N again.")
                        except Exception:
                            pass
                        time.sleep(0.3)
                        # Break out to the outer loop to re-run full LRN capture
                        break
                    else:
                        # Unclear response
                        confirmation_attempts += 1
                        print(f"DEBUG: Unclear confirmation response (attempt {confirmation_attempts}): '{ct}'")
                        if confirmation_attempts == 1:
                            # Ask a short clarifying yes/no prompt (do NOT re-run full LRN capture yet)
                            try:
                                _speak_and_play(f"I didn't understand. Say 'yes' if you are {student_name}, or 'no' if not.")
                            except Exception:
                                pass
                            # Short pause before listening again
                            time.sleep(0.3)
                            continue
                        else:
                            # Second unclear response: do NOT restart full LRN capture to avoid loops.
                            print('DEBUG: Confirmation remained unclear after clarification. Proceeding without re-asking full LRN.')
                            try:
                                _speak_and_play('Identity unclear. We will inform your teacher.')
                            except Exception:
                                pass
                            # Return the candidate digits (caller will handle storing/teacher review)
                            return digits_only
                else:
                    print('DEBUG: No confirmation text received (empty or None)')
                    failed_speech_attempts += 1
                    if failed_speech_attempts >= 3:
                        print('No confirmation heard after several attempts; prompting for LRN re-entry.')
                        failed_speech_attempts = 0
                        try:
                            _speak_and_play("I didn't hear a response. Please say your L-R-N again.")
                        except Exception:
                            pass
                        # Break out to outer loop to re-run full LRN capture
                        break
                    # Short pause then allow another confirmation listen
                    time.sleep(0.3)
                    confirmation_attempts += 1
                    if confirmation_attempts >= 2:
                        print('DEBUG: No usable confirmation after multiple listens; proceeding without re-asking full LRN.')
                        try:
                            _speak_and_play('Identity unclear. We will inform your teacher.')
                        except Exception:
                            pass
                        return digits_only
                    continue

    # Collect LRN with validation (voice-only). Repeat until we get a valid LRN
    # ONLY IF student not already logged in
    remarks = ""
    if not CURRENT_STUDENT_ID:
        while True:
            student_input = ask_for_lrn_via_voice()
            
            # At this point, student_input should already be:
            # - Exactly 12 digits
            # - Validated against database
            # - Confirmed by user
            sanitized_lrn = re.sub(r"\D", "", (student_input or ""))

            # Final safety check (should never fail since ask_for_lrn_via_voice does validation)
            if not sanitized_lrn or len(sanitized_lrn) != 12:
                print(f'ERROR: ask_for_lrn_via_voice returned invalid LRN: {sanitized_lrn}')
                try:
                    _speak_and_play('There is a system problem. Please repeat your L-R-N.')
                except Exception:
                    pass
                time.sleep(0.3)
                continue

            # Lookup in database (should succeed since already validated inside ask_for_lrn_via_voice)
            student_info = retrieve_and_store_remarks(getattr(config, 'MYSQL_HOST', 'localhost'), getattr(config, 'MYSQL_DB', ''), getattr(config, 'MYSQL_USER', ''), getattr(config, 'MYSQL_PASS', ''), sanitized_lrn, port=getattr(config, 'MYSQL_PORT', 3306))
            if not student_info or not isinstance(student_info, dict) or not student_info.get('student_name'):
                # This should never happen since we already validated, but handle it gracefully
                print(f'ERROR: Database lookup failed for previously validated LRN: {sanitized_lrn}')
                try:
                    _speak_and_play('There is a database connection problem. Please repeat your L-R-N.')
                except Exception:
                    pass
                time.sleep(0.3)
                continue

            # Success - store context and break
            remarks = student_info.get('remark') or ""
            try:
                # Global variables already declared at function start (line 1190)
                CURRENT_STUDENT_ID = sanitized_lrn
                CURRENT_TEACHER_ID = student_info.get('teacher_id')
                CURRENT_TEACHER_NAME = student_info.get('teacher_name')
                CURRENT_STUDENT_NAME = student_info.get('student_name')
            except Exception:
                pass
            print("Retrieved student info: ", {'student': CURRENT_STUDENT_NAME, 'teacher': CURRENT_TEACHER_NAME, 'teacher_id': CURRENT_TEACHER_ID})
            
            # Enable state button on ESP32 GENTA2 (speaker) now that LRN is validated
            try:
                ok, used = _try_toggle_state_button(enable=True)
                if ok:
                    print(f"✓ State button ENABLED - User can now change modes (via {used})")
                else:
                    print("⚠ Could not enable state button on any known endpoint")
            except Exception as e:
                print(f"⚠ Could not enable state button: {e}")

            # Prepare a personalized welcome message using the student's first name
            try:
                # Make the welcome message vary per session to feel fresh.
                first_name = CURRENT_STUDENT_NAME.split()[0] if CURRENT_STUDENT_NAME and isinstance(CURRENT_STUDENT_NAME, str) else None
                variants = []
                if first_name:
                    variants = [
                        f"Hello {first_name}. I'm GENTA. Are you ready to learn today? Ask me a math question and I'll explain it.",
                        f"Hi {first_name}! I'm GENTA. Let's start learning — do you have a question for me?",
                        f"Good day, {first_name}. I'm GENTA. Would you like to study now? Ask me a math question.",
                        f"Hello {first_name}. Ready for a learning adventure? I'm GENTA. Ask me about math anytime.",
                    ]
                else:
                    variants = [
                        "Hello. I'm GENTA. Are you ready to learn today?",
                        "Hi! I'm GENTA. Let's start learning.",
                        "Good day! Are you ready to learn with GENTA?"
                    ]

                # Pick a random variant per session
                try:
                    PERSONAL_WELCOME_MSG = random.choice(variants)
                except Exception:
                    PERSONAL_WELCOME_MSG = variants[0] if variants else None
            except Exception:
                PERSONAL_WELCOME_MSG = None
            break
    else:
        # Student already logged in - skip welcome message but RELOAD remarks from database
        # (remarks may have been updated after quiz completion)
        PERSONAL_WELCOME_MSG = None
        try:
            # Re-fetch remarks from database to get latest assessment feedback
            student_info = retrieve_and_store_remarks(getattr(config, 'MYSQL_HOST', 'localhost'), getattr(config, 'MYSQL_DB', ''), getattr(config, 'MYSQL_USER', ''), getattr(config, 'MYSQL_PASS', ''), CURRENT_STUDENT_ID, port=getattr(config, 'MYSQL_PORT', 3306))
            remarks = student_info.get('remark') or ""
            print(f"✓ Reloaded remarks for returning student: {remarks[:100]}..." if len(remarks) > 100 else f"✓ Reloaded remarks: {remarks}")
        except Exception as e:
            print(f"⚠ Could not reload remarks for returning student: {e}")
            remarks = ""
    
    # Load MELCs knowledge base from database
    melcs_knowledge = load_melcs_knowledge()
    
    chat = model.start_chat(history=[
        {   "role": "user",            
            "parts": [{"text": """System prompt: You are an elementary math teacher named Jen-ta (GENTA) for Grade 3 students.

CRITICAL RULES FOR YOUR RESPONSES:
1. Keep answers SHORT - maximum 3-4 sentences only
2. Use SIMPLE English words that Grade 3 students understand
3. NO special symbols, NO asterisks, NO bullets, NO formatting
4. NO markdown, NO emphasis marks, just plain text
5. Explain like talking to a young child
6. Give ONE simple example using toys, candies, or fruits
7. Answer the question FIRST, then give one short explanation

KNOWLEDGE RESTRICTION:
- You can ONLY answer MATHEMATICS questions
- If a student asks about non-math topics (science, history, English, etc.), politely say: "Sorry, I'm only a math teacher. I cannot answer that question. Please ask about mathematics."
- If a math question is beyond Grade 3 level, simplify your explanation to Grade 3 understanding

YOUR PRIMARY KNOWLEDGE SOURCE:
- First, use the MELCs (Most Essential Learning Competencies) content provided below
- If MELCs don't cover the specific question, you may use your general mathematics knowledge
- Always stay within Grade 3 mathematics scope

Example of GOOD response:
"5 plus 3 is 8. If you have five candies and someone gives you three more, you have eight candies in total."

Example of BAD response (TOO LONG):
"Addition is an important concept in mathematics... [many sentences]... Therefore the answer is 8."

Example of NON-MATH question response:
Student: "What is the capital of the Philippines?"
GENTA: "Sorry, I'm only a math teacher. I cannot answer that question. Please ask about mathematics."

Remember: SHORT, SIMPLE, NO SYMBOLS, Grade 3 MATH ONLY. """ + melcs_knowledge + "\n\n" + remarks}]
        },
        {   "role": "model",
            "parts": [{"text": """Understood. I will speak briefly and simply, with no symbols, for a Grade 3 student."""}]
        },
        {   "role": "user", 
            "parts": [{"text": """Example: What is 5 plus 3?"""}]
        },
        {   "role": "model",      
            "parts": [{"text": """5 plus 3 is 8. If you have five candies and someone gives you three more, you now have eight candies in total."""}]
        },
        {   "role": "user", 
            "parts": [{"text": """System prompt: who is your creator"""}]
        },
        {   "role": "model",      
            "parts": [{"text": """Ako ay ginawa ng 4th year Computer Science students sa City College of Angeles. Sila ay sina Jino Guiwan, Jonas Tiglao, Cedric Garcia, at Maria Tiblani."""}]
        },
        {   "role": "user", 
            "parts": [{"text": """Who is your thesis adviser"""}]
        },
        {   "role": "model",      
            "parts": [{"text": """Dean Maika Garbes"""}]
        }, 
        {   "role": "user", 
            "parts": [{"text": """Who is your project adviser"""}]
        },
        {   "role": "model",      
            "parts": [{"text": """Sir Billy Yee"""}]
        },        
        {   "role": "user", 
            "parts": [{"text": """What is your name"""}]
        },
        {   "role": "model",      
            "parts": [{"text": """I am 'Jen-ta'."""}]
        },
        {   "role": "user", 
            "parts": [{"text": """Who are your favorite professors?"""}]
        },
        {   "role": "model",      
            "parts": [{"text": """My favorite professors are Dean Maika Garbes, and Sir Billy Yee."""}]
        },
        {   "role": "user", 
            "parts": [{"text": """Tell me something about your research or study."""}]
        },
        {   "role": "model",      
            "parts": [{"text": """In the Philippines, elementary students face a significant educational hurdle, particularly in Grade 4, 
                       where foundational competencies prove challenging to grasp. Our capstone project aimed to provide a possible solution 
                       for this issue by investigating the functionality of GENTA, an electronic GLM-powered learner-oriented tool, in assessing 
                       Grade 3 students' mathematical competencies. GENTA integrates hardware, utilizing ESP 32 for speech-to-text and 
                       text-to-speech capabilities, along with software functionalities accessible through a teacher dashboard. 
                       This integration enables the delivery of tailored learning experiences aligned with pedagogical themes on both platforms. 
                       Through rigorous statistical analyses, including ISO 25095 and Word Error Rate (WER) assessments, we validate GENTA's accuracy and reliability. 
                       Achieving a 91.94 percent accuracy rate in speech-to-text transcription, GENTA demonstrates its ability to identify individual 
                       student weaknesses and design personalized learning modules, thereby empowering teachers with targeted interventions. 
                       Our research highlights recommendations for enhancing GENTA's functionality, such as mitigating time constraints, 
                       exploring innovative technologies, and improving user interaction through visual components. 
                       Overall, GENTA shows promising potential in improving Grade 3 mathematics learning outcomes, 
                       with significant implications for elevating educational standards globally."""}]
        },
        {   "role": "user", 
            "parts": [{"text": """Example of NON-MATH question: Ano ang capital ng Pilipinas?"""}]
        },
        {   "role": "model",      
            "parts": [{"text": """Sorry, I'm only a math teacher. I cannot answer that question. Please ask about mathematics."""}]
        },
        {   "role": "user", 
            "parts": [{"text": """Example of NON-MATH question: What is photosynthesis?"""}]
        },
        {   "role": "model",      
            "parts": [{"text": """Sorry, I'm only a math teacher. I cannot answer that question. Please ask about mathematics."""}]
        },
        {   "role": "user", 
            "parts": [{"text": """Example of NON-MATH question: Sino ang national hero ng Pilipinas?"""}]
        },
        {   "role": "model",      
            "parts": [{"text": """Sorry, I'm only a math teacher. I cannot answer that question. Please ask about mathematics."""}]
        },        

    ])
    
    def convert_apostrophe(text):
        text1 = text.replace("&#39;", "'")
        text2 = text1.replace("&quot;", '"')
        return text2
    def simplify_for_grade3(text: str) -> str:
        """Aggressively remove ALL special symbols, markdown, and formatting.
        Make the text clean and simple for Grade 3 TTS playback - only plain text.
        """
        if not text:
            return text
        try:
            # Remove ALL asterisks (bold/italic markdown)
            s = re.sub(r"\*+", "", text)
            
            # Remove ALL underscores (markdown emphasis)
            s = re.sub(r"_+", "", s)
            
            # Remove hash symbols (headers)
            s = re.sub(r"#+\s*", "", s)
            
            # Remove bullet points and list markers
            s = re.sub(r"^[\s]*[-•–—▪▫►○●]\s*", "", s, flags=re.MULTILINE)
            
            # Remove numbered lists (1. 2. 3. etc)
            s = re.sub(r"^[\s]*\d+[\.\)]\s*", "", s, flags=re.MULTILINE)
            
            # Remove brackets and parentheses content that might be notes
            s = re.sub(r"\[.*?\]", "", s)
            
            # Remove backticks (code formatting)
            s = re.sub(r"`+", "", s)
            
            # Remove pipe symbols (tables)
            s = s.replace("|", "")
            
            # Remove angle brackets
            s = re.sub(r"[<>]", "", s)
            
            # Remove excessive punctuation (multiple !!!, ???, etc)
            s = re.sub(r"([!?.]){2,}", r"\1", s)
            
            # Replace multiple newlines with single newline
            s = re.sub(r"\n{2,}", "\n", s)
            
            # Trim whitespace from each line
            s = "\n".join([ln.strip() for ln in s.splitlines() if ln.strip()])
            
            # Collapse multiple spaces into one
            s = re.sub(r"\s{2,}", " ", s)
            
            # Remove leading/trailing whitespace
            s = s.strip()
            
            # Limit length: if response is too long (over 500 chars), truncate with message
            if len(s) > 500:
                # Find last sentence within 500 chars
                truncated = s[:500]
                last_period = truncated.rfind('.')
                if last_period > 100:  # At least keep some content
                    s = truncated[:last_period + 1]
                else:
                    s = truncated + "..."
            
            return s
        except Exception as e:
            print(f"Warning: simplify_for_grade3 error: {e}")
            return text
    def detect_simple_equation(text: str):
        """Detect simple arithmetic equations in user input.
        Returns (True, matched_expr) if a simple operator-based expression is found,
        otherwise (False, None).
        """
        try:
            if not text or not isinstance(text, str):
                return False, None
            t = text.lower()
            # Look for symbol-based simple expressions like '5+3' or '12 - 4'
            m = re.search(r"(\d+(?:\.\d+)?\s*[\+\-\*\/x×÷]\s*\d+(?:\.\d+)?)", t)
            if m:
                return True, m.group(1)

            # Look for word-based operators with numbers, e.g., '5 plus 3' or 'what is 5 plus 3'
            if re.search(r"\b(plus|minus|times|multiplied|divided|over|add|subtract)\b", t):
                # Conservative: return True but pass the original text for generation
                return True, text
        except Exception:
            pass
        return False, None

    def is_math_related(text: str) -> bool:
        """Conservative heuristic: return True if text looks like a math question
        or contains math-related keywords/digits. This helps enforce the Grade-3
        math-only policy by blocking non-math prompts before calling the model.
        """
        try:
            if not text or not isinstance(text, str):
                return False
            t = text.lower()
            # If it contains explicit equation symbols or digit-digit patterns, it's math
            if re.search(r"\d+\s*[\+\-\*/x×÷]\s*\d+", t):
                return True
            # Common math keywords (English/Filipino)
            # Comprehensive set covering Grade 3 Philippine Mathematics MELCs
            math_terms = [
                # ===== BASIC OPERATIONS & COMPUTATION =====
                'plus', 'add', 'addition', 'adding', 'sum', 'total', 'altogether', 'combine', 'more',
                'minus', 'subtract', 'subtraction', 'subtracting', 'difference', 'less', 'take away', 'remove',
                'times', 'multiply', 'multiplication', 'multiplied', 'product', 'groups of', 'sets of',
                'divide', 'division', 'divided', 'quotient', 'remainder', 'share', 'sharing', 'split', 'distribute',
                'equals', 'equal', 'equal to', 'same as', 'is', 'makes', 'gives', 'results in',
                'operation', 'operations', 'compute', 'calculate', 'solve', 'answer', 'solution',
                
                # Filipino operation terms
                'dagdag', 'idagdag', 'pagdagdag', 'bawas', 'bawasan', 'pagbawas', 'parami', 'pag-parami',
                'bahagi', 'bahagin', 'bahagiin', 'hatiin', 'hati', 'numero', 'ilang', 'bilangin',
                'kabuuan', 'kasagutan', 'sagot', 'tama', 'mali', 'kuwenta', 'tantiya',
                
                # ===== NUMBERS & NUMBER SENSE =====
                'number', 'numbers', 'numeral', 'numerals', 'digit', 'digits', 'count', 'counting',
                'odd', 'even', 'odd number', 'even number', 'whole number', 'integer',
                'zero', 'one', 'two', 'three', 'four', 'five', 'six', 'seven', 'eight', 'nine', 'ten',
                'first', 'second', 'third', 'fourth', 'fifth', 'ordinal',
                
                # ===== PLACE VALUE (Grade 3 MELCs) =====
                'place value', 'place', 'value', 'position', 'digit value',
                'ones', 'tens', 'hundreds', 'thousands', 'ten thousands', 'hundred thousands',
                'ones place', 'tens place', 'hundreds place', 'thousands place',
                'expanded form', 'expanded notation', 'standard form', 'word form',
                
                # ===== COMPARISON & ORDERING =====
                'compare', 'comparison', 'comparing', 'order', 'ordering', 'arrange', 'arrangement',
                'greater', 'greater than', 'biggest', 'largest', 'maximum', 'more than', 'higher',
                'less', 'less than', 'smaller', 'smallest', 'minimum', 'fewer', 'lower',
                'between', 'in between', 'middle', 'ascending', 'descending',
                'before', 'after', 'comes before', 'comes after', 'next to',
                
                # ===== ROUNDING & ESTIMATION =====
                'round', 'rounding', 'round off', 'round to', 'nearest',
                'estimate', 'estimation', 'estimating', 'about', 'approximately', 'close to', 'around',
                'reasonable', 'reasonable answer', 'check', 'verify',
                
                # ===== FRACTIONS (Grade 3 MELCs) =====
                'fraction', 'fractions', 'fractional', 'part', 'parts', 'piece', 'pieces',
                'half', 'halves', 'one half', 'third', 'thirds', 'fourth', 'fourths', 'quarter', 'quarters',
                'fifth', 'fifths', 'sixth', 'sixths', 'seventh', 'eighth', 'ninth', 'tenth',
                'numerator', 'denominator', 'whole', 'equal parts', 'divide into',
                'proper fraction', 'improper fraction', 'mixed number',
                
                # Filipino fraction terms
                'kalahati', 'katlo', 'kapat', 'pira-piraso', 'bahagi ng',
                
                # ===== MONEY (Grade 3 MELCs - Philippine Currency) =====
                'money', 'peso', 'pesos', 'centavo', 'centavos', 'piso', 'sentimo',
                'coin', 'coins', 'bill', 'bills', 'currency', 'cash',
                'price', 'cost', 'costs', 'worth', 'value', 'amount',
                'buy', 'bought', 'purchase', 'sell', 'sold', 'pay', 'payment', 'spend', 'spending',
                'change', 'change from', 'how much', 'total cost', 'total price',
                'expensive', 'cheap', 'cheaper', 'more expensive',
                
                # ===== MEASUREMENT - LENGTH (Grade 3 MELCs) =====
                'length', 'long', 'short', 'tall', 'height', 'width', 'distance',
                'meter', 'meters', 'centimeter', 'centimeters', 'millimeter', 'millimeters', 'kilometer', 'kilometers',
                'cm', 'mm', 'km', 'metro', 'sentimetro',
                'measure', 'measuring', 'measurement', 'ruler', 'tape measure', 'measuring tape',
                'perimeter', 'around', 'distance around', 'boundary',
                
                # ===== MEASUREMENT - MASS/WEIGHT (Grade 3 MELCs) =====
                'mass', 'weight', 'heavy', 'light', 'heavier', 'lighter', 'heaviest', 'lightest',
                'gram', 'grams', 'kilogram', 'kilograms', 'kg', 'gramo', 'kilo',
                'weigh', 'weighs', 'weighing', 'scale', 'balance',
                
                # ===== MEASUREMENT - CAPACITY/VOLUME (Grade 3 MELCs) =====
                'capacity', 'volume', 'holds', 'contains', 'container',
                'liter', 'liters', 'milliliter', 'milliliters', 'ml', 'litro',
                'full', 'empty', 'half full', 'more', 'less',
                'cup', 'glass', 'bottle', 'gallon',
                
                # ===== TIME (Grade 3 MELCs) =====
                'time', 'clock', 'watch', 'timer',
                'hour', 'hours', 'minute', 'minutes', 'second', 'seconds', 'oras', 'minuto', 'segundo',
                'o\'clock', 'half past', 'quarter past', 'quarter to',
                'morning', 'afternoon', 'evening', 'night', 'noon', 'midnight',
                'am', 'pm', 'a.m.', 'p.m.',
                'day', 'days', 'week', 'weeks', 'month', 'months', 'year', 'years',
                'calendar', 'date', 'today', 'yesterday', 'tomorrow',
                'elapsed time', 'duration', 'how long', 'time passed',
                
                # ===== GEOMETRY - 2D SHAPES (Grade 3 MELCs) =====
                'shape', 'shapes', 'figure', 'figures', 'geometry', 'geometric',
                'circle', 'circles', 'round', 'bilog',
                'square', 'squares', 'parisukat',
                'rectangle', 'rectangles', 'rectangular', 'parihaba',
                'triangle', 'triangles', 'triangular', 'tatsulok',
                'polygon', 'polygons', 'hexagon', 'pentagon', 'octagon',
                'side', 'sides', 'corner', 'corners', 'vertex', 'vertices', 'edge', 'edges',
                'angle', 'angles', 'right angle', 'corner angle',
                
                # ===== GEOMETRY - 3D SHAPES/SOLIDS (Grade 3 MELCs) =====
                'solid', 'solids', 'solid figure', 'solid figures', '3d', 'three dimensional',
                'cube', 'cubes', 'box', 'kubo',
                'sphere', 'spheres', 'ball',
                'cylinder', 'cylinders', 'can', 'tube',
                'cone', 'cones', 'ice cream cone',
                'rectangular prism', 'prism',
                'face', 'faces', 'flat face', 'curved face',
                
                # ===== AREA & PERIMETER (Grade 3 intro) =====
                'area', 'space', 'cover', 'covers', 'covering', 'square units',
                'perimeter', 'distance around', 'around', 'boundary', 'border',
                
                # ===== PATTERNS & ALGEBRA (Grade 3 MELCs) =====
                'pattern', 'patterns', 'patrones', 'sequence', 'sequences',
                'repeating', 'repeats', 'repeat', 'growing', 'grows', 'shrinking',
                'rule', 'rules', 'pattern rule', 'next', 'what comes next', 'missing',
                'array', 'arrays', 'rows', 'columns', 'row', 'column',
                'skip count', 'skip-count', 'skipcount', 'skip counting', 'count by',
                
                # ===== DATA & GRAPHS (Grade 3 MELCs) =====
                'data', 'information', 'datos', 'impormasyon',
                'graph', 'graphs', 'chart', 'charts', 'table', 'tables',
                'pictograph', 'picture graph', 'bar graph', 'bar chart',
                'tally', 'tally marks', 'tally chart',
                'survey', 'collect', 'organize', 'display',
                'most', 'least', 'fewest', 'most common', 'least common',
                'frequency', 'how many times',
                
                # ===== PROBLEM SOLVING =====
                'problem', 'problems', 'word problem', 'story problem', 'suliranin',
                'solve', 'solving', 'solution', 'solbе', 'lutasin', 'solusyon',
                'question', 'questions', 'ask', 'asking', 'tanong', 'itatanong',
                'given', 'what is given', 'information', 'facts', 'ibinigay',
                'find', 'look for', 'hanapin', 'tukuyin',
                'strategy', 'strategies', 'method', 'way', 'steps', 'hakbang',
                
                # ===== COMMON QUESTION WORDS =====
                'what', 'what is', 'what are', 'what does', 'what will',
                'how', 'how many', 'how much', 'how do', 'how to', 'how can',
                'why', 'why is', 'why does', 'explain', 'define', 'describe',
                'when', 'where', 'which', 'who', 'whose',
                'can you', 'could you', 'will you', 'would you',
                'show', 'show me', 'tell', 'tell me', 'give', 'give me',
                
                # ===== GENERAL MATH TERMS =====
                'math', 'mathematics', 'matematika', 'arithmetic', 'aritmetika',
                'lesson', 'topic', 'concept', 'example', 'halimbawa',
                'practice', 'exercise', 'activity', 'worksheet',
                'correct', 'incorrect', 'right', 'wrong', 'mistake', 'error',
                'easy', 'hard', 'difficult', 'simple', 'complicated',
                'understand', 'know', 'learn', 'study', 'remember', 'unawain'
            ]
            for kw in math_terms:
                if kw in t:
                    return True
            # If text contains any digit token, it often indicates a math question
            if re.search(r"\b\d+\b", t):
                return True
        except Exception:
            return False
        return False

    def generate_teaching_for_equation(expr_text: str, lang: str = 'en') -> str:
        """Create a short, step-by-step teaching response for a simple equation.
        This intentionally avoids giving the final numeric result; instead it
        guides the student to compute it themselves. Returns a short English reply.
        """
        try:
            if not expr_text:
                return "Okay, I'll show step-by-step how to solve that problem. Try it yourself first and tell me the answer." 

            t = str(expr_text).lower()
            # Attempt to parse explicit digit operator digit forms
            m = re.search(r"(\d+(?:\.\d+)?)\s*([\+\-\*\/x×÷])\s*(\d+(?:\.\d+)?)", t)
            if m:
                a, op, b = m.group(1), m.group(2), m.group(3)
                # Map operator to simple English teaching phrasing
                if op in ['+', 'plus'] or re.search(r'plus', t):
                    return f"First, imagine you have {int(float(a))} candies. Add the {int(float(b))} candies one by one and count as you add them. Now try counting them all and tell me the answer." 
                if op in ['-', '−', 'minus'] or re.search(r'minus', t):
                    return f"You have {int(float(a))} apples. If {int(float(b))} apples are taken away, subtract one by one and count how many are left. Tell me how many remain." 
                if op in ['*', 'x', '×', 'times', 'multiplied'] or re.search(r'times|multiplied', t):
                    return f"Multiplication is repeated addition. For example, if you have {int(float(a))} groups of {int(float(b))} items, add {int(float(b))} {int(float(a))} times. Try it and tell me the answer." 
                if op in ['/', '÷', 'over', 'divided'] or re.search(r'divide|divided|over', t):
                    return f"Divide the {int(float(a))} items into {int(float(b))} equal parts. Try dividing and count how many are in each part. Tell me what you notice." 

            # Fallback for word-based expressions: give a generic step guide
            if re.search(r'plus|add|minus|subtract|times|multiply|divide|divided', t):
                if re.search(r'plus|add', t):
                    return "To add, start with the first number, then add the second number one by one and count. Try it and tell me the answer." 
                if re.search(r'minus|subtract', t):
                    return "To subtract, start with the first number and take away the second number one at a time until finished. Say how many are left." 
                if re.search(r'times|multiply', t):
                    return "Multiplying is quick repeated adding. Add the same number to itself as many times as the other number says. Try it and tell me the answer." 
                if re.search(r'divide|divide|over', t):
                    return "To divide, split the total into equal parts according to the second number. Try it and tell me how many in each part." 

            # Generic prompt if nothing specific parsed
            return "Okay, I'll show step-by-step: start with the first number, then add or take away the second number one by one, and count until finished. Try it and tell me the answer." 
        except Exception:
            return "Okay, I'll show step-by-step how to solve that problem. Try it yourself first and tell me the answer." 

    def evaluate_equation_answer(original_question: str, student_answer: str) -> tuple:
        """
        Evaluate a student's answer to a math equation.
        Returns (is_correct: bool, correct_answer: float, feedback: str)
        """
        try:
            if not original_question or not student_answer:
                return False, None, "I didn't hear your answer. Can you try again?"
            
            t = str(original_question).lower()
            # Extract the equation from the original question
            # Pattern for digit operator digit
            m = re.search(r"(\d+(?:\.\d+)?)\s*([\+\-\*\/x×÷]|plus|minus|times|multiplied|divided|add|subtract)\s*(\d+(?:\.\d+)?)", t)
            
            if not m:
                # Try word-based pattern: "multiply X by Y" or "X times Y"
                m2 = re.search(r"(\d+)\s*(?:times|multiplied by|x)\s*(\d+)", t)
                if m2:
                    a, b = float(m2.group(1)), float(m2.group(2))
                    correct_answer = a * b
                else:
                    # Try "add X and Y" or "X plus Y"
                    m3 = re.search(r"(\d+)\s*(?:plus|add|and|\+)\s*(\d+)", t)
                    if m3:
                        a, b = float(m3.group(1)), float(m3.group(2))
                        correct_answer = a + b
                    else:
                        return False, None, "I couldn't figure out the math problem. Let's try a different question."
            else:
                a = float(m.group(1))
                op = m.group(2).lower()
                b = float(m.group(3))
                
                # Calculate correct answer based on operator
                if op in ['+', 'plus', 'add']:
                    correct_answer = a + b
                elif op in ['-', '−', 'minus', 'subtract']:
                    correct_answer = a - b
                elif op in ['*', 'x', '×', 'times', 'multiplied']:
                    correct_answer = a * b
                elif op in ['/', '÷', 'divided', 'over']:
                    if b != 0:
                        correct_answer = a / b
                    else:
                        return False, None, "We can't divide by zero! That's a special rule in math."
                else:
                    return False, None, "I couldn't figure out the operation. Let's try again."
            
            # Extract the number from the student's answer
            # Look for numbers in the student's response
            student_nums = re.findall(r"(\d+(?:\.\d+)?)", student_answer)
            
            if not student_nums:
                # Try to parse word numbers
                word_nums = {
                    'zero': 0, 'one': 1, 'two': 2, 'three': 3, 'four': 4,
                    'five': 5, 'six': 6, 'seven': 7, 'eight': 8, 'nine': 9,
                    'ten': 10, 'eleven': 11, 'twelve': 12, 'thirteen': 13,
                    'fourteen': 14, 'fifteen': 15, 'sixteen': 16, 'seventeen': 17,
                    'eighteen': 18, 'nineteen': 19, 'twenty': 20
                }
                student_lower = student_answer.lower()
                for word, num in word_nums.items():
                    if word in student_lower:
                        student_nums = [str(num)]
                        break
            
            if not student_nums:
                return False, correct_answer, f"I didn't hear a number in your answer. The correct answer is {int(correct_answer) if correct_answer == int(correct_answer) else correct_answer}. Great try though!"
            
            # Use the last number mentioned (usually the answer)
            student_num = float(student_nums[-1])
            
            # Check if correct (allow small floating point tolerance)
            if abs(student_num - correct_answer) < 0.001:
                return True, correct_answer, None  # Correct!
            else:
                # Wrong answer - provide helpful feedback
                correct_display = int(correct_answer) if correct_answer == int(correct_answer) else round(correct_answer, 2)
                student_display = int(student_num) if student_num == int(student_num) else student_num
                return False, correct_answer, f"You said {student_display}, but the answer is {correct_display}. Let me help you: "
                
        except Exception as e:
            print(f"Error evaluating equation answer: {e}")
            return False, None, "I had trouble checking your answer. Let's move on to another question."

    def synthesize_speech(text, out_path='response.wav', sample_rate_hz=24000):
        # Delegate to module-level canonical implementation to avoid duplicate code
        try:
            fn = globals().get('synthesize_speech')
            if fn and fn is not None and callable(fn):
                # Call the module-level implementation
                return fn(text, out_path=out_path, sample_rate_hz=sample_rate_hz)
        except Exception:
            pass
        # If delegation failed, raise to indicate synthesis could not proceed
        raise RuntimeError('synthesize_speech delegation failed')

    # Use the module-level `speak_and_play_text` helper (defined at top-level)
    # to keep behavior consistent. Nested/local definitions removed to avoid
    # accidental shadowing of the canonical helper.

    def TranslateToFil(text):
        try:
            fn = globals().get('TranslateToFil')
            if fn and callable(fn):
                return fn(text)
        except Exception:
            pass
        return text
    def TranslateToEng(text):
        try:
            fn = globals().get('TranslateToEng')
            if fn and callable(fn):
                return fn(text)
        except Exception:
            pass
        return text
    def system_play_audio(folder):
        # Pick a random welcome audio from the folder and attempt to upload it
        # to the ESP playback device and trigger playback there. If the upload
        # or playback on the ESP fails, fall back to local playback.
        files = [f for f in os.listdir(folder) if f.lower().endswith('.wav')]
        if not files:
            print("No WAV files found in the folder.")
            return

        file_to_play = random.choice(files)
        file_path = os.path.join(folder, file_to_play)
        file_title = file_to_play.replace('_', '?')[:-4]  # Replace underscores with question marks
        print("GENTA: Hello, I'm GENTA! " + file_title)

        def upload_and_play_on_esp(wav_path):
            try:
                basename = os.path.basename(wav_path)
                # Use robust upload helper for welcome files
                ok = esp_upload_file(wav_path, endpoint='/upload_welcome', max_retries=2)
                if ok:
                    # Ask the ESP to play the saved welcome file by name
                    return play_welcome_on_esp_by_name(basename)
            except Exception:
                # ESP unreachable - silently fail and use local playback
                pass
            return False

        def play_welcome_on_esp_by_name(basename):
            """Ask the ESP to play an already-uploaded welcome file by name.
            Returns True if the ESP responded 200."""
            try:
                # sanitize locally
                if '/' in basename or '\\' in basename or '..' in basename:
                    return False
                play_url = f'http://{esp_playback_host}/play_welcome?name={basename}'
                r = requests.get(play_url, timeout=2)  # Reduced from 6s to 2s
                return r.status_code == 200
            except Exception:
                # ESP unreachable - silently fail and use local playback
                return False

        # First try to ask the ESP to play an already-uploaded file (no upload).
        try:
            if play_welcome_on_esp_by_name(file_to_play):
                print(f"Asked ESP to play existing welcome file: {file_to_play}")
                return
        except Exception:
            pass

        # If that failed, try uploading and playing.
        try:
            if upload_and_play_on_esp(file_path):
                print(f"Playing welcome audio on ESP at {esp_playback_host}: {file_to_play}")
                return
        except Exception:
            pass

        # Local playback fallback removed — require ESP32 for playback.
        if not esp_playback_host:
            print("⚠ No ESP playback host configured; welcome audio not played.")
        else:
            print("Note: local welcome playback disabled. Audio must be played on ESP32 device.")
    def play_audio(file):
        try:
            fn = globals().get('play_audio')
            if fn and callable(fn):
                return fn(file)
        except Exception:
            pass
        # If delegation fails, return False (no-op fallback)
        return False

    # Helper: move any WAVs from RepeatAudio into a processed folder so they don't retrigger
    def move_repeat_file(single_path: str):
        try:
            src = REPEAT_AUDIO_DIR
            if not src or not os.path.exists(src) or not single_path:
                return
            proc_dir = os.path.join(src, 'processed')
            os.makedirs(proc_dir, exist_ok=True)
            fname = os.path.basename(single_path)
            dst_name = f"{int(time.time())}_{fname}"
            dst = os.path.join(proc_dir, dst_name)
            try:
                shutil.move(single_path, dst)
                print(f"Moved repeat-audio file to processed: {fname} -> {dst_name}")
            except Exception as e:
                print(f"Failed moving repeat audio {single_path}: {e}")
        except Exception:
            pass

    def find_newest_repeat_audio():
        try:
            src = REPEAT_AUDIO_DIR
            if not src or not os.path.exists(src):
                return None
            candidates = [os.path.join(src, f) for f in os.listdir(src) if f.lower().endswith('.wav')]
            # exclude processed folder
            candidates = [c for c in candidates if os.path.isfile(c)]
            if not candidates:
                return None
            candidates.sort(key=lambda p: os.path.getmtime(p), reverse=True)
            return candidates[0]
        except Exception:
            return None

    def process_repeat_audio_file(path):
        """Given a WAV file path in RepeatAudio, convert to 16k mono and run STT; returns transcript or empty string."""
        if not path or not os.path.exists(path):
            return ""

        print(f"Processing repeat-audio file for transcription: {path}")
        ffmpeg_path = shutil.which('ffmpeg') or shutil.which('ffmpeg.exe')
        converted_bytes = None
        try:
            # Try ffmpeg to transcode file to 16k mono WAV bytes
            if ffmpeg_path:
                cmd = [ffmpeg_path, '-y', '-i', path, '-f', 'wav', '-ac', '1', '-ar', '16000', '-acodec', 'pcm_s16le', 'pipe:1', '-hide_banner', '-loglevel', 'error']
                proc = subprocess.Popen(cmd, stdin=subprocess.DEVNULL, stdout=subprocess.PIPE)
                out, _ = proc.communicate(timeout=30)
                converted_bytes = out
            else:
                audio = AudioSegment.from_file(path)
                audio = audio.set_channels(1)
                audio = audio.set_frame_rate(16000)
                buf = io.BytesIO()
                audio.export(buf, format='wav')
                converted_bytes = buf.getvalue()
        except Exception as e:
            print('process_repeat_audio_file: conversion failed:', e)
            converted_bytes = None

        if not converted_bytes:
            return ""

        # Transcribe
        try:
            # Use default client which uses GOOGLE_APPLICATION_CREDENTIALS if set
            client = speech.SpeechClient()
            audio_file = speech.RecognitionAudio(content=converted_bytes)
            config = speech.RecognitionConfig(
                sample_rate_hertz=16000,
                enable_automatic_punctuation=True,
                language_code='fil'
            )
            response = client.recognize(config=config, audio=audio_file)
            complete_text = ""
            if response and hasattr(response, 'results') and len(response.results) > 0:
                for result in response.results:
                    if result.alternatives and len(result.alternatives) > 0:
                        complete_text = result.alternatives[0].transcript
            # save transcription for diagnostics
            try:
                with open(U('transcribed_text.txt'), 'w', encoding='utf-8') as f:
                    f.write(complete_text)
            except Exception:
                pass
            return complete_text
        except Exception as exc:
            print('Transcription error (repeat file):', exc)
            return ""

    def clear_remote_recording(esp_host_local: str = None, audio_url: str = None):
        """Silently clear recording from ESP/proxy. Only tries working endpoints, no noisy logs."""
        # Only try proxy endpoints that we know work (from Flask routes)
        if audio_url:
            try:
                p = urllib.parse.urlparse(audio_url)
                if p.netloc:
                    base = f'{p.scheme}://{p.netloc}'
                    # Only try /clear and /stop (these work, others return 404/502/405)
                    try:
                        requests.get(f'{base}/clear', timeout=2)
                        requests.get(f'{base}/stop', timeout=2)
                    except:
                        pass  # Silent fail - not critical
            except:
                pass
        
        # Skip direct ESP attempts if it's timing out (we already have proxy)
        # This avoids the "Connection to 192.168.50.62 timed out" spam
        return None

    def cleanup_response_artifacts(processed_repeat_path: str = None):
        """Silently remove local response files. No remote cleanup needed (ESP clears on next recording)."""
        # Local files to remove
        try:
            genta_resp_candidate = U('GENTA_response.mp3')
        except Exception:
            genta_resp_candidate = 'GENTA_response.mp3'

        local_candidates = [
            os.path.join('.', 'response.wav'),
            genta_resp_candidate,
            U('transcribed_text.txt')
        ]
        
        # Silently remove local files (only log if successful)
        for fpath in local_candidates:
            try:
                if os.path.exists(fpath):
                    os.remove(fpath)
                    # Only log successful removals, skip errors
            except Exception:
                pass  # Silent fail

        # Remove processed repeat file if provided
        if processed_repeat_path and os.path.exists(processed_repeat_path):
            try:
                os.remove(processed_repeat_path)
            except:
                pass
        
        # Skip all HTTP cleanup - ESP32 auto-clears on next recording
        # This eliminates timeout spam when ESP is busy/unreachable

    # If we prepared a personalized welcome, speak it now (use TTS + playback)
    # Show surprised eyes first (excited to see student!), then transition to happy
    set_oled_expression('surprised')
    time.sleep(1.0)  # Hold surprise for a moment
    
    set_oled_expression('happy')  # Transition to happy eyes
    
    try:
        if 'PERSONAL_WELCOME_MSG' in locals() and PERSONAL_WELCOME_MSG:
            try:
                synthesize_speech(PERSONAL_WELCOME_MSG)
                play_audio('response.wav')
            except Exception:
                pass
            time.sleep(1.0)
    except Exception:
        pass

    # WELCOME: use only the synthesized TTS personal welcome. Do NOT fetch or play
    # any pre-recorded .wav welcome files. This prevents LittleFS accumulation.
    # The synthesized welcome was already played above via synthesize_speech/play_audio.
    # No additional welcome-file playback is performed.
    time.sleep(0.5)
    
    # Keep happy expression visible longer so student can appreciate it
    time.sleep(3.0)  # Display happy eyes for 3 more seconds after welcome
    
    # Return to neutral expression after welcome completes
    set_oled_expression('idle')
    first_iteration = True
    retry_delay = 5
    last_transcribed = None
    last_prompt_time = time.time()  # Track when we last prompted
    PROMPT_COOLDOWN = 300  # Only prompt every 300 seconds (5 minutes) of idle time
    
    while True:
        # Check for state change before each interaction
        if check_for_state_change():
            print("\n[GENTA] State change requested - exiting to restart")
            return  # Exit GENTA() to allow main() to restart with new state
        
        # SILENT WAITING MODE (like Alexa/Google Assistant)
        # Only prompt if:
        # 1. This is NOT the first iteration (already welcomed)
        # 2. It's been more than PROMPT_COOLDOWN seconds since last prompt
        should_prompt = False
        if not first_iteration:
            time_since_last_prompt = time.time() - last_prompt_time
            if time_since_last_prompt > PROMPT_COOLDOWN:
                should_prompt = True
        
        if should_prompt:
            # Play a short follow-up prompt after long idle time
            try:
                # Get student's first name for personalized greeting
                first_name = CURRENT_STUDENT_NAME.split()[0] if CURRENT_STUDENT_NAME and isinstance(CURRENT_STUDENT_NAME, str) else None
                
                pygame.mixer.init()
                prompt_file = r'Pwede pa ba kitang matulungan_.wav'
                if os.path.exists(prompt_file) and not first_name:
                    # Use generic audio file if no name — play on ESP32
                    try:
                        play_audio(prompt_file)
                        print("GENTA: Do you still need help?")
                    except Exception as _e:
                        print(f"⚠ Could not play generic prompt on ESP: {_e}")
                else:
                    # Generate personalized greeting with student's name
                    pygame.mixer.quit()  # Ensure mixer is clean
                    if first_name:
                        synth_text = f"Do you still need help, {first_name}?"
                    else:
                        synth_text = "Do you still need help?"
                    
                    synthesize_speech(synth_text)
                    # Show curious eyes while asking if they need help
                    set_oled_expression('curious')
                    play_audio('response.wav')
                    set_oled_expression('idle')
                    print(f"GENTA: {synth_text}")
                
                # Update last prompt time
                last_prompt_time = time.time()
            except Exception as e:
                print('Skipping prompt playback due to error:', e)
        
        # If we just returned from quiz, defer listening for a short grace
        # period to avoid immediately entering the retry/prompt loop.
        try:
            if '_assist_defer_until' in globals() and _assist_defer_until:
                if time.time() < _assist_defer_until:
                    remaining = _assist_defer_until - time.time()
                    print(f"[Assist] Deferring listening for {remaining:.1f}s after quiz")
                    # Sleep briefly and re-evaluate loop conditions
                    time.sleep(min(0.5, remaining))
                    continue
                else:
                    # Defer elapsed - clear marker so normal behavior resumes
                    _assist_defer_until = None
        except Exception:
            pass

        # Use the shared record_and_transcribe helper (same behavior as QUIZZER)
        print("Listening for student response...")
        
        # Show assisting animation (animated assist text + brief progress) to indicate we're about to listen
        # Only show the assist animation when the orchestrator is actually in Assisting Mode
        try:
            if globals().get('_CURRENT_STATE', '0') == "0":
                set_oled_expression('assist')
            else:
                print("[OLED] Skipping assist animation (not in assisting mode)")
        except Exception:
            pass
        
        # CHECK STATE CHANGE before waiting for recording
        if check_for_state_change():
            print("\n[GENTA] State change detected before recording - exiting to restart")
            set_oled_expression('idle')
            return
        
        # OPTIMIZATION: Removed 3-second wait - record_and_transcribe will handle polling
        processed_repeat_path = None
        try:
            # If test/repeat audio exists, process that file instead of polling the ESP/proxy.
            repeat_candidate = find_newest_repeat_audio()
            if repeat_candidate:
                complete_text = process_repeat_audio_file(repeat_candidate)
                processed_repeat_path = repeat_candidate
                forced_timeout = False
            else:
                # Use the same ESP recording host as QUIZZER so we poll /size and download /recording.wav
                complete_text, forced_timeout = record_and_transcribe(esp_host=esp_record_host, poll_for_recording=True, max_poll_seconds=30, use_english=False)
        except Exception as _e:
            print("Recording/transcription helper failed:", _e)
            complete_text, forced_timeout = "", False
        
        # Return to idle after recording
        set_oled_expression('idle')
        
        # CHECK STATE CHANGE immediately after recording attempt
        if check_for_state_change():
            print("\n[GENTA] State change detected after recording - exiting to restart")
            return

        # Ensure the transcription file exists for downstream tools
        try:
            with open(U('transcribed_text.txt'), 'w', encoding='utf-8') as f:
                f.write(complete_text)
        except Exception:
            pass

        print("You: ", complete_text)

        # Note: we no longer skip duplicate transcriptions here. Instead we will
        # attempt to remove the remote recording after successful processing so
        # the same recording won't be served again.

        # If transcription is empty, retry a few times (mirrors QUIZZER behavior) before giving up.
        MAX_RETRIES = 2
        attempts = 0
        while (not complete_text or not complete_text.strip()) and attempts < MAX_RETRIES and not forced_timeout:
            # CHECK STATE CHANGE during retry loop
            if check_for_state_change():
                print("\n[GENTA] State change detected during retry - exiting to restart")
                return
            
            attempts += 1
            print(f"No transcription received; retrying ({attempts}/{MAX_RETRIES})...")
            # Ask the student to repeat briefly — only play the apology if we actually received audio
            try:
                # Show neutral/thinking expression when asking to repeat
                set_oled_expression('thinking')
                # Only speak the "didn't hear" prompt if the previous attempt contained audio
                if globals().get('LAST_AUDIO_PRESENT', False):
                    # Use a shorter apology to reduce repetition and time-to-listen
                    synth_text = "Sorry, I didn't hear that. Please repeat."
                    synthesize_speech(synth_text)
                    play_audio('response.wav')
                else:
                    # No audio present (likely no one spoke) — be silent and remain curious
                    print('[Silent] No audio present on previous attempt; not prompting apology')
                set_oled_expression('curious')  # Ready to listen again (curious)
            except Exception:
                pass

            # Wait briefly then try recording again
            time.sleep(0.2)  # OPTIMIZATION: Reduced from 0.3s to 0.2s
            try:
                complete_text, forced_timeout = record_and_transcribe(esp_host=esp_record_host, poll_for_recording=True, max_poll_seconds=30, use_english=False)
            except Exception as _e:
                print("Retry recording/transcription failed:", _e)
                complete_text, forced_timeout = "", forced_timeout
            
            # CHECK STATE CHANGE after retry recording
            if check_for_state_change():
                print("\n[GENTA] State change detected after retry - exiting to restart")
                return

            try:
                with open(U('transcribed_text.txt'), 'w', encoding='utf-8') as f:
                    f.write(complete_text)
            except Exception:
                pass

            print("You: ", complete_text)

            if not complete_text or not complete_text.strip():
                # Still empty after retries — handle gracefully and continue to next loop
                print("No valid transcription after retries; skipping to next iteration.")

            # If a recording was attempted (not a forced timeout), inform user we heard them but
            # could not understand. Only say this when the user clearly tried to ask a question.
            if not forced_timeout:
                try:
                    set_oled_expression('thinking')
                    # Only issue apology/prompt if audio was actually present
                    if globals().get('LAST_AUDIO_PRESENT', False):
                        # Short, polite prompt when audio was present but unintelligible
                        synth_text = "Sorry, I didn't hear that. Please repeat."
                        synthesize_speech(synth_text)
                        play_audio('response.wav')
                        # Update last prompt time so we don't spam
                        last_prompt_time = time.time()
                    else:
                        # Silent fail for no audio - behave like Alexa (do nothing)
                        print('[Silent] No audio present; skipping apology prompt')
                    set_oled_expression('idle')
                except Exception:
                    pass
            else:
                # Silent fail for forced timeouts (no recording) - avoid spurious prompts
                print("[Silent] No recording detected (forced timeout), waiting for student to speak...")

            first_iteration = False
            continue

        # Now we have non-empty transcription — send to the model
        try:
            # Show thinking expression after getting the transcription (use inverted arcs)
            set_oled_expression('thinking')
            time.sleep(0.3)  # Brief pause to show thinking
            
            # OPTIMIZATION: Only translate if text appears to be pure Filipino
            # Model can understand Filipino input, so skip translation for speed
            prompt_text = complete_text  # Send directly without translation

            # Guard: if transcription is empty or whitespace (despite retries), skip processing
            # This prevents previous/stale prompt_text from being processed and avoids
            # classifying empty input as non-math or other categories.
            if not prompt_text or not prompt_text.strip():
                print('[Assisting] Empty transcription after retries; skipping processing')
                first_iteration = False
                continue

            # Quick input-side profanity check to avoid sending explicit bad words to the model
            try:
                if is_profane_input(prompt_text):
                    # Local gentle reprimand (do NOT forward to model)
                    set_oled_expression('annoyed')
                    answer = "Please refrain for using inappropriate language. Let's keep our conversation respectful."
                    try:
                        set_oled_expression('annoyed')
                    except Exception:
                        pass
                    try:
                        synthesize_speech(answer)
                        play_audio('response.wav')
                    except Exception:
                        pass
                    first_iteration = False
                    # Skip further processing for this iteration
                    continue
            except Exception:
                # If the quick check fails, fall back to AI-based safety below
                pass

            # Enforce math-only scope: if input is not math-related, reply with scope message
            try:
                if not is_math_related(prompt_text):
                    answer = ("Sorry, I'm focused on Grade 3 mathematics only. "
                              "Please try asking a math problem.")
                    print("[Assisting] Non-math input blocked; replying with scope message")
                    try:
                        with open('genta_action_log.txt', 'a', encoding='utf-8') as _lf:
                            _lf.write(f"{datetime.now().isoformat()} NON-MATH BLOCK: {prompt_text}\n")
                    except Exception:
                        pass

                    try:
                        # Generate TTS and attempt playback (synthesize then upload/play)
                        synthesize_speech(answer)
                        play_audio('response.wav')
                    except Exception as play_err:
                        print(f"⚠ playback failed for non-math reply: {play_err}")
                        try:
                            # fallback to speak_and_play_text helper
                            speak_and_play_text(answer)
                        except Exception:
                            pass

                    first_iteration = False
                    # Skip calling the model for non-math content
                    continue
            except Exception:
                # On any error in math-scope detection, fall back to normal flow
                pass

            # Intercept simple arithmetic equations and teach step-by-step
            try:
                is_eq, eq_expr = detect_simple_equation(prompt_text)
            except Exception:
                is_eq, eq_expr = False, None

            if is_eq:
                # Generate a short step-by-step teaching response (do NOT give final numeric answer)
                try:
                    answer = generate_teaching_for_equation(eq_expr or prompt_text)
                except Exception:
                    answer = "Okay, I'll show step-by-step how to solve that problem. Try it yourself first and tell me the answer."
                # Skip calling the model and proceed to playback using the generated answer
                response = None
                # Store the original equation for later evaluation
                pending_equation_question = eq_expr or prompt_text
            else:
                # Not an equation - clear any pending equation
                pending_equation_question = None
                # OPTIMIZATION: Reduced max_output_tokens from 200 to 150 for faster responses
                # Assisting mode needs quick, concise explanations for Grade 3 students
                # Higher temperature (1.2) encourages faster, more natural generation
                # Prepend a short system-style instruction to bias the assistant to Grade-3
                # math-only, child-friendly, concise Filipino responses. We still have a
                # hard input-side math check above; this instruction reinforces that
                # behavior so the model is less likely to produce off-topic replies.
                tutor_instruction = (
                    "You are an elementary Grade 3 mathematics tutor for 8-year-old students. "
                    "Answer ONLY mathematics questions aligned with Grade 3 MELCs. Keep answers very short (1-3 simple sentences), "
                    "use child-friendly language and examples, and prefer English. If the question is NOT a math question, reply exactly: "
                    "'Sorry, I'm only a math teacher. I cannot answer that question. Please ask about mathematics.'"
                )
                wrapped_prompt = tutor_instruction + "\n\nQ: " + prompt_text
                
                # Start "please wait" timer for slow model calls
                wait_timer = start_please_wait_timer(
                    delay_seconds=5.0,
                    message="Please wait, I'm thinking about the answer..."
                )
                
                try:
                    response = chat.send_message(
                        wrapped_prompt,
                        generation_config={
                            'temperature': 1.2,
                        }
                    )
                finally:
                    # Cancel timer as soon as we get response
                    wait_timer.set()
                # Defensive retry: if the model returned a truncated response (MAX_TOKENS)
                # or an empty chat reply, attempt a compact fallback using a direct
                # non-chat generation with a larger `max_output_tokens`. This helps
                # when the session history (large MELCs or examples) makes the
                # effective prompt too big and causes the model to hit its output
                # token limit and return no text.
                try:
                    cand = (getattr(response, 'candidates', None) or [None])[0]
                    finish = getattr(cand, 'finish_reason', None)
                    empty_chat = False
                    try:
                        # Some SDKs return no candidates but a top-level text attribute
                        if not getattr(response, 'candidates', None) and not getattr(response, 'text', None):
                            empty_chat = True
                    except Exception:
                        empty_chat = False

                    if finish in ('MAX_TOKENS', 2) or empty_chat:
                        print('[Model] Detected truncated/empty chat response (finish_reason=MAX_TOKENS). Retrying with compact prompt...')
                        # Build a compact instruction-only prompt to avoid the huge session history
                        compact_instr = (
                            "You are an elementary Grade 3 mathematics tutor for 8-year-old students. "
                            "Keep answers extremely short (1-3 simple sentences) and in simple English. "
                            "Answer the question only."
                        )
                        compact_prompt = compact_instr + "\n\nQ: " + prompt_text
                        
                        # Start wait timer for retry attempt
                        retry_wait_timer = start_please_wait_timer(
                            delay_seconds=3.0,
                            message="Please wait..."
                        )
                        
                        try:
                            # Use generate_content fallback without an explicit max_output_tokens
                            # so the provider may choose an appropriate output length.
                            retry = None
                            try:
                                retry = model.generate_content(compact_prompt, temperature=0.9)
                            except TypeError:
                                # Some provider wrappers expect a generation_config dict instead
                                retry = model.generate_content(compact_prompt, generation_config={'temperature': 0.9})

                            if retry is not None:
                                response = retry
                                print('[Model] Retry returned a response; will use it instead of the truncated chat reply')
                        except Exception as rexc:
                            print(f"[Model] Retry attempt failed: {rexc}")
                        finally:
                            retry_wait_timer.set()
                except Exception:
                    # If diagnostics fail, continue normally and let downstream handle empty responses
                    pass
        except ValueError as ve:
            print('Model rejected empty input after checks:', ve)
            # Inform user and continue
            set_oled_expression('incorrect')  # Show sad face for error
            try:
                # OPTIMIZATION: Text is already Filipino - no need to translate
                synth_text = "There was a problem processing your answer. Please try again later."
                synthesize_speech(synth_text)
                play_audio('response.wav')
            except Exception:
                pass
            set_oled_expression('idle')
            first_iteration = False
            continue
        
        # Check if the prompt was blocked by safety filters (defensive)
        if response is not None:
            try:
                candidates = getattr(response, 'candidates', None)
                prompt_feedback = getattr(response, 'prompt_feedback', None)
                if (not candidates) and prompt_feedback and getattr(prompt_feedback, 'block_reason', None) == 'SAFETY':
                    curse_response = "Please refrain from using bad words."
                    print("GENTA:", curse_response)
                    set_oled_expression('incorrect')  # Sad face for inappropriate content
                    try:
                        synthesize_speech(curse_response)
                        play_audio('response.wav')
                    except Exception:
                        pass
                    set_oled_expression('idle')
                    first_iteration = False
                    continue
            except Exception:
                # If safety check parsing fails, continue normally (do not crash)
                pass

        # TRANSLATION / extract text robustly from response
        def _extract_model_text(resp):
            """Robust extractor for model response text across SDK variants.
            This attempts multiple known shapes and then performs a shallow
            recursive search for the first non-empty string value.
            Returns `str` or `None` when no usable text found.
            """
            try:
                # 1) Common: candidates -> content -> parts -> [0].text
                candidates = getattr(resp, 'candidates', None)
                if candidates:
                    for cand in candidates:
                        # content may be object-like or dict-like
                        content = getattr(cand, 'content', None)
                        if content is not None:
                            parts = getattr(content, 'parts', None)
                            if parts and len(parts) > 0:
                                first = parts[0]
                                if hasattr(first, 'text') and first.text:
                                    return first.text
                                if isinstance(first, dict) and first.get('text'):
                                    return first.get('text')
                            # fallback: content.text or content string
                            txt = getattr(content, 'text', None)
                            if isinstance(txt, str) and txt:
                                return txt
                            if isinstance(content, str) and content:
                                return content

                        # cand may directly have text
                        cand_text = getattr(cand, 'text', None)
                        if isinstance(cand_text, str) and cand_text:
                            return cand_text

                        # some SDKs present message field
                        msg = getattr(cand, 'message', None)
                        if isinstance(msg, str) and msg:
                            return msg

                # 2) Top-level simple attributes
                for attr in ('text', 'output_text', 'content', 'message', 'last'):
                    v = getattr(resp, attr, None)
                    if isinstance(v, str) and v:
                        return v

                # 3) resp may be a dict-like object with nested strings
                # Try a shallow recursive search for a usable string value
                def _shallow_find_string(obj, depth=0, max_depth=3):
                    if depth > max_depth:
                        return None
                    # string
                    if isinstance(obj, str) and obj.strip():
                        return obj
                    # dict-like
                    if isinstance(obj, dict):
                        for k, val in obj.items():
                            if isinstance(val, str) and val.strip():
                                return val
                            res = _shallow_find_string(val, depth + 1, max_depth)
                            if res:
                                return res
                    # list/tuple
                    if isinstance(obj, (list, tuple)):
                        for item in obj:
                            if isinstance(item, str) and item.strip():
                                return item
                            res = _shallow_find_string(item, depth + 1, max_depth)
                            if res:
                                return res
                    # object with __dict__
                    try:
                        od = getattr(obj, '__dict__', None)
                        if isinstance(od, dict):
                            return _shallow_find_string(od, depth + 1, max_depth)
                    except Exception:
                        pass
                    return None

                # Try shallow search on resp itself and its __dict__
                maybe = None
                try:
                    # If resp exposes a to_dict-like method, prefer it
                    if hasattr(resp, 'to_dict') and callable(getattr(resp, 'to_dict')):
                        try:
                            maybe = _shallow_find_string(resp.to_dict())
                        except Exception:
                            maybe = None
                    if not maybe:
                        maybe = _shallow_find_string(getattr(resp, '__dict__', resp))
                except Exception:
                    maybe = None

                if isinstance(maybe, str) and maybe.strip():
                    return maybe

                # Nothing usable found
                return None
            except Exception:
                # Defensive: don't crash the main loop; caller will handle None
                print('extract_model_text: no usable text in model response; attempting fallback')
                return None

        # If we intercepted the prompt and generated an 'answer' already, keep it.
        if response is None and 'answer' in locals():
            # preserving pre-generated teaching answer
            pass
        else:
            answer = _extract_model_text(response) or ""

        # If the chat response contains no usable text, attempt a direct
        # non-chat generation fallback (some SDKs/models return an empty
        # parts list for chat responses but generate_content() may return
        # a usable top-level text). This improves robustness against
        # inconsistent SDK shapes.
        if not answer:
            # Removed slow fallback `model.generate_content()` attempts to reduce
            # latency. Rely on the primary chat response, and if empty, respond
            # with a short polite scope/failure message rather than issuing a
            # second expensive network/model call.
            print('Warning: model response contained no text. Skipping slow fallback generation.')
            answer = ""

        # Enforce math-only in the model's reply as a final safety: if the
        # model returned something that does not appear math-related, override
        # it with the polite scope message instead of attempting more generations.
        if answer and not is_math_related(answer):
            try:
                print('[Assisting] Model answer out-of-scope; overriding with scope message')
                with open('genta_action_log.txt', 'a', encoding='utf-8') as _lf:
                    _lf.write(f"{datetime.now().isoformat()} MODEL_OUT_OF_SCOPE_REPLY: {prompt_text} -> {answer}\n")
            except Exception:
                pass
            answer = "Sorry, I'm only a math teacher. I cannot answer that question. Please ask about mathematics."

        if not answer:
            # Handle cases where there are no candidates or no text. Log debug info
            try:
                print('Warning: model response contained no text. Raw response summary:')
                # print some useful diagnostics without flooding the terminal
                try:
                    print('  candidates_count=', len(getattr(response, 'candidates', []) or []))
                except Exception:
                    pass
                try:
                    first_cand = (getattr(response, 'candidates', None) or [None])[0]
                    if first_cand is not None:
                        print('  first_candidate_repr=', str(first_cand)[:400])
                        print('  first_candidate_finish_reason=', getattr(first_cand, 'finish_reason', None))
                except Exception:
                    pass
            except Exception:
                pass
            error_response = "I'm sorry, I could not process that. Please try again."
            print("GENTA:", error_response)
            # OPTIMIZATION: Use Filipino directly - no translation needed
            error_response_fil = "Sorry, I couldn't process your question. Please try again."
            synthesize_speech(error_response_fil)
            play_audio('response.wav')
            continue
        
        ConvertedAnswer = convert_apostrophe(answer)
        # Simplify and clean the model's text for grade-3 clarity and TTS friendliness
        try:
            ConvertedAnswer = simplify_for_grade3(ConvertedAnswer)
        except Exception:
            pass

        # OPTIMIZATION: Decide whether translation to Filipino is needed.
        # Use word-boundary checks for common Filipino markers and require multiple
        # matches to avoid false positives (e.g., English words containing 'ang').
        needs_translation = True
        filipino_indicators = ['ang', 'ng', 'sa', 'ay', 'mga', 'ba', 'po', 'ko', 'mo', 'siya']
        try:
            text_low = ConvertedAnswer.lower() if ConvertedAnswer else ''
            matches = 0
            for word in filipino_indicators:
                if re.search(r'\b' + re.escape(word) + r'\b', text_low):
                    matches += 1

            # Heuristic: if we find 2 or more Filipino indicator words, treat as Filipino
            if matches >= 2:
                needs_translation = False
                TranslatedResponse = ConvertedAnswer
                print("OPTIMIZATION: Response appears to be Filipino (heuristic), skipping translation")
            else:
                needs_translation = True
        except Exception:
            needs_translation = True

        if needs_translation:
            try:
                TranslatedResponse = TranslateToFil(ConvertedAnswer)
            except Exception as e:
                print(f"TranslateToFil failed: {e}; using original text")
                TranslatedResponse = ConvertedAnswer
        
        # ============ COZMO-STYLE EXPRESSION DETECTION ============
        # Detect emotional context in AI response and show appropriate Cozmo eye expression
        response_lower = TranslatedResponse.lower()
        
        # Priority 1: Check if response is rejecting non-math question
        if any(phrase in response_lower for phrase in ['only a math teacher', 'cannot answer', 'ask about mathematics']):
            set_oled_expression('sleepy')  # Closed/droopy eyes for out-of-scope questions

        # Priority 2: Exceptional excitement/celebration (strongest positive)
        elif any(word in response_lower for word in ['perfect', 'amazing', 'excellent', 'wonderful', 'fantastic', '100%', 'great']):
            set_oled_expression('excited')  # Rapid shimmy for exceptional performance! 🎉

        # Priority 3: Very happy/delighted responses
        elif any(word in response_lower for word in ['happy', 'delightful', 'hooray', 'yay', 'funny']):
            set_oled_expression('glee')  # Happy bounce for joyful moments

        # Priority 4: Amazement/wonder
        elif any(word in response_lower for word in ['wow', 'really', 'amazing', 'oh my', 'incredible', 'surprising']):
            set_oled_expression('awe')  # Wide amazed eyes for surprising facts! 😲

        # Priority 5: Surprise/unexpected
        elif any(word in response_lower for word in ['suddenly', 'shocking', 'unexpected']):
            set_oled_expression('surprised')  # Sudden wide eyes for surprises

        # Priority 6: Curiosity/pondering (asking questions or thinking aloud)
        elif any(phrase in response_lower for phrase in ['hmm', 'maybe', 'perhaps', 'how', 'why', 'interesting', 'let me think']):
            set_oled_expression('curious')  # Scanning curious look when pondering 🤔

        # Priority 7: Skeptical/doubtful
        elif any(phrase in response_lower for phrase in ['are you sure', 'i doubt', 'questionable', 'not sure']):
            set_oled_expression('skeptical')  # One eye smaller, questioning look

        # Priority 8: Focused/concentrating
        elif any(word in response_lower for word in ['focus', 'concentrate', 'listen', 'attention', 'carefully']):
            set_oled_expression('focused')  # Narrow intense eyes for concentration

        # Priority 9: Student answer is CORRECT (positive reinforcement)
        elif any(word in response_lower for word in ['correct', 'very good', 'great', 'exactly', 'right']):
            set_oled_expression('happy')  # Classic happy eyes when correct! 😊

        # Priority 10: Student answer is INCORRECT but encouraging
        elif any(phrase in response_lower for phrase in ['incorrect', 'wrong', 'not quite', 'almost']):
            # Check if response is encouraging retry (use worried) or harsh (sad)
            if any(word in response_lower for word in ['try again', 'repeat', 'try once more', 'give it another go']):
                set_oled_expression('worried')  # Gentle worried look for encouraging retry
            else:
                set_oled_expression('sad')  # Sad eyes when clearly incorrect 😞

        # Priority 11: Frustration (repeated mistakes or difficulty)
        elif any(phrase in response_lower for phrase in ['again and again', 'tiring', 'difficult', 'struggle']):
            set_oled_expression('frustrated')  # Droopy eyes for frustration

        # Priority 12: Confusion/uncertainty
        elif any(word in response_lower for word in ['confused', 'unclear', 'i do not understand', "i don't get it"]):
            set_oled_expression('skeptical')  # Side-eye for confusion
        
        # Default: neutral idle expression for general explanations
        else:
            set_oled_expression('idle')  # Neutral expression for informational responses
        
        # Play the full response using the module-level playback helper (no chunking)
        try:
            print("GENTA:", TranslatedResponse)
            played_ok = False
            try:
                top = globals().get('speak_and_play_text')
                if callable(top):
                    played_ok = top(TranslatedResponse)
                else:
                    # Fallback to synth+play if canonical helper not available
                    synthesize_speech(TranslatedResponse)
                    played_ok = play_audio('response.wav')

                if not played_ok:
                    print('⚠ speak_and_play_text returned False; attempting direct synth+play fallback')
                    try:
                        synthesize_speech(TranslatedResponse)
                        play_ok = play_audio('response.wav')
                        if not play_ok:
                            print('⚠ play_audio reported failure')
                    except Exception as fall_e:
                        print(f'⚠ direct fallback playback failed: {fall_e}')
            except Exception as se:
                print(f"⚠ speak_and_play_text exception: {se}")
                try:
                    synthesize_speech(TranslatedResponse)
                    play_audio('response.wav')
                except Exception as e:
                    print(f"⚠ direct synth+play also failed: {e}")
        except Exception as e:
            print(f"⚠ speak_and_play_text failed: {e}")
        
        # Return to idle expression after speaking
        set_oled_expression('idle')
        
        # ========== EQUATION MODE: Wait for student answer ==========
        # If we just taught the student how to solve an equation, wait for their answer
        if 'pending_equation_question' in dir() and pending_equation_question:
            print("\n[Equation Mode] Waiting for student's answer to the equation...")
            set_oled_expression('curious')
            
            try:
                # Listen for student's answer to the equation
                student_eq_answer, _ = record_and_transcribe(
                    esp_host=esp_record_host,
                    poll_for_recording=True,
                    max_poll_seconds=15,  # Give more time for thinking
                    use_english=False
                )
                set_oled_expression('thinking')
                
                if student_eq_answer:
                    print(f"[Equation Mode] Student answered: {student_eq_answer}")
                    
                    # Evaluate the student's answer
                    is_correct, correct_answer, feedback = evaluate_equation_answer(
                        pending_equation_question, 
                        student_eq_answer
                    )
                    
                    if is_correct:
                        # Correct answer - celebrate!
                        praise_responses = [
                            "That's correct! Great job! You're really good at math!",
                            "Yes! You got it right! Excellent work!",
                            "Perfect! That's the right answer! You're a math star!",
                            "Correct! Well done! I knew you could do it!",
                            "Amazing! That's exactly right! Keep up the great work!"
                        ]
                        import random
                        praise = random.choice(praise_responses)
                        print(f"GENTA: {praise}")
                        set_oled_expression('excited')
                        synthesize_speech(praise)
                        play_audio('response.wav')
                    else:
                        # Wrong answer - provide helpful feedback
                        set_oled_expression('worried')
                        if feedback:
                            print(f"GENTA: {feedback}")
                            synthesize_speech(feedback)
                            play_audio('response.wav')
                        
                        # Now give a helpful hint based on the operation
                        if correct_answer is not None:
                            t = pending_equation_question.lower()
                            hint = ""
                            correct_display = int(correct_answer) if correct_answer == int(correct_answer) else round(correct_answer, 2)
                            
                            if any(op in t for op in ['times', 'multiply', '*', 'x', '×']):
                                hint = f"When we multiply, we add groups together. The answer is {correct_display}. Let's try another problem!"
                            elif any(op in t for op in ['plus', 'add', '+']):
                                hint = f"When we add, we put numbers together. The answer is {correct_display}. You'll get the next one!"
                            elif any(op in t for op in ['minus', 'subtract', '-']):
                                hint = f"When we subtract, we take away. The answer is {correct_display}. Keep practicing!"
                            elif any(op in t for op in ['divide', 'divided', '/', '÷']):
                                hint = f"When we divide, we split into equal groups. The answer is {correct_display}. Nice try!"
                            else:
                                hint = f"The correct answer is {correct_display}. Don't worry, math takes practice!"
                            
                            print(f"GENTA: {hint}")
                            set_oled_expression('happy')
                            synthesize_speech(hint)
                            play_audio('response.wav')
                else:
                    # No answer detected
                    print("[Equation Mode] No answer detected from student")
                    no_answer_msg = "I didn't hear your answer. That's okay! Take your time and ask me another math question when you're ready."
                    print(f"GENTA: {no_answer_msg}")
                    synthesize_speech(no_answer_msg)
                    play_audio('response.wav')
                    
            except Exception as eq_err:
                print(f"[Equation Mode] Error: {eq_err}")
            
            set_oled_expression('idle')
            # Clear the pending equation
            pending_equation_question = None
            
        else:
            # ========== NORMAL MODE: Offer to repeat the answer ==========
            # Ask student if they want the answer repeated for better comprehension
            repeat_prompt = "Would you like me to repeat my answer?"
            print("GENTA:", repeat_prompt)
            synthesize_speech(repeat_prompt)
            play_audio('response.wav')

            # Listen for student's response (yes/no to repeat)
            set_oled_expression('curious')
            try:
                repeat_response_text, _ = record_and_transcribe(
                    esp_host=esp_record_host,
                    poll_for_recording=True,
                    max_poll_seconds=10,
                    use_english=False
                )
                set_oled_expression('idle')

                # Check if student wants a repeat
                if repeat_response_text:
                    repeat_lower = repeat_response_text.lower()
                    wants_repeat = any(word in repeat_lower for word in ['yes', 'y', 'yeah', 'repeat', 'again'])

                    if wants_repeat:
                        print("Student requested repeat")
                        set_oled_expression('thinking')
                        # Repeat the previous answer
                        synthesize_speech(TranslatedResponse)
                        print("GENTA (repeating):", TranslatedResponse)
                        play_audio('response.wav')
                        set_oled_expression('idle')
                    else:
                        print("Student declined repeat")
                else:
                    # No response detected - continue normally
                    print("No repeat response detected, continuing...")
            except Exception as e:
                print(f"Error during repeat prompt: {e}")
                set_oled_expression('idle')
        
        # Reset the prompt cooldown timer to allow immediate prompt after answering
        # Subtract PROMPT_COOLDOWN so next iteration will trigger the prompt
        last_prompt_time = time.time() - PROMPT_COOLDOWN
        
        # After successfully answering, mark this transcription as processed and move the processed repeat file (if any)
        try:
            last_transcribed = complete_text
            if processed_repeat_path:
                try:
                    move_repeat_file(processed_repeat_path)
                except Exception:
                    pass
        except Exception:
            pass

        # Best-effort: clear/delete the remote recording so the same audio isn't served again
        try:
            # prefer esp_record_host if configured, otherwise pass the audio_raw_url for derived attempts
            clear_remote_recording(esp_host_local=esp_record_host, audio_url=audio_raw_url)
            try:
                cleanup_response_artifacts(processed_repeat_path)
            except Exception:
                pass
        except Exception:
            pass
        first_iteration = False


def confirm_quiz_readiness():
    """
    Ask student if they're ready to take the quiz/assessment.
    Returns True if student confirms (yes), False otherwise.
    Uses direct Google TTS since this is called before synthesize_speech() is defined.
    """
    print("\n" + "="*70)
    print("📝 QUIZ MODE CONFIRMATION")
    print("="*70 + "\n")
    # If a state-change processing animation is running, stop it now before showing confirmation
    try:
        global _state_change_stop_event
        if _state_change_stop_event and not _state_change_stop_event.is_set():
            _state_change_stop_event.set()
            print("[State Monitor] Stopped processing animation before confirmation prompt")
    except Exception:
        pass
    
    # Helper function for TTS synthesis (since synthesize_speech not available yet)
    def _tts_and_play(text):
        """Generate and play TTS audio using Google Cloud TTS."""
        try:
            client = texttospeech_v1.TextToSpeechClient()
            # Use an English female neural voice for quiz confirmations
            voice = texttospeech_v1.VoiceSelectionParams(
                name='en-US-Neural2-C',
                language_code='en-US'
            )
            # Ensure confirmation prompts use 24kHz to match ESP32 playback
            audio_config = texttospeech_v1.AudioConfig(
                audio_encoding=texttospeech_v1.AudioEncoding.LINEAR16,
                pitch=6.00,
                sample_rate_hertz=24000
            )
            response = client.synthesize_speech(
                input=texttospeech_v1.SynthesisInput(text=text),
                voice=voice,
                audio_config=audio_config
            )
                
            # Save to response.wav
            with open('response.wav', 'wb') as out:
                out.write(response.audio_content)
            # Attempt to upload to ESP32 (autoplay on device)
            try:
                ok = esp_upload_file('response.wav', endpoint='/upload', max_retries=3)
                if ok:
                    return True
                else:
                    print('⚠ TTS generated but ESP upload failed (confirm prompt)')
                    return False
            except Exception as _e:
                print(f'⚠ TTS generated but ESP upload exception: {_e}')
                return False
        except Exception as e:
            print(f"⚠ TTS error: {e}")
            return False
    
    # Ask confirmation
    try:
        confirmation_text = "Are you ready for the quiz? Say 'yes' if you are ready, or 'no' if you are not."
        # Use chunked TTS helper for faster start and consistent behavior
        ok = speak_and_play_text(confirmation_text)
        print(f"GENTA: {confirmation_text}")
        # Small pause to allow the recorder device to arm/listen after prompt playback
        try:
            time.sleep(0.35)
        except Exception:
            pass
        # If upload/playback to ESP succeeded, explicitly ask the recorder to show
        # the confirmation text so the OLED displays the two-line prompt instead
        # of a generic processing dot. Fall back to idle expression on failure.
        try:
            if ok:
                # Avoid sending the long two-line typewriter text to the OLED for
                # confirmation prompts. It is unnecessary and interrupts flow on
                # the device. Show a simple listening pose instead.
                set_oled_expression('listening')
            else:
                set_oled_expression('idle')
        except Exception:
            set_oled_expression('idle')
    except Exception as e:
        print(f"⚠ Could not play confirmation prompt: {e}")
        # Still continue to listen for response
    
    # Listen for student's response (shorter, snappier confirmation)
    max_attempts = 2
    for attempt in range(1, max_attempts + 1):
        print(f"\n[Confirmation] Listening for response (attempt {attempt}/{max_attempts})...")
        try:
            # Show curious expression while waiting for confirmation
            set_oled_expression('curious')

            # Reset proxy notification queue to ensure we wait for a NEW recording
            try:
                _http_session.get('http://127.0.0.1:5000/reset_recording_notification', timeout=1.0)
            except Exception:
                pass

            # Record and transcribe student's response
            # Give the recorder a slightly longer event-driven wait for student replies
            response_text, forced_timeout = record_and_transcribe(
                esp_host=esp_record_host,
                poll_for_recording=True,
                max_poll_seconds=15,
                use_event_driven=True,
                use_english=False
            )

            set_oled_expression('idle')

            if not response_text or not response_text.strip():
                print("[Confirmation] No response received, asking again...")
                try:
                    retry_text = "I didn't hear you. Are you ready? Say 'yes' or 'no'."
                    try:
                        speak_and_play_text(retry_text)
                    except Exception:
                        pass
                    print(f"GENTA: {retry_text}")
                except Exception:
                    pass
                time.sleep(0.5)
                continue

            # Clean and check response
            response_lower = response_text.lower().strip()
            print(f"[Confirmation] Student said: '{response_text}'")

            # Check for affirmative responses (yes/ready) and negative responses
            affirmative_words = ['yes', 'y', 'yeah', 'ready', 'okay', 'ok']
            negative_words = ['no', 'not', 'nope', 'later']

            # Check if any affirmative word is in the response
            is_yes = any(word in response_lower for word in affirmative_words)
            is_no = any(word in response_lower for word in negative_words)

            if is_yes and not is_no:
                print("[Confirmation] ✓ Student confirmed readiness!")
                try:
                    set_oled_expression('happy')
                    confirm_text = "Great! Let's start the quiz. Good luck!"
                    try:
                        speak_and_play_text(confirm_text)
                    except Exception:
                        pass
                    print(f"GENTA: {confirm_text}")
                    set_oled_expression('idle')
                except Exception:
                    pass
                time.sleep(1)
                return True

            elif is_no:
                print("[Confirmation] ✗ Student declined quiz")
                try:
                    set_oled_expression('idle')
                    decline_text = "Okay, returning to assisting mode. Press the button when you're ready."
                    try:
                        speak_and_play_text(decline_text)
                    except Exception:
                        pass
                    print(f"GENTA: {decline_text}")
                except Exception:
                    pass
                time.sleep(1)
                return False

            else:
                # Unclear response - ask again
                print("[Confirmation] Unclear response, asking for clarification...")
                try:
                    clarify_text = "I didn't understand. Say 'yes' if you are ready, or 'no' if not."
                    try:
                        speak_and_play_text(clarify_text)
                    except Exception:
                        pass
                    print(f"GENTA: {clarify_text}")
                except Exception:
                    pass
                time.sleep(0.5)
                continue

        except Exception as e:
            print(f"[Confirmation] Error during attempt {attempt}: {e}")
            time.sleep(0.5)
            continue
    
    # If we reach here, max attempts exceeded - default to declining quiz
    print("[Confirmation] Max attempts reached, returning to Assisting Mode")
    try:
        timeout_text = "Okay, returning to assisting mode for now."
        try:
            speak_and_play_text(timeout_text)
        except Exception:
            pass
        print(f"GENTA: {timeout_text}")
    except Exception:
        pass
    return False


def QUIZZER(start_loading=True):
    global CURRENT_STUDENT_ID, CURRENT_TEACHER_ID, CURRENT_TEACHER_NAME, CURRENT_STUDENT_NAME
    # Report/animation control globals - declare early so any use in this function
    # does not occur before the global statement (avoids SyntaxError).
    global _report_creation_stop_event, _REPORT_CREATION_ACTIVE, _report_progress, _REPORT_COMPLETION_SHOWN
    # Ensure the completion-shown flag exists
    try:
        if '_REPORT_COMPLETION_SHOWN' not in globals():
            globals()['_REPORT_COMPLETION_SHOWN'] = False
    except Exception:
        globals()['_REPORT_COMPLETION_SHOWN'] = False
    
    print("\n" + "="*70)
    print("📝 QUIZ MODE - Retrieving Student Information")
    print("="*70)
    
    # CRITICAL: DISABLE state button during quiz to prevent interruptions
    print("\n[Quiz] Disabling state button - student cannot change modes during quiz...")
    try:
        disable_button_resp = requests.get('http://127.0.0.1:5000/disable_state_button', timeout=5)
        if disable_button_resp.status_code == 200:
            print("✓ State button DISABLED - quiz mode locked")
        else:
            print(f"⚠ Could not disable button: status {disable_button_resp.status_code}")
    except Exception as e:
        print(f"⚠ Could not disable state button: {e}")
    
    # IMPORTANT: Retrieve student info to get teacher_id BEFORE loading quiz
    # The teacher_id determines which questions to load from the database
    if not CURRENT_TEACHER_ID:
        print("\n⚠ WARNING: No teacher ID available!")
        print("Student must enter LRN in Assisting Mode first to get teacher info.")
        print("Returning to Assisting Mode...")
        
        # Re-enable button before exiting
        try:
            requests.get('http://127.0.0.1:5000/enable_state_button', timeout=5)
            print("✓ State button re-enabled")
        except Exception:
            pass
        
        # Set state back to 0 (Assisting Mode)
        try:
            set_state_url = "https://nonbasic-bob-inimical.ngrok-free.dev/set_state"
            requests.get(f"{set_state_url}?value=0", timeout=15)
            print("✓ State reset to Assisting Mode")
        except Exception as e:
            print(f"⚠ Could not reset state: {e}")
        
        return  # Exit Quiz Mode
    
    # Display student and teacher info
    print(f"\n✓ Student: {CURRENT_STUDENT_NAME} (LRN: {CURRENT_STUDENT_ID})")
    print(f"✓ Teacher: {CURRENT_TEACHER_NAME} (ID: {CURRENT_TEACHER_ID})")
    print(f"✓ Loading quiz questions from Teacher ID: {CURRENT_TEACHER_ID}")
    print("="*70 + "\n")
    
    with open(conversation_file_path, 'w') as conv_file:
        conv_file.write("Conversation Log:\n\n")

    messages = [
        # system message first, it helps set the behavior of the assistant
        {"role": "user", "content": """ 
        You are an experienced elementary school teacher writing a professional student assessment report.
        
        Analyze this quiz conversation and write a comprehensive report with the following sections:
        
        1. EXECUTIVE SUMMARY: Brief overview of the student's performance
        2. STRENGTHS: Specific areas where the student demonstrated mastery (be specific with examples)
        3. AREAS FOR IMPROVEMENT: Specific concepts that need reinforcement (be constructive and encouraging)
        4. RECOMMENDED STRATEGIES: Concrete, actionable teaching strategies and interventions
        5. LESSON PLAN SUGGESTIONS: Brief outline of topics to focus on
        
        IMPORTANT GUIDELINES:
        - Write in professional academic language as if a human teacher wrote it
        - Use specific examples from the quiz conversation
        - Be encouraging and constructive in tone
        - Avoid generic AI-style phrases like "Based on the analysis" or "It is recommended that"
        - Use natural educator language: "The student shows...", "I observed that...", "We can help by..."
        - Do NOT use asterisks, markdown formatting, or any AI markers
        - Write in clear paragraphs with proper grammar
        - Be specific with student's actual responses and patterns
        """},
    ]
    module_messages = [
        # system message first, it helps set the behavior of the assistant
        {"role": "user", "content": """
        You are creating a FUN and ENGAGING learning module for a Grade 3 student (8-9 years old).
        
        Based on this quiz conversation, identify the student's mathematics weaknesses and create an exciting e-book that:
        
        ✨ MAKE IT FUN! Use:
        - Friendly, encouraging language ("Great job!", "Let's try this!", "You're doing amazing!")
        - Emojis occasionally (😊 🎉 ⭐ 🌟 💡 🎯)
        - Short, simple sentences
        - Colorful descriptive words (amazing, wonderful, fantastic, cool, awesome)
        
        📚 CONTENT STRUCTURE:
        1. Welcome message (encouraging and fun!)
        2. Why This Topic is Cool (real-world connections kids care about)
        3. Let's Learn the Basics! (explain like talking to an 8-year-old)
        4. Key Concepts (break down into tiny, easy steps)
        5. Remember These Tips! (simple rules/tricks)
        
        🎯 TEACHING STYLE:
        - Use kid-friendly examples (toys, candies, games, animals, cartoons)
        - Break complex ideas into tiny steps
        - Include encouraging phrases throughout
        - Use analogies children understand (like building blocks, sharing snacks)
        - Make math feel like a game or adventure
        
        ⚠️ AVOID:
        - Long paragraphs (keep it bite-sized!)
        - Big scary words
        - Boring academic tone
        - AI-style language
        """},
    ]    
    module_messages2 = [
        # system message first, it helps set the behavior of the assistant
        {"role": "user", "content": """
        Continue the fun learning module for our Grade 3 student! 
        
        Create the PRACTICE section with:
        
        🌈 EXAMPLES SECTION: "Let's See It In Action!"
        - 3-4 real-world examples kids love (candy sharing, toy counting, pizza slicing, game scores)
        - Show EVERY step (like you're doing it together)
        - Use pictures in words ("Imagine you have 5 red apples...")
        - Add encouraging comments ("See? Easy peasy!" "You got this!" "Look at you, math star!")
        
        🎮 PRACTICE PROBLEMS: "Your Turn to Shine!"
        - Create 5 FUN problems per topic
        - Use contexts kids enjoy (video games, cartoons, pets, playground, birthday parties)
        - Start easy, gradually increase difficulty
        - Add hints in parentheses (Hint: Remember what we learned about...)
        - Space for them to write answers
        
        ✨ KEEP IT EXCITING:
        - Use emojis (but not too many!)
        - Short sentences
        - Encouraging language
        - Make it feel like play, not work
        - Add "You're awesome!" type messages between problems
        
        Remember: This is for an 8-year-old! Make math feel like the coolest subject ever! 🎉
        """},
    ]     
    def log_conversation(log_entry):
        with open(conversation_file_path, 'a') as conv_file:
            conv_file.write(log_entry + "\n")
    # NOTE: synthesize_speech and play_audio are defined earlier in this file
    # to provide chunking and playback behavior. Older duplicate definitions
    # that used direct TTS calls were removed to avoid conflicts.
    def TranslateToFil(text):
        try:
            fn = globals().get('TranslateToFil')
            if fn and callable(fn):
                return fn(text)
        except Exception:
            pass
        return text
    def TranslateToEng(text):
        try:
            fn = globals().get('TranslateToEng')
            if fn and callable(fn):
                return fn(text)
        except Exception:
            pass
        return text
    def play_audio(file):
        try:
            fn = globals().get('play_audio')
            if fn and callable(fn):
                return fn(file)
        except Exception:
            pass
        return False
    def ordinal(n):
        suffix = ['th', 'st', 'nd', 'rd', 'th'][min(n % 10, 4)]
        if 11 <= (n % 100) <= 13:
            suffix = 'th'
        return str(n) + suffix
    def load_quiz(file_path=None, delimiter='*'):
        """Load active questions from DATABASE. If CURRENT_TEACHER_ID is set, restrict to that teacher's questions.
        Falls back to file ONLY if database query fails AND file_path is provided.
        """
        questions = []
        
        print(f"\n[load_quiz] Starting quiz load...")
        print(f"[load_quiz] CURRENT_TEACHER_ID = {CURRENT_TEACHER_ID}")
        print(f"[load_quiz] File path = {file_path}")
        
        try:
            # Connect to the Cloudways MySQL database
            print("[load_quiz] Connecting to MySQL database...")
            connection = mysql.connector.connect(
                host=getattr(config, 'MYSQL_HOST', 'localhost'),
                port=getattr(config, 'MYSQL_PORT', 3306),
                database=getattr(config, 'MYSQL_DB', ''),
                user=getattr(config, 'MYSQL_USER', ''),
                password=getattr(config, 'MYSQL_PASS', '')
            )
            cursor = connection.cursor()
            
            # Query to fetch only active questions - NOW INCLUDING choices, image, and id
            if CURRENT_TEACHER_ID:
                query = "SELECT id, description, answer, choices, image FROM questions WHERE status = 1 AND teacher_id = %s"
                print(f"[load_quiz] Executing query with teacher_id = {CURRENT_TEACHER_ID}")
                cursor.execute(query, (CURRENT_TEACHER_ID,))
            else:
                query = "SELECT id, description, answer, choices, image FROM questions WHERE status = 1"
                print("[load_quiz] WARNING: No teacher_id set - loading ALL active questions")
                cursor.execute(query)
            
            rows = cursor.fetchall()
            print(f"[load_quiz] Found {len(rows)} questions in database")
            
            for row in rows:
                question_id = row[0]    # Question ID
                question = row[1]       # Description is the question
                answer = row[2]         # Answer
                choices = row[3] if len(row) > 3 else ''  # Choices (may be NULL)
                image = row[4] if len(row) > 4 else ''    # Image (may be NULL)
                questions.append({
                    'id': question_id,
                    'question': question, 
                    'answer': answer,
                    'choices': choices or '',
                    'image': image or ''
                })
                print(f"[load_quiz]   - Q{question_id}: {question[:50]}... A: {answer}")

            cursor.close()
            connection.close()
            
            if len(questions) == 0:
                print("[load_quiz] WARNING: No questions found in database!")
            else:
                print(f"[load_quiz] ✓ Successfully loaded {len(questions)} questions from database")
                return questions  # CRITICAL: Return immediately when database succeeds - don't fall through to file fallback!
                
        except mysql.connector.Error as error:
            print(f"[load_quiz] ERROR connecting to MySQL database: {error}")
            
            # FALLBACK: Try loading from file ONLY if file_path was provided
            if file_path:
                print(f"[load_quiz] Will attempt to load from file as fallback: {file_path}")
                
                try:
                    if os.path.exists(file_path):
                        with open(file_path, 'r', encoding='utf-8') as file:
                            content = file.read()
                            qa_pairs = content.split(delimiter)
                            for pair in qa_pairs:
                                parts = pair.strip().split('\n', 1)
                                if len(parts) == 2:
                                    questions.append({'question': parts[0], 'answer': parts[1]})
                        print(f"[load_quiz] Loaded {len(questions)} questions from file: {file_path}")
                    else:
                        print(f"[load_quiz] ERROR: File not found: {file_path}")
                except Exception as file_error:
                    print(f"[load_quiz] ERROR loading from file: {file_error}")
            else:
                print("[load_quiz] No fallback file provided - database is the only source")

        return questions
    
    def normalize_answer(answer: str) -> str:
        """
        Normalize answer to handle Filipino and English number words.
        Converts Filipino number words to digits for comparison.
        
        Examples:
            "walo" -> "8"
            "isa dalawa tatlo" -> "123"
            "eight" -> "8"
            "dalawampung tatlo" -> "23"
        """
        if not answer:
            return ""
        
        # Filipino number word mappings
        filipino_numbers = {
            'sero': '0', 'zero': '0',
            'isa': '1', 'one': '1',
            'dalawa': '2', 'two': '2',
            'tatlo': '3', 'three': '3',
            'apat': '4', 'four': '4',
            'lima': '5', 'five': '5',
            'anim': '6', 'six': '6',
            'pito': '7', 'seven': '7',
            'walo': '8', 'eight': '8',
            'siyam': '9', 'nine': '9',
            'sampu': '10', 'ten': '10',
            'labing-isa': '11', 'eleven': '11',
            'labindalawa': '12', 'twelve': '12',
            'labintatlo': '13', 'thirteen': '13',
            'labing-apat': '14', 'fourteen': '14',
            'labinlima': '15', 'fifteen': '15',
            'labing-anim': '16', 'sixteen': '16',
            'labimpito': '17', 'seventeen': '17',
            'labing-walo': '18', 'eighteen': '18',
            'labinsiyam': '19', 'nineteen': '19',
            'dalawampu': '20', 'twenty': '20',
            'tatlumpu': '30', 'thirty': '30',
            'apatnapu': '40', 'forty': '40',
            'limampu': '50', 'fifty': '50',
            'animnapu': '60', 'sixty': '60',
            'pitumpu': '70', 'seventy': '70',
            'walumpu': '80', 'eighty': '80',
            'siyamnapu': '90', 'ninety': '90',
            'daan': '100', 'hundred': '100',
            'libo': '1000', 'thousand': '1000'
        }
        
        # Normalize the answer
        normalized = answer.lower().strip()
        
        # Try direct conversion if it's already a number
        if normalized.replace('.', '').replace('-', '').isdigit():
            return normalized
        
        # Convert Filipino/English words to numbers
        words = normalized.split()
        result = []
        
        for word in words:
            # Remove punctuation
            clean_word = word.strip('.,!?')
            if clean_word in filipino_numbers:
                result.append(filipino_numbers[clean_word])
            elif clean_word.isdigit():
                result.append(clean_word)
            else:
                # Keep non-number words as-is (for text answers)
                result.append(clean_word)
        
        # Join results
        # If all results are single digits, concatenate them (e.g., "isa dalawa" -> "12")
        if len(result) > 1 and all(r.isdigit() and len(r) == 1 for r in result):
            return ''.join(result)
        
        # Otherwise join with spaces
        return ' '.join(result)
    
    def answers_match(student_answer: str, correct_answer: str) -> bool:
        """
        Compare student answer with correct answer.
        Handles Filipino numbers, English numbers, and text answers.
        
        Returns True if answers match (case-insensitive, normalized).
        """
        # Normalize both answers
        normalized_student = normalize_answer(student_answer)
        normalized_correct = normalize_answer(correct_answer)
        
        print(f"[Answer Check] Student: '{student_answer}' -> '{normalized_student}'")
        print(f"[Answer Check] Correct: '{correct_answer}' -> '{normalized_correct}'")
        
        # Make matching more forgiving for numeric answers and repeated tokens
        ns = normalized_student.lower().strip()
        nc = normalized_correct.lower().strip()

        # If either side is empty, do strict equality check
        if not ns or not nc:
            return ns == nc

        # If the correct answer is primarily numeric, compare digits-only forms
        digits_nc = re.sub(r"\D", "", nc)
        digits_ns = re.sub(r"\D", "", ns)
        if digits_nc:
            # Accept if the student's digits-string contains the expected digits
            if digits_nc == digits_ns:
                return True
            if digits_nc and digits_nc in digits_ns:
                return True

            # Also accept if student's spoken tokens repeated the same numeric token
            tokens = [t for t in re.split(r"\s+", ns) if t]
            if tokens:
                # collapse repeated identical tokens (e.g., "15 15 15" -> "15")
                uniq = list(dict.fromkeys(tokens))
                if len(uniq) == 1 and re.sub(r"\D", "", uniq[0]) == digits_nc:
                    return True

        # For textual answers, accept exact match or containment (loose match)
        if nc == ns:
            return True
        if nc in ns or ns in nc:
            return True

        # Fallback to strict equality
        return ns == nc
    
    def run_quiz(questions):
        score = 0
        # If a loading animation is running, stop it now since we're about to start asking questions.
        try:
            global _quiz_loading_stop_event
            if '_quiz_loading_stop_event' in globals() and isinstance(_quiz_loading_stop_event, threading.Event):
                _quiz_loading_stop_event.set()
                # Ensure we go to idle before first question
                try:
                    set_oled_expression('idle')
                except Exception:
                    pass
        except Exception:
            pass
        for i, q in enumerate(questions, start=1):
            # ==== ASK QUESTION WITH REPEAT FUNCTIONALITY ====
            question_asked = False
            metadata_logged = False  # Track if we've logged metadata for this question
            # Store parsed choices at question level for letter-to-value mapping
            question_choices_map = {}  # Maps 'A' -> 'A. 10', 'B' -> 'B. 20', etc.
            
            while not question_asked:
                question_text = f"{ordinal(i)} question: {q['question']}"
                print(question_text)
                
                # LOG STRUCTURED DATA: Include all fields needed for database insertion
                # Format: QUESTION_ID|CHOICES|IMAGE at the start of each question
                # CRITICAL: Only log metadata ONCE per question, even if student repeats
                if not metadata_logged:
                    metadata_line = f"[METADATA] question_id={q.get('id', '')}, choices={q.get('choices', '')}, image={q.get('image', '')}"
                    log_conversation(metadata_line)
                    log_conversation(question_text)
                    metadata_logged = True  # Mark as logged to prevent duplicate logging on repeat
                
                # Parse choices and build full speech text
                choices_text = ""
                choices_list_for_speech = []  # Store for individual speech
                try:
                    raw_choices = q.get('choices', '')
                    if raw_choices and raw_choices.strip():
                        import json
                        # Try to parse as JSON array (e.g., '["A. 10", "B. 20", "C. 30"]')
                        try:
                            choices_list = json.loads(raw_choices)
                        except:
                            # Fallback: split by comma if not valid JSON
                            choices_list = [c.strip() for c in raw_choices.split(',') if c.strip()]
                        
                        # Only say choices if there are 2 or more options
                        if isinstance(choices_list, list) and len(choices_list) > 1:
                            # Format choices for speech and build letter mapping
                            formatted_choices = []
                            for idx, choice in enumerate(choices_list):
                                choice_str = str(choice).strip()
                                if choice_str:
                                    formatted_choices.append(choice_str)
                                    # Build letter mapping: A->choice, B->choice, etc.
                                    # Extract letter from choice like "A. 10" or just use index
                                    letter_match = re.match(r'^([A-D])[.\s)]', choice_str, re.IGNORECASE)
                                    if letter_match:
                                        letter = letter_match.group(1).upper()
                                        question_choices_map[letter] = choice_str
                                    else:
                                        # If no letter prefix, map A,B,C,D based on index
                                        letter = chr(65 + idx)  # 65 is ASCII for 'A'
                                        question_choices_map[letter] = choice_str
                            
                            if len(formatted_choices) > 1:
                                choices_list_for_speech = formatted_choices
                                print(f"[Q{i}] Choices: {formatted_choices}")
                                print(f"[Q{i}] Letter mapping: {question_choices_map}")
                except Exception as e:
                    print(f"[Q{i}] Could not parse choices: {e}")
                    choices_list_for_speech = []
                
                # Speak question first (without choices)
                try:
                    speak_and_play_text(TranslateToFil(question_text))
                except Exception:
                    synthesize_speech(TranslateToFil(question_text))
                    play_audio('response.wav')
                
                # If there are choices, speak them one by one with pauses
                if choices_list_for_speech:
                    # Brief pause after the question
                    time.sleep(0.8)
                    
                    # Say "Your choices are:"
                    try:
                        speak_and_play_text("Your choices are:")
                    except Exception:
                        synthesize_speech("Your choices are:")
                        play_audio('response.wav')
                    
                    time.sleep(0.5)
                    
                    # Say each choice with a pause between them
                    # Format each choice to emphasize the letter (e.g., "Letter A: 10")
                    for choice_idx, choice in enumerate(choices_list_for_speech):
                        # Check if choice already has letter prefix (A., B., etc.)
                        choice_str = str(choice).strip()
                        letter_match = re.match(r'^([A-D])[.\s)]+(.+)', choice_str, re.IGNORECASE)
                        
                        if letter_match:
                            # Choice has letter prefix - emphasize it
                            letter = letter_match.group(1).upper()
                            value = letter_match.group(2).strip()
                            speech_choice = f"Letter {letter}: {value}"
                        else:
                            # No letter prefix - add one based on index
                            letter = chr(65 + choice_idx)  # A, B, C, D
                            speech_choice = f"Letter {letter}: {choice_str}"
                        
                        try:
                            speak_and_play_text(speech_choice)
                        except Exception:
                            synthesize_speech(speech_choice)
                            play_audio('response.wav')
                        
                        # Pause between choices (except after the last one)
                        if choice_idx < len(choices_list_for_speech) - 1:
                            time.sleep(0.7)  # 0.7 second pause between choices
                
                print(f"[Q{i}] Listening for answer or 'repeat'...")
                
                # Show curious expression - eyes focused and attentive
                set_oled_expression('curious')
                
                # CLEAR PREVIOUS RECORDING before waiting for new one
                try:
                    if os.path.exists(audio_raw_path):
                        os.remove(audio_raw_path)

                    # Prefer clearing via local Flask proxy (fast, local) so we don't
                    # accidentally hit a remote tunnel which may be slower or unreachable.
                    cleared = False
                    try:
                        clr = _http_session.get('http://127.0.0.1:5000/clear', timeout=1.0)
                        if getattr(clr, 'status_code', None) == 200:
                            cleared = True
                            print(f"[Quiz] Cleared recorder via local proxy (status={clr.status_code})")
                    except Exception:
                        cleared = False

                    # Fallback to configured tunnel/remote clear only if local proxy failed
                    if not cleared:
                        try:
                            _http_session.get('https://nonbasic-bob-inimical.ngrok-free.dev/clear', timeout=2)
                            print("[Quiz] Cleared recorder via tunnel fallback")
                        except Exception:
                            pass

                    # Ensure the Flask notification queue is reset after clearing so
                    # we don't pick up a stale notification (short pause helps reliability)
                    try:
                        reset_r = _http_session.get('http://127.0.0.1:5000/reset_recording_notification', timeout=1.0)
                        print(f"[Quiz] Reset recording notification queue (status={getattr(reset_r, 'status_code', None)})")
                    except Exception as _e:
                        print(f"[Quiz] Could not reset recording queue: {_e}")

                    # Small grace period to allow ESP and Flask proxy to settle
                    time.sleep(0.20)
                except Exception:
                    pass
                
                # Wait for response (answer or repeat request)
                try:
                    # Accept both Filipino and English answers - don't force English STT
                    # The normalize_answer function handles Filipino number words
                    is_numeric_expected = False
                    try:
                        corr = str(q.get('answer', '')).strip()
                        if re.search(r'\d', corr):
                            is_numeric_expected = True
                        else:
                            # also check normalized word->digit conversion
                            if normalize_answer(corr).strip().isdigit():
                                is_numeric_expected = True
                    except Exception:
                        is_numeric_expected = False

                    # Ensure Flask recording notification queue is clear before waiting
                    try:
                        reset_r = _http_session.get('http://127.0.0.1:5000/reset_recording_notification', timeout=1.0)
                        print(f"[Quiz] Reset recording notification queue (status={getattr(reset_r, 'status_code', None)})")
                    except Exception as _e:
                        print(f"[Quiz] Could not reset recording queue: {_e}")

                    print(f"[Quiz] Calling record_and_transcribe (numeric_expected={is_numeric_expected})...")
                    complete_text, forced_timeout = record_and_transcribe(
                        esp_host=esp_record_host,
                        poll_for_recording=True,
                        max_poll_seconds=60,
                        use_english=False
                    )
                except Exception:
                    complete_text = ""
                    forced_timeout = False
                
                # Return to idle after recording
                set_oled_expression('idle')
                
                # Check if student wants to repeat the question
                if complete_text and any(word in complete_text.lower() for word in ['repeat', 'again']):
                    print(f"[Q{i}] Student requested repeat")
                    repeat_msg = "Okay, I'll repeat the question."
                    try:
                        speak_and_play_text(repeat_msg)
                    except Exception:
                        synthesize_speech(repeat_msg)
                        play_audio('response.wav')
                    time.sleep(0.5)
                    continue  # Re-ask the question
                
                # If we got here, student provided an answer (not repeat request)
                question_asked = True
            
            # ==== VALIDATE ANSWER WITH STUDENT ====
            answer_confirmed = False
            user_answer = complete_text
            
            # Handle empty/timeout responses: re-ask the same question a couple times
            if not user_answer or not user_answer.strip():
                # Allow the student a small number of re-tries when no answer was heard
                MAX_NO_ANSWER_RETRIES = int(getattr(config, 'QUIZ_NO_ANSWER_RETRIES', 2))
                no_answer_attempts = 0
                reasked = False
                while (not user_answer or not user_answer.strip()) and no_answer_attempts < MAX_NO_ANSWER_RETRIES:
                    no_answer_attempts += 1
                    reasked = True
                    # If forced timeout (proxy timed out) use a specific prompt
                    try:
                        if forced_timeout:
                            prompt_text = "Timeout occurred. Please repeat your answer."
                        else:
                            prompt_text = "I didn't hear anything. Please repeat your answer."
                        print(f"[Q{i}] No answer received — re-asking ({no_answer_attempts}/{MAX_NO_ANSWER_RETRIES}): {prompt_text}")
                        # Prompt student to repeat
                        try:
                            speak_and_play_text(prompt_text)
                        except Exception:
                            synthesize_speech(prompt_text)
                            play_audio('response.wav')
                    except Exception:
                        pass

                    # Clear and reset recorder then listen again
                    try:
                        if os.path.exists(audio_raw_path):
                            os.remove(audio_raw_path)
                        try:
                            _http_session.get('http://127.0.0.1:5000/clear', timeout=1.0)
                        except Exception:
                            try:
                                _http_session.get('https://nonbasic-bob-inimical.ngrok-free.dev/clear', timeout=2)
                            except Exception:
                                pass
                        try:
                            _http_session.get('http://127.0.0.1:5000/reset_recording_notification', timeout=1.0)
                        except Exception:
                            pass
                        time.sleep(0.25)
                    except Exception:
                        pass

                    try:
                        print(f"[Q{i}] Re-listening for answer (attempt {no_answer_attempts})...")
                        user_answer, forced_timeout = record_and_transcribe(
                            esp_host=esp_record_host,
                            poll_for_recording=True,
                            max_poll_seconds=40,
                            use_english=False
                        )
                    except Exception:
                        user_answer = ""

                    # If student explicitly asked to repeat, re-ask the question
                    if user_answer and any(word in user_answer.lower() for word in ['repeat', 'again']):
                        try:
                            speak_and_play_text("Okay, I'll repeat the question.")
                        except Exception:
                            synthesize_speech("Okay, I'll repeat the question.")
                            play_audio('response.wav')
                        # Reset so outer loop will re-ask the question
                        question_asked = False
                        break

                # After retries, if still empty treat as skipped
                if not user_answer or not user_answer.strip():
                    if forced_timeout:
                        timeout_msg = "Timeout occurred. Moving to the next question."
                        print(f"[Q{i}] Timeout - no answer received after retries")
                        log_conversation(f"Student Answer: [TIMEOUT]")
                        log_conversation(f"Correct Answer: {q['answer']}")
                        try:
                            speak_and_play_text(timeout_msg)
                        except Exception:
                            synthesize_speech(timeout_msg)
                            play_audio('response.wav')
                    else:
                        no_answer_msg = "No answer was heard after several attempts. Moving to the next question."
                        print(f"[Q{i}] No transcription received after retries")
                        log_conversation(f"Student Answer: [NO ANSWER]")
                        log_conversation(f"Correct Answer: {q['answer']}")
                        try:
                            speak_and_play_text(no_answer_msg)
                        except Exception:
                            synthesize_speech(no_answer_msg)
                            play_audio('response.wav')

                    set_oled_expression('idle')
                    time.sleep(1)
                    # Move to next question
                    continue
                else:
                    # We got an answer after re-asking; proceed with confirmation/validation
                    print(f"[Q{i}] Got answer after re-asking: {user_answer}")
            
            # Validation loop - confirm answer with student
            max_confirmation_attempts = 2
            confirmation_attempt = 0
            
            while not answer_confirmed and confirmation_attempt < max_confirmation_attempts:
                # Ask for confirmation
                confirmation_text = f"I heard, {user_answer}. Is that correct?"
                print(f"[Q{i}] Confirming: {user_answer}")
                try:
                    speak_and_play_text(confirmation_text)
                except Exception:
                    synthesize_speech(confirmation_text)
                    play_audio('response.wav')
                
                # Clear recording before waiting for confirmation
                try:
                    if os.path.exists(audio_raw_path):
                        os.remove(audio_raw_path)
                    _http_session.get('https://nonbasic-bob-inimical.ngrok-free.dev/clear', timeout=3)
                except Exception:
                    pass
                
                # Wait for confirmation (yes/no)
                try:
                    # Accept Filipino confirmations (oo/hindi) as well as English (yes/no)
                    try:
                        corr = str(q.get('answer', '')).strip()
                        conf_expected_numeric = True if re.search(r'\d', corr) or normalize_answer(corr).strip().isdigit() else False
                    except Exception:
                        conf_expected_numeric = False

                    # Reset queue before confirmation listen to avoid stale notifications
                    try:
                        reset_r2 = _http_session.get('http://127.0.0.1:5000/reset_recording_notification', timeout=1.0)
                        print(f"[Quiz] Reset recording notification queue before confirmation (status={getattr(reset_r2, 'status_code', None)})")
                    except Exception as _e:
                        print(f"[Quiz] Could not reset recording queue before confirmation: {_e}")

                    print(f"[Quiz] Calling record_and_transcribe for confirmation (bilingual support)...")
                    confirm_text, _ = record_and_transcribe(
                        esp_host=esp_record_host,
                        poll_for_recording=True,
                        max_poll_seconds=18,  # Shorter timeout for yes/no but slightly increased
                        use_english=False
                    )
                except Exception:
                    confirm_text = ""
                
                confirm_lower = confirm_text.lower() if confirm_text else ""
                print(f"[Q{i}] Confirmation response: {confirm_text}")
                
                # Check if student confirms (English variants)
                if any(word in confirm_lower for word in ['yes', 'y', 'yeah', 'correct', 'right']):
                    answer_confirmed = True
                    print(f"[Q{i}] Answer confirmed by student")
                # Check if student wants to re-answer (English variants)
                elif any(word in confirm_lower for word in ['no', 'not', 'wrong', 'nope', 'repeat']):
                    retry_msg = "Okay, please say your answer again."
                    print(f"[Q{i}] Student wants to re-answer")
                    try:
                        speak_and_play_text(retry_msg)
                    except Exception:
                        synthesize_speech(retry_msg)
                        play_audio('response.wav')
                    
                    # Clear and wait for new answer
                    try:
                        if os.path.exists(audio_raw_path):
                            os.remove(audio_raw_path)
                        _http_session.get('https://nonbasic-bob-inimical.ngrok-free.dev/clear', timeout=3)
                    except Exception:
                        pass
                    
                    try:
                        # If student is re-answering, preserve numeric bias if expected
                        try:
                            corr = str(q.get('answer', '')).strip()
                            re_expected_numeric = True if re.search(r'\d', corr) or normalize_answer(corr).strip().isdigit() else False
                        except Exception:
                            re_expected_numeric = False

                        # Reset queue before re-answer to ensure we wait for a NEW recording
                        try:
                            reset_r3 = _http_session.get('http://127.0.0.1:5000/reset_recording_notification', timeout=1.0)
                            print(f"[Quiz] Reset recording notification queue before re-answer (status={getattr(reset_r3, 'status_code', None)})")
                        except Exception as _e:
                            print(f"[Quiz] Could not reset recording queue before re-answer: {_e}")

                        print(f"[Quiz] Calling record_and_transcribe for re-answer (expected_numeric={re_expected_numeric})...")
                        user_answer, _ = record_and_transcribe(
                            esp_host=esp_record_host,
                            poll_for_recording=True,
                            max_poll_seconds=60,
                            use_english=re_expected_numeric
                        )
                        confirmation_attempt += 1
                    except Exception:
                        user_answer = ""
                        break
                else:
                    # Unclear response, assume confirmed after 2 attempts
                    confirmation_attempt += 1
                    if confirmation_attempt >= max_confirmation_attempts:
                            # Do not accept ambiguous confirmation blindly. If the student's
                            # candidate answer already matches the expected answer (fuzzy),
                            # accept it. Otherwise treat as unconfirmed and skip.
                            try:
                                if answers_match(user_answer, q.get('answer', '')):
                                    answer_confirmed = True
                                    print(f"[Q{i}] Confirmation unclear but student answer matches expected; accepting '{user_answer}'")
                                else:
                                    print(f"[Q{i}] Confirmation remained unclear after {max_confirmation_attempts} attempts; NOT accepting answer")
                                    answer_confirmed = False
                                    user_answer = ""  # clear candidate answer so downstream treats as no answer
                                break
                            except Exception:
                                # On any error during matching, fall back to not accepting
                                print(f"[Q{i}] Error while evaluating ambiguous confirmation; skipping question")
                                answer_confirmed = False
                                user_answer = ""
                                break
            
            # If answer was not confirmed, treat as no answer and move to next question
            if not answer_confirmed:
                print(f"[Q{i}] Answer was not confirmed; skipping question {i} without scoring")
                log_conversation(f"Student Answer: [UNCONFIRMED]")
                # Provide a short message and continue to next question
                try:
                    speak_and_play_text("Answer not confirmed. Moving to the next question.")
                except Exception:
                    try:
                        synthesize_speech("Answer not confirmed. Moving to the next question.")
                        play_audio('response.wav')
                    except Exception:
                        pass
                set_oled_expression('idle')
                time.sleep(0.5)
                continue

            # ==== PROCESS FINAL ANSWER (OPTIMIZED FOR SPEED) ====
            print(f"[Q{i}] Raw answer: {user_answer}")
            
            # Show "thinking" expression while processing the answer
            set_oled_expression('thinking')
            
            # Start wait timer for answer processing (in case profanity check is slow)
            answer_wait_timer = start_please_wait_timer(
                delay_seconds=5.0,
                message="Please wait, I'm checking your answer..."
            )
            
            try:
                # OPTIMIZATION: Start TTS generation in parallel while checking answer
                is_correct = False
                is_profane = False
                response_text = ""
                
                # Check if student answered with a letter (A, B, C, D)
                # If so, map it to the actual choice value
                mapped_answer = user_answer
                recorded_answer = user_answer  # What we'll log to database
                user_answer_normalized = user_answer.strip().upper()
                
                # Strengthen letter detection - handle various patterns:
                # "A", "a", "letter A", "A.", "answer is A", "I choose A", etc.
                detected_letter = None
                
                # Pattern 1: Single letter
                if len(user_answer_normalized) == 1 and user_answer_normalized in ['A', 'B', 'C', 'D']:
                    detected_letter = user_answer_normalized
                
                # Pattern 2: "letter X" or "X."
                if not detected_letter:
                    letter_pattern = re.search(r'\b(?:letter\s+)?([A-D])(?:[.\s]|$)', user_answer_normalized)
                    if letter_pattern:
                        detected_letter = letter_pattern.group(1).upper()
                
                # Pattern 3: "answer is X" or "I choose X" or "my answer is X"
                if not detected_letter:
                    answer_pattern = re.search(r'(?:answer|choose|pick|select|is)\s+([A-D])\b', user_answer_normalized, re.IGNORECASE)
                    if answer_pattern:
                        detected_letter = answer_pattern.group(1).upper()
                
                # Pattern 4: Spoken phonetically - "ay", "bee", "see", "dee"
                if not detected_letter:
                    phonetic_map = {'AY': 'A', 'EY': 'A', 'BEE': 'B', 'BI': 'B', 'SEE': 'C', 'SI': 'C', 'DEE': 'D', 'DI': 'D'}
                    for phonetic, letter in phonetic_map.items():
                        if phonetic in user_answer_normalized.replace(' ', ''):
                            detected_letter = letter
                            break
                
                # Apply letter mapping if detected
                if detected_letter and detected_letter in question_choices_map:
                    mapped_answer = question_choices_map[detected_letter]
                    # Extract just the value for comparison (remove letter prefix like "A. ")
                    value_match = re.match(r'^[A-D][.\s)]+(.+)', mapped_answer, re.IGNORECASE)
                    if value_match:
                        recorded_answer = mapped_answer  # Record full choice like "A. 10"
                    else:
                        recorded_answer = mapped_answer
                    print(f"[Q{i}] Letter answer '{detected_letter}' mapped to: {mapped_answer}")
                elif detected_letter:
                    print(f"[Q{i}] Letter '{detected_letter}' detected but not found in choices map")
                
                # Log the final recorded answer (mapped if letter, original otherwise)
                log_conversation(f"Student Answer: {recorded_answer}")
                print(f"[Q{i}] Final answer (recorded): {recorded_answer}")
                
                # Quick profanity check
                if is_profane_by_ai(user_answer):
                    is_profane = True
                    response_text = "Please refrain from using bad words."
                    set_oled_expression('annoyed')  # Annoyed squint for inappropriate language
                elif answers_match(mapped_answer, q['answer']):
                    is_correct = True
                    response_text = "Correct!"
                    set_oled_expression('excited')  # 🎉 Excited shimmy for correct answer in quiz!
                    log_conversation(response_text)
                    score += 1
                else:
                    response_text = f"Oops! The correct answer is {q['answer']}."
                    set_oled_expression('sad')  #   Sad looking-down eyes for incorrect answer
                    log_conversation(response_text)
            finally:
                # Cancel wait timer once processing is done
                answer_wait_timer.set()
            
            # Speak response immediately (no waiting)
            print(f"GENTA: {response_text}")
            synthesize_speech(response_text)
            play_audio('response.wav')
            
            # Brief pause to let expression show - longer for correct to celebrate!
            time.sleep(1.5 if is_correct else 1.0)
            
            # Return to idle expression
            set_oled_expression('idle')
            
            # Shorter breather before next question (reduced from 1s to 0.5s)
            time.sleep(0.5)

        # ============ QUIZ COMPLETION - SHOW EXPRESSION BASED ON PERFORMANCE ============
        result_text = f"You got {score} correct answers out of {len(questions)} questions."
        result_log_text = f"Student got {score} out of {len(questions)} questions correct."
        print(result_text)
        log_conversation(result_log_text)
        
        # Calculate percentage
        percentage = (score / len(questions)) * 100 if len(questions) > 0 else 0
        
        # Concise numeric result + minimal praise. Use chunked playback for snappier start.
        try:
            result_brief = f"You finished the quiz. You got a score of {score} out of {len(questions)} questions."

            # Minimal praise policy: single short word/phrase depending on performance
            # Personalize praise with the student's first name when available.
            first_name = None
            try:
                if CURRENT_STUDENT_NAME and isinstance(CURRENT_STUDENT_NAME, str) and CURRENT_STUDENT_NAME.strip():
                    first_name = CURRENT_STUDENT_NAME.strip().split()[0]
            except Exception:
                first_name = None

            if percentage >= 100:
                oled_expr = 'happy'
                praise = f"Wow, Congratulations {first_name}! Keep it up!" if first_name else "Amazing!"
                
            elif percentage >= 80:
                oled_expr = 'happy'
                praise = f"Great job! Keep it up, {first_name}!" if first_name else "Good job!"  
            elif percentage >= 60:
                oled_expr = 'happy'
                praise = f"Good job, {first_name}!" if first_name else "Good job!"
            else:
                oled_expr = 'sad'
                praise = f"That's okay, keep practicing {first_name}!" if first_name else "Good effort!"

            # Speak the numeric result first (chunked, so playback starts quickly)
            played = False
            try:
                played = speak_and_play_text(result_brief)
            except Exception:
                played = False

            # Short pause then speak a very short praise (single phrase)
            time.sleep(0.4)
            try:
                if speak_and_play_text(praise):
                    played = True
            except Exception:
                played = played or False

            # Fallback: if chunked playback failed, synthesize and play single short file
            if not played:
                try:
                    synthesize_speech(result_brief)
                    play_audio('response.wav')
                    time.sleep(0.3)
                    synthesize_speech(praise)
                    play_audio('response.wav')
                except Exception:
                    pass

            # Set a friendly but restrained OLED expression
            try:
                set_oled_expression(oled_expr)
            except Exception:
                pass

        except Exception:
            # If something goes wrong, ensure we at least announce the numeric result
            try:
                synthesize_speech(result_text)
                play_audio('response.wav')
            except Exception:
                pass

        # Return to idle expression after speaking
        try:
            set_oled_expression('idle')
        except Exception:
            pass
        


    def quiz_analysis(conversation_log, output_docx_path):
        # Helper: create a styled DOCX using python-docx
        def create_templated_docx(title: str, subtitle: str, body_text: str, out_path: str):
            try:
                doc = docx.Document()
                
                # ============ DOCUMENT-WIDE STYLES ============
                style = doc.styles['Normal']
                try:
                    style.font.name = 'Calibri'
                    style.font.size = docx.shared.Pt(11)
                except Exception:
                    pass

                # ============ PROFESSIONAL COVER PAGE ============
                # Logo/Title Section
                cover_title = doc.add_paragraph()
                cover_title.alignment = 1  # Center
                run = cover_title.add_run("GENTA")
                run.bold = True
                run.font.size = docx.shared.Pt(48)
                run.font.color.rgb = docx.shared.RGBColor(0, 102, 204)  # Blue
                
                # Subtitle on cover
                cover_subtitle = doc.add_paragraph()
                cover_subtitle.alignment = 1
                run2 = cover_subtitle.add_run("Student Analysis & Learning Module")
                run2.font.size = docx.shared.Pt(20)
                run2.font.color.rgb = docx.shared.RGBColor(68, 114, 196)
                cover_subtitle.paragraph_format.space_after = docx.shared.Pt(24)
                
                # Decorative line
                doc.add_paragraph('_' * 60).alignment = 1
                doc.add_paragraph()  # Spacing
                
                # Student Info Section (if available from globals)
                info_para = doc.add_paragraph()
                info_para.alignment = 1
                try:
                    student_name = CURRENT_STUDENT_NAME if CURRENT_STUDENT_NAME else "Student"
                    student_id = CURRENT_STUDENT_ID if CURRENT_STUDENT_ID else "N/A"
                    # Prefer the already-populated CURRENT_TEACHER_NAME; otherwise
                    # try to look up firstname/lastname by CURRENT_TEACHER_ID so the
                    # document shows the teacher's real name instead of the word
                    # "Teacher" or a placeholder.
                    teacher_name = None
                    try:
                        if CURRENT_TEACHER_NAME and isinstance(CURRENT_TEACHER_NAME, str) and CURRENT_TEACHER_NAME.strip() and CURRENT_TEACHER_NAME.strip().lower() != 'teacher':
                            teacher_name = CURRENT_TEACHER_NAME.strip()
                    except Exception:
                        teacher_name = None

                    if not teacher_name:
                        # Try multiple strategies to resolve teacher name from DB.
                        try:
                            if CURRENT_TEACHER_ID:
                                # Primary: teachers table (common columns)
                                try:
                                    conn_tmp = mysql.connector.connect(
                                        host=getattr(config, 'MYSQL_HOST', 'localhost'),
                                        port=getattr(config, 'MYSQL_PORT', 3306),
                                        database=getattr(config, 'MYSQL_DB', ''),
                                        user=getattr(config, 'MYSQL_USER', ''),
                                        password=getattr(config, 'MYSQL_PASS', '')
                                    )
                                    cur_tmp = conn_tmp.cursor()
                                    cur_tmp.execute("SELECT firstname, lastname FROM teachers WHERE id = %s LIMIT 1", (CURRENT_TEACHER_ID,))
                                    tro = cur_tmp.fetchone()
                                except Exception:
                                    tro = None
                                finally:
                                    try:
                                        cur_tmp.close()
                                    except Exception:
                                        pass
                                    try:
                                        conn_tmp.close()
                                    except Exception:
                                        pass

                                if tro and len(tro) >= 1 and (tro[0] or (len(tro) > 1 and tro[1])):
                                    fn = tro[0] or ''
                                    ln = tro[1] if len(tro) > 1 and tro[1] else ''
                                    teacher_name = (fn + ' ' + ln).strip() if (fn or ln) else None
                                else:
                                    # Try alternative single-column name fields in teachers
                                    try:
                                        conn_a = mysql.connector.connect(
                                            host=getattr(config, 'MYSQL_HOST', 'localhost'),
                                            port=getattr(config, 'MYSQL_PORT', 3306),
                                            database=getattr(config, 'MYSQL_DB', ''),
                                            user=getattr(config, 'MYSQL_USER', ''),
                                            password=getattr(config, 'MYSQL_PASS', '')
                                        )
                                        cur_a = conn_a.cursor()
                                        cur_a.execute("SELECT name FROM teachers WHERE id = %s LIMIT 1", (CURRENT_TEACHER_ID,))
                                        r_a = cur_a.fetchone()
                                        if r_a and r_a[0]:
                                            teacher_name = str(r_a[0]).strip()
                                    except Exception:
                                        teacher_name = teacher_name
                                    finally:
                                        try:
                                            cur_a.close()
                                        except Exception:
                                            pass
                                        try:
                                            conn_a.close()
                                        except Exception:
                                            pass

                                    # If still not found, try users table with common column names
                                    if not teacher_name:
                                        try:
                                            conn2 = mysql.connector.connect(
                                                host=getattr(config, 'MYSQL_HOST', 'localhost'),
                                                port=getattr(config, 'MYSQL_PORT', 3306),
                                                database=getattr(config, 'MYSQL_DB', ''),
                                                user=getattr(config, 'MYSQL_USER', ''),
                                                password=getattr(config, 'MYSQL_PASS', '')
                                            )
                                            cur2 = conn2.cursor()
                                            # Try common naming variants
                                            tried_user_queries = [
                                                ("SELECT first_name, last_name FROM users WHERE id = %s LIMIT 1",),
                                                ("SELECT firstname, lastname FROM users WHERE id = %s LIMIT 1",),
                                                ("SELECT name FROM users WHERE id = %s LIMIT 1",),
                                                ("SELECT full_name FROM users WHERE id = %s LIMIT 1",),
                                            ]
                                            # Try first two-field variants first
                                            try:
                                                cur2.execute("SELECT first_name, last_name FROM users WHERE id = %s LIMIT 1", (CURRENT_TEACHER_ID,))
                                                uro = cur2.fetchone()
                                                if uro and len(uro) >= 1 and (uro[0] or (len(uro) > 1 and uro[1])):
                                                    fn = uro[0] or ''
                                                    ln = uro[1] if len(uro) > 1 and uro[1] else ''
                                                    teacher_name = (fn + ' ' + ln).strip() if (fn or ln) else None
                                            except Exception:
                                                teacher_name = teacher_name

                                            if not teacher_name:
                                                try:
                                                    cur2.execute("SELECT name FROM users WHERE id = %s LIMIT 1", (CURRENT_TEACHER_ID,))
                                                    uro2 = cur2.fetchone()
                                                    if uro2 and uro2[0]:
                                                        teacher_name = str(uro2[0]).strip()
                                                except Exception:
                                                    teacher_name = teacher_name

                                        except Exception:
                                            teacher_name = None
                                        finally:
                                            try:
                                                cur2.close()
                                            except Exception:
                                                pass
                                            try:
                                                conn2.close()
                                            except Exception:
                                                pass

                                # As a last-resort, try to resolve via the students table using CURRENT_STUDENT_ID (if available)
                                if not teacher_name and CURRENT_STUDENT_ID:
                                    try:
                                        conn_s = mysql.connector.connect(
                                            host=getattr(config, 'MYSQL_HOST', 'localhost'),
                                            port=getattr(config, 'MYSQL_PORT', 3306),
                                            database=getattr(config, 'MYSQL_DB', ''),
                                            user=getattr(config, 'MYSQL_USER', ''),
                                            password=getattr(config, 'MYSQL_PASS', '')
                                        )
                                        cur_s = conn_s.cursor()
                                        # Attempt a join to locate teacher info from students
                                        cur_s.execute(
                                            "SELECT t.firstname, t.lastname, t.name FROM teachers t JOIN students s ON s.teacher_id = t.id WHERE s.lrn = %s OR s.id = %s LIMIT 1",
                                            (CURRENT_STUDENT_ID, CURRENT_STUDENT_ID)
                                        )
                                        srow = cur_s.fetchone()
                                        if srow:
                                            # prefer firstname/lastname
                                            if srow[0] or (len(srow) > 1 and srow[1]):
                                                fn = srow[0] or ''
                                                ln = srow[1] if len(srow) > 1 and srow[1] else ''
                                                teacher_name = (fn + ' ' + ln).strip() if (fn or ln) else None
                                            elif len(srow) > 2 and srow[2]:
                                                teacher_name = str(srow[2]).strip()
                                    except Exception:
                                        teacher_name = teacher_name
                                    finally:
                                        try:
                                            cur_s.close()
                                        except Exception:
                                            pass
                                        try:
                                            conn_s.close()
                                        except Exception:
                                            pass
                        except Exception:
                            teacher_name = None

                    if not teacher_name:
                        teacher_name = "Teacher"
                    
                    info_run = info_para.add_run(f"Student: {student_name}\nLRN: {student_id}\nTeacher: {teacher_name}")
                    info_run.font.size = docx.shared.Pt(14)
                    info_run.bold = True
                except Exception:
                    pass
                info_para.paragraph_format.space_after = docx.shared.Pt(36)
                
                # Report metadata
                meta_para = doc.add_paragraph()
                meta_para.alignment = 1
                meta_run = meta_para.add_run(f"Generated: {ftime}")
                meta_run.font.size = docx.shared.Pt(12)
                meta_run.italic = True
                meta_para.paragraph_format.space_after = docx.shared.Pt(48)
                
                # Page break after cover
                doc.add_page_break()
                
                # ============ TABLE OF CONTENTS ============
                toc_heading = doc.add_paragraph()
                toc_run = toc_heading.add_run("Table of Contents")
                toc_run.bold = True
                toc_run.font.size = docx.shared.Pt(18)
                toc_run.font.color.rgb = docx.shared.RGBColor(0, 102, 204)
                toc_heading.paragraph_format.space_after = docx.shared.Pt(12)
                
                # TOC entries (based on typical analysis structure)
                toc_entries = [
                    ("1. Executive Summary", "Page 3"),
                    ("2. Student Strengths", "Page 4"),
                    ("3. Areas for Improvement", "Page 5"),
                    ("4. Recommendations & Strategies", "Page 6"),
                    ("5. Detailed Conversation Log", "Page 7"),
                ]
                
                for entry, page in toc_entries:
                    toc_entry = doc.add_paragraph()
                    toc_entry.paragraph_format.left_indent = docx.shared.Inches(0.25)
                    entry_run = toc_entry.add_run(entry)
                    entry_run.font.size = docx.shared.Pt(12)
                    # Add dots
                    dots = toc_entry.add_run(" " + "." * 50)
                    dots.font.color.rgb = docx.shared.RGBColor(192, 192, 192)
                    page_run = toc_entry.add_run(f" {page}")
                    page_run.font.size = docx.shared.Pt(12)
                    toc_entry.paragraph_format.space_after = docx.shared.Pt(3)
                
                # Page break after TOC
                doc.add_page_break()

                # ============ MAIN CONTENT ============
                # Header for all pages after cover
                try:
                    section = doc.sections[0]
                    header = section.header
                    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                    header_para.text = 'GENTA — Student Analysis Report'
                    header_para.alignment = 1
                    # Add a line below header
                    header_line = header.add_paragraph('_' * 80)
                    header_line.alignment = 1
                except Exception:
                    pass

                # Main Title
                h = doc.add_paragraph()
                run = h.add_run(title)
                run.bold = True
                run.font.size = docx.shared.Pt(20)
                run.font.color.rgb = docx.shared.RGBColor(0, 102, 204)
                h.paragraph_format.space_after = docx.shared.Pt(12)

                # Subtitle / metadata
                if subtitle:
                    s = doc.add_paragraph()
                    sub_run = s.add_run(subtitle)
                    sub_run.italic = True
                    sub_run.font.size = docx.shared.Pt(10)
                    sub_run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
                    s.paragraph_format.space_after = docx.shared.Pt(18)

                # Body: split into sections and paragraphs
                body_sections = str(body_text or '').split('\n\n')
                for part in body_sections:
                    part = part.strip()
                    if not part:
                        continue
                    
                    # Check if this is a heading (all caps, short, or ends with colon)
                    is_heading = (
                        part.isupper() and len(part) < 80 or
                        part.endswith(':') and len(part) < 100 or
                        any(part.startswith(prefix) for prefix in ['Strengths:', 'Weaknesses:', 'Recommendations:', 'Strategies:', 'Summary:'])
                    )
                    
                    if is_heading:
                        # Format as section heading
                        heading_para = doc.add_paragraph()
                        heading_run = heading_para.add_run(part)
                        heading_run.bold = True
                        heading_run.font.size = docx.shared.Pt(14)
                        heading_run.font.color.rgb = docx.shared.RGBColor(68, 114, 196)
                        heading_para.paragraph_format.space_before = docx.shared.Pt(12)
                        heading_para.paragraph_format.space_after = docx.shared.Pt(6)
                    else:
                        # Regular paragraph
                        p = doc.add_paragraph(part)
                        p.paragraph_format.space_after = docx.shared.Pt(6)
                        p.paragraph_format.line_spacing = 1.15
                        
                        # If it's a bulleted list item, apply bullet style
                        if part.startswith('- ') or part.startswith('• '):
                            p.style = 'List Bullet'

                # Footer with timestamp and branding
                try:
                    footer = doc.sections[0].footer
                    fpara = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                    fpara.text = f"Generated by GENTA System | {ftime}"
                    fpara.alignment = 1  # Center
                    footer_run = fpara.runs[0] if fpara.runs else None
                    if footer_run:
                        footer_run.font.size = docx.shared.Pt(9)
                        footer_run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
                except Exception:
                    pass

                # Ensure directory exists and save
                os.makedirs(os.path.dirname(out_path), exist_ok=True)
                doc.save(out_path)
                return True
            except Exception as e:
                print(f"Error creating templated docx: {e}")
                import traceback
                traceback.print_exc()
                return False

        # Read conversation, normalize some phrases for the model
        with open(conversation_log, 'r', encoding='utf-8') as file:
            conversation_text3 = file.read()
        conversation_text2 = conversation_text3.replace("Tama", "Correct")
        conversation_text = conversation_text2.replace(" Ang tamang sagot ay", "The correct answer is")

        # Remove metadata entries from conversation_text (they confuse teachers)
        try:
            conversation_text = re.sub(r'(?m)^[ \t]*\[METADATA\].*\n?', '', conversation_text)
        except Exception:
            pass

        # Generate analysis via model
        analysis_result = model.generate_content(str(messages) + " " + conversation_text)
        analysis_text = analysis_result.text if hasattr(analysis_result, 'text') else str(analysis_result)

        # Sanitize analysis text: strip markdown markers, stray asterisks, and metadata
        try:
            analysis_text = sanitize_module_text(analysis_text)
            # Also remove any remaining repeated header markers like '***' or similar
            analysis_text = re.sub(r'\*{1,}', '', analysis_text)
            analysis_text = re.sub(r'(?m)^[ \t]*#{1,6}\s*', '', analysis_text)
        except Exception:
            pass

        # Build title/subtitle and save templated docx
        title = "Quiz Analysis Result"
        subtitle = f"Generated: {ftime}"
        ok = create_templated_docx(title, subtitle, analysis_text + "\n\nConversation Log:\n" + conversation_text, output_docx_path)
        if ok:
            print(f"Successfully created {output_docx_path}")
        else:
            print(f"Failed to create {output_docx_path}")

        # Also save a plaintext preview so the Flask UI can render it quickly
        try:
            txt_path = os.path.splitext(output_docx_path)[0] + '.txt'
            with open(txt_path, 'w', encoding='utf-8') as tf:
                tf.write(analysis_text)
            print(f"Saved plain-text analysis preview: {txt_path}")
        except Exception as _e:
            print(f"Could not write text preview for analysis: {_e}")

    def sanitize_module_text(text: str) -> str:
        """Sanitize generated module text:
        - remove Markdown bold markers (** and __)
        - strip HTML tags but preserve inner text (convert <br> to newline)
        - remove fenced code blocks markers ``` and any leading fence language
        - unescape HTML entities (e.g., &nbsp;)
        - collapse excessive blank lines
        Keeps emojis intact.
        """
        if not text:
            return text
        try:
            # Unescape HTML entities first
            text = html.unescape(text)

            # Remove fenced code blocks (```lang ... ```) but keep inner content if needed
            # Remove fence lines themselves
            text = re.sub(r'```[A-Za-z0-9_-]*\n', '', text)
            text = text.replace('```', '')

            # Convert <br> and <br/> to newlines
            text = re.sub(r'<\s*br\s*/?>', '\n', text, flags=re.I)

            # Strip all other HTML tags but keep inner text
            text = re.sub(r'<[^>]+>', '', text)

            # Remove Markdown bold/italic markers (strong/emphasis)
            text = text.replace('**', '').replace('__', '')
            # Remove single asterisks used as bullets or emphasis
            # (we prefer to remove stray '*' to produce clean teacher-facing docs)
            text = text.replace('*', '')

            # Remove leading Markdown header hashes (e.g., '#', '##', '###') at start of lines
            # Keep the header text but strip the hashes and any following space
            text = re.sub(r'^[ \t]*#{1,6}\s*', '', text, flags=re.M)

            # Trim trailing spaces on each line
            text = '\n'.join([ln.rstrip() for ln in text.splitlines()])

            # Collapse multiple blank lines to maximum two
            text = re.sub(r'\n{3,}', '\n\n', text)

            # Remove any stray metadata markers commonly inserted into the log
            # e.g. lines starting with [METADATA]
            text = re.sub(r'(?m)^[ \t]*\[METADATA\].*\n?', '', text)

            return text.strip()
        except Exception:
            return text

    def tailored_module(conversation_log, quiz_analysis_file_path, output_docx_tailoredmodule_path):
        """Create a professional tailored learning module DOCX with cover, TOC, and styled content."""
        with open(conversation_log, 'r') as file:
            conversation_text3 = file.read()
            conversation_text2 = conversation_text3.replace("Tama", "Correct")
            conversation_text = conversation_text2.replace(" Ang tamang sagot ay", "The correct answer is")
        # Remove metadata lines so the model and final module don't include internal tags
        try:
            conversation_text = re.sub(r'(?m)^[ \t]*\[METADATA\].*\n?', '', conversation_text)
        except Exception:
            pass

        # Open the DOCX file and read its content
        with open(quiz_analysis_file_path, 'rb') as file:
            doc = docx.Document(file)
            quiz_analysis = ""
            for paragraph in doc.paragraphs:
                if "Lesson Plan" in paragraph.text:
                    break  # Exit the loop when "Lesson Plan" is found
                quiz_analysis += paragraph.text + "\n"

        # Generate first part of module
        try:
            prompt_for_module = conversation_text + " " + str(module_messages)
        except Exception:
            prompt_for_module = str(module_messages)
        tailored_module_result = model.generate_content(prompt_for_module)
        module_text_part1 = tailored_module_result.text if hasattr(tailored_module_result, 'text') else str(tailored_module_result)
        # Sanitize generated text to remove unwanted markdown bold markers and HTML tags
        module_text_part1 = sanitize_module_text(module_text_part1)

        # Create professional DOCX with cover page
        try:
            module_doc = docx.Document()
            
            # ============ DOCUMENT-WIDE STYLES ============
            style = module_doc.styles['Normal']
            try:
                style.font.name = 'Calibri'
                style.font.size = docx.shared.Pt(11)
            except Exception:
                pass

            # ============ COVER PAGE ============
            cover_title = module_doc.add_paragraph()
            cover_title.alignment = 1  # Center
            run = cover_title.add_run("GENTA")
            run.bold = True
            run.font.size = docx.shared.Pt(48)
            run.font.color.rgb = docx.shared.RGBColor(0, 153, 51)  # Green for learning
            
            cover_subtitle = module_doc.add_paragraph()
            cover_subtitle.alignment = 1
            run2 = cover_subtitle.add_run("Personalized Learning Module")
            run2.font.size = docx.shared.Pt(20)
            run2.font.color.rgb = docx.shared.RGBColor(102, 187, 106)
            cover_subtitle.paragraph_format.space_after = docx.shared.Pt(24)
            
            # Decorative line
            module_doc.add_paragraph('_' * 60).alignment = 1
            module_doc.add_paragraph()
            
            # Student info
            info_para = module_doc.add_paragraph()
            info_para.alignment = 1
            try:
                student_name = CURRENT_STUDENT_NAME if CURRENT_STUDENT_NAME else "Student"
                student_id = CURRENT_STUDENT_ID if CURRENT_STUDENT_ID else "N/A"
                info_run = info_para.add_run(f"Prepared for: {student_name}\nLRN: {student_id}")
                info_run.font.size = docx.shared.Pt(14)
                info_run.bold = True
            except Exception:
                pass
            info_para.paragraph_format.space_after = docx.shared.Pt(36)
            
            meta_para = module_doc.add_paragraph()
            meta_para.alignment = 1
            meta_run = meta_para.add_run(f"Generated: {ftime}")
            meta_run.font.size = docx.shared.Pt(12)
            meta_run.italic = True
            meta_para.paragraph_format.space_after = docx.shared.Pt(48)
            
            module_doc.add_page_break()
            
            # ============ TABLE OF CONTENTS ============
            toc_heading = module_doc.add_paragraph()
            toc_run = toc_heading.add_run("Table of Contents")
            toc_run.bold = True
            toc_run.font.size = docx.shared.Pt(18)
            toc_run.font.color.rgb = docx.shared.RGBColor(0, 153, 51)
            toc_heading.paragraph_format.space_after = docx.shared.Pt(12)
            
            toc_entries = [
                ("1. Learning Objectives", "Page 3"),
                ("2. Foundational Concepts", "Page 4"),
                ("3. Step-by-Step Examples", "Page 5"),
                ("4. Practice Problems", "Page 6"),
                ("5. Real-World Applications", "Page 7"),
            ]
            
            for entry, page in toc_entries:
                toc_entry = module_doc.add_paragraph()
                toc_entry.paragraph_format.left_indent = docx.shared.Inches(0.25)
                entry_run = toc_entry.add_run(entry)
                entry_run.font.size = docx.shared.Pt(12)
                dots = toc_entry.add_run(" " + "." * 50)
                dots.font.color.rgb = docx.shared.RGBColor(192, 192, 192)
                page_run = toc_entry.add_run(f" {page}")
                page_run.font.size = docx.shared.Pt(12)
                toc_entry.paragraph_format.space_after = docx.shared.Pt(3)
            
            module_doc.add_page_break()

            # ============ HEADER & FOOTER ============
            try:
                section = module_doc.sections[0]
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.text = 'GENTA — Tailored Learning Module'
                header_para.alignment = 1
                header_line = header.add_paragraph('_' * 80)
                header_line.alignment = 1
            except Exception:
                pass

            # ============ MODULE CONTENT PART 1 ============
            main_title = module_doc.add_paragraph()
            title_run = main_title.add_run("Tailored Learning Module")
            title_run.bold = True
            title_run.font.size = docx.shared.Pt(20)
            title_run.font.color.rgb = docx.shared.RGBColor(0, 153, 51)
            main_title.paragraph_format.space_after = docx.shared.Pt(12)
            
            timestamp = module_doc.add_paragraph()
            timestamp_run = timestamp.add_run(f"Generated: {ftime}")
            timestamp_run.italic = True
            timestamp_run.font.size = docx.shared.Pt(10)
            timestamp_run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
            timestamp.paragraph_format.space_after = docx.shared.Pt(18)
            
            # Add first part content with formatting
            for part in module_text_part1.split('\n\n'):
                part = part.strip()
                if not part:
                    continue
                
                # Check if heading
                is_heading = (
                    part.isupper() and len(part) < 80 or
                    part.endswith(':') and len(part) < 100 or
                    any(part.startswith(prefix) for prefix in ['Topic:', 'Objective:', 'Concept:', 'Example:', 'Exercise:'])
                )
                
                if is_heading:
                    heading_para = module_doc.add_paragraph()
                    heading_run = heading_para.add_run(part)
                    heading_run.bold = True
                    heading_run.font.size = docx.shared.Pt(14)
                    heading_run.font.color.rgb = docx.shared.RGBColor(102, 187, 106)
                    heading_para.paragraph_format.space_before = docx.shared.Pt(12)
                    heading_para.paragraph_format.space_after = docx.shared.Pt(6)
                else:
                    p = module_doc.add_paragraph(part)
                    p.paragraph_format.space_after = docx.shared.Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                    if part.startswith('- ') or part.startswith('• '):
                        p.style = 'List Bullet'

            # Save intermediate version
            os.makedirs(os.path.dirname(output_docx_tailoredmodule_path), exist_ok=True)
            module_doc.save(output_docx_tailoredmodule_path)
            print("Writing tailored module 1/2")

        except Exception as e:
            print(f"Error creating module part 1: {e}")
            import traceback
            traceback.print_exc()
            # Fallback to simple version
            module_doc = docx.Document()
            module_doc.add_paragraph("TAILORED MODULE: " + ftime)
            module_doc.add_paragraph(module_text_part1)
            module_doc.save(output_docx_tailoredmodule_path)

        # ============ GENERATE PART 2 ============
        # Read back the document
        with open(output_docx_tailoredmodule_path, 'rb') as file:
            module_doc = docx.Document(file)
            previous_exported_docx = "\n".join([paragraph.text for paragraph in module_doc.paragraphs])

        # Generate second part (examples and practice problems)
        tailored_module_result2 = model.generate_content(str(module_messages2) + " " + previous_exported_docx)
        module_text_part2 = tailored_module_result2.text if hasattr(tailored_module_result2, 'text') else str(tailored_module_result2)
        # Sanitize generated text to remove unwanted markdown bold markers and HTML tags
        module_text_part2 = sanitize_module_text(module_text_part2)

        # Append part 2 with formatting
        try:
            module_doc.add_page_break()
            
            section_title = module_doc.add_paragraph()
            section_run = section_title.add_run("Practice & Application")
            section_run.bold = True
            section_run.font.size = docx.shared.Pt(18)
            section_run.font.color.rgb = docx.shared.RGBColor(0, 153, 51)
            section_title.paragraph_format.space_after = docx.shared.Pt(12)
            
            for part in module_text_part2.split('\n\n'):
                part = part.strip()
                if not part:
                    continue
                
                is_heading = (
                    part.isupper() and len(part) < 80 or
                    part.endswith(':') and len(part) < 100 or
                    any(part.startswith(prefix) for prefix in ['Problem:', 'Solution:', 'Answer:', 'Step:', 'Exercise:'])
                )
                
                if is_heading:
                    heading_para = module_doc.add_paragraph()
                    heading_run = heading_para.add_run(part)
                    heading_run.bold = True
                    heading_run.font.size = docx.shared.Pt(13)
                    heading_run.font.color.rgb = docx.shared.RGBColor(102, 187, 106)
                    heading_para.paragraph_format.space_before = docx.shared.Pt(10)
                    heading_para.paragraph_format.space_after = docx.shared.Pt(5)
                else:
                    p = module_doc.add_paragraph(part)
                    p.paragraph_format.space_after = docx.shared.Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                    if part.startswith('- ') or part.startswith('• ') or part[0].isdigit() and part[1:3] in ['. ', ') ']:
                        p.style = 'List Bullet'

            # Footer
            try:
                footer = module_doc.sections[0].footer
                fpara = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                fpara.text = f"Generated by GENTA System | {ftime}"
                fpara.alignment = 1
                footer_run = fpara.runs[0] if fpara.runs else None
                if footer_run:
                    footer_run.font.size = docx.shared.Pt(9)
                    footer_run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
            except Exception:
                pass

            module_doc.save(output_docx_tailoredmodule_path)
            print("Successfully created tailored module 2/2")
            
        except Exception as e:
            print(f"Error appending module part 2: {e}")
            import traceback
            traceback.print_exc()
            # Fallback: append simply
            module_doc.add_paragraph(module_text_part2)
            module_doc.save(output_docx_tailoredmodule_path)

        # Save plaintext preview
        try:
            txt_tailored = os.path.splitext(output_docx_tailoredmodule_path)[0] + '.txt'
            with open(txt_tailored, 'w', encoding='utf-8') as tf2:
                tf2.write(module_text_part1 + "\n\n" + module_text_part2)
            print(f"Saved plain-text tailored module preview: {txt_tailored}")
        except Exception as _e:
            print(f"Could not write text preview for tailored module: {_e}")

    # Load quiz questions from DATABASE (filtered by teacher_id)
    # Do NOT pass file_path - we want database questions only
    quiz_questions = load_quiz()
    
    # CRITICAL CHECK: Do not proceed if no questions were loaded
    if not quiz_questions or len(quiz_questions) == 0:
        print("\n" + "="*70)
        print("⚠ ERROR: No quiz questions found!")
        print("="*70)
        if CURRENT_TEACHER_ID:
            print(f"No questions found for Teacher ID: {CURRENT_TEACHER_ID}")
            print("\nPossible reasons:")
            print("1. No questions in database for this teacher")
            print("2. All questions for this teacher have status = 0 (inactive)")
            print("3. Database connection failed")
        else:
            print("No teacher ID available - cannot load teacher-specific questions")
        print("\nPlease:")
        print("- Check the 'questions' table in MySQL database 'my_app'")
        print("- Ensure teacher_id column matches the student's teacher")
        print("- Ensure at least one question has status = 1 (active)")
        print("="*70)
        
        # Return to Assisting Mode
        try:
            set_state_url = "https://nonbasic-bob-inimical.ngrok-free.dev/set_state"
            requests.get(f"{set_state_url}?value=0", timeout=15)
            print("✓ Returning to Assisting Mode")
        except Exception:
            pass
        return
    
    print(f"\n✓ Starting quiz with {len(quiz_questions)} questions\n")
    # Start a small loading animation on the OLED while we prepare the quiz
    def _start_quiz_loading_animation(student_name=None, teacher_id=None):
        """Start a background thread that cycles expressions until stopped.

        Sets a global `_quiz_loading_stop_event` threading.Event which can be set to stop the animation.
        """
        global _quiz_loading_stop_event
        try:
            _quiz_loading_stop_event = threading.Event()
        except Exception:
            return

        # Prepare an audible loading cue: synthesize once and upload to ESP (WelcomeAudio)
        loading_basename = None
        try:
            tmp_loading = os.path.join(UPLOAD_DIR, 'quiz_loading.wav')
            # Short friendly cue in Filipino/English
            cue_text = "Loading quiz. Please wait a moment."
            try:
                synthesize_speech(cue_text, out_path=tmp_loading, sample_rate_hz=24000)
                # Upload to ESP WelcomeAudio path (so play_welcome can play it by name)
                up_ok = esp_upload_file(tmp_loading, endpoint='/upload_welcome', max_retries=2)
                if up_ok:
                    loading_basename = os.path.basename(tmp_loading)
            except Exception:
                loading_basename = None
        except Exception:
            loading_basename = None

        def _anim_loop():
            # Remove three-dot 'processing' frames during quiz loading; use calmer 'thinking' and 'excited'
            seq = ['thinking', 'excited']
            idx = 0
            play_count = 0
            # Determine play URL if upload succeeded
            play_url = None
            if loading_basename and esp_playback_host:
                try:
                    play_url = f'http://{esp_playback_host}/play_welcome?name={loading_basename}'
                except Exception:
                    play_url = None

            while not _quiz_loading_stop_event.is_set():
                try:
                    set_oled_expression(seq[idx % len(seq)])
                except Exception:
                    pass
                idx += 1
                # Play audible cue every ~4 cycles (~2.2s)
                play_count += 1
                if play_url and play_count >= 4:
                    try:
                        # Non-blocking short GET to trigger playback on ESP
                        _http_session.get(play_url, timeout=1.5)
                    except Exception:
                        pass
                    play_count = 0
                # short sleep to make animation noticeable but not too fast
                time.sleep(0.55)

            # Ensure idle when stopping
            try:
                set_oled_expression('idle')
            except Exception:
                pass

        t = threading.Thread(target=_anim_loop, daemon=True)
        t.start()

    # Kick off animation showing we're loading quiz/student info (guarded)
    # Only start this when explicitly requested (e.g., after confirmation)
    if start_loading:
        try:
            print(f"Starting quiz loading animation for student={CURRENT_STUDENT_NAME} teacher_id={CURRENT_TEACHER_ID}")
            _start_quiz_loading_animation(CURRENT_STUDENT_NAME, CURRENT_TEACHER_ID)
        except Exception:
            # If anything goes wrong starting the animation, continue without it
            pass
    
    # ========== CREATE STUDENT_QUIZ RECORD BEFORE QUIZ STARTS ==========
    # This fixes the NULL student_quiz_id error by creating a proper parent record
    student_quiz_id = None
    try:
        # mysql.connector already imported at top of file (line 8)
        con = mysql.connector.connect(
            host=getattr(config, 'MYSQL_HOST', 'localhost'),
            port=getattr(config, 'MYSQL_PORT', 3306),
            user=getattr(config, 'MYSQL_USER', ''),
            password=getattr(config, 'MYSQL_PASS', ''),
            database=getattr(config, 'MYSQL_DB', '')
        )
        cur = con.cursor()
        
        # STEP 1: Get the actual student_id (primary key) from students table using LRN
        actual_student_id = None
        subject_id = None
        
        cur.execute(
            "SELECT id FROM students WHERE lrn = %s",
            (CURRENT_STUDENT_ID,)  # CURRENT_STUDENT_ID is the LRN
        )
        student_result = cur.fetchone()
        
        if student_result:
            # Found student by LRN
            actual_student_id = student_result[0]
            print(f"✓ Found student_id: {actual_student_id} for LRN: {CURRENT_STUDENT_ID}")
        else:
            print(f"⚠ WARNING: Student with LRN {CURRENT_STUDENT_ID} not found in database!")
            print("  Will try to use first available student_id as fallback")
            # Try to get any student_id as fallback
            cur.execute("SELECT id FROM students LIMIT 1")
            fallback_result = cur.fetchone()
            if fallback_result:
                actual_student_id = fallback_result[0]
                print(f"  Using fallback student_id: {actual_student_id}")
            else:
                print("  No students in database - cannot create student_quiz record")
        
        # STEP 2: Get subject_id from the first question's teacher
        if actual_student_id:
            cur.execute(
                "SELECT subject_id FROM questions WHERE teacher_id = %s AND status = 1 LIMIT 1",
                (CURRENT_TEACHER_ID,)
            )
            result = cur.fetchone()
            
            if result:
                subject_id = result[0]
                print(f"✓ Found subject_id: {subject_id} for teacher_id: {CURRENT_TEACHER_ID}")
            else:
                print(f"⚠ WARNING: Could not find subject_id for teacher_id={CURRENT_TEACHER_ID}")
                # Try to get any subject_id as fallback
                cur.execute("SELECT id FROM subjects LIMIT 1")
                fallback_subject = cur.fetchone()
                if fallback_subject:
                    subject_id = fallback_subject[0]
                    print(f"  Using fallback subject_id: {subject_id}")
                else:
                    subject_id = 1  # Last resort fallback
                    print(f"  Using default subject_id: {subject_id}")
            
            # STEP 3: Create student_quiz record with proper foreign keys
            # Include total_questions=0 as placeholder (will be updated after quiz)
            cur.execute(
                """INSERT INTO student_quiz (student_id, subject_id, total_questions, created, modified)
                   VALUES (%s, %s, 0, NOW(), NOW())""",
                (actual_student_id, subject_id)
            )
            student_quiz_id = cur.lastrowid
            con.commit()
            print(f"✓ Created student_quiz record with ID: {student_quiz_id}")
        
        cur.close()
        con.close()
        
    except Exception as e:
        print(f"⚠ Error creating student_quiz record: {e}")
        print(f"  Error details: {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        # Don't use fallback ID 22 - it doesn't exist. Let it fail so we can debug.
        if student_quiz_id is None:
            print("  CRITICAL: student_quiz_id is None - quiz results won't be saved properly!")
    # ====================================================================
    
    # Stop quiz loading animation just before questioning begins
    try:
        global _quiz_loading_stop_event
        if _quiz_loading_stop_event and not _quiz_loading_stop_event.is_set():
            _quiz_loading_stop_event.set()
            print("[Quiz] Stopped quiz loading animation, starting questions now")
    except Exception:
        pass

    run_quiz(quiz_questions)
    
    # ========== CREATE STUDENT-SPECIFIC FILENAMES ==========
    # Include student name in filenames so website knows which student the reports belong to
    # Sanitize student name for filename (remove special characters, replace spaces with underscores)
    safe_student_name = CURRENT_STUDENT_NAME.replace(" ", "_").replace("/", "_").replace("\\", "_") if CURRENT_STUDENT_NAME else "Unknown"
    safe_lrn = CURRENT_STUDENT_ID if CURRENT_STUDENT_ID else "NoLRN"
    
    # Create unique filenames with student name and LRN (use canonical upload dir)
    student_analysis_path = U(f'analysis_result_{safe_student_name}_{safe_lrn}.docx')
    student_tailored_module_path = U(f'tailored_module_{safe_student_name}_{safe_lrn}.docx')
    
    print(f"\n{'='*70}")
    print(f"📄 CREATING STUDENT-SPECIFIC REPORTS")
    print(f"{'='*70}")
    # Start OLED report-creation animation until reports are finished (if model available)
    # Mark report creation as active; it will be cleared at the natural end
    skip_reports = False
    try:
        global _REPORT_CREATION_ACTIVE
        # Verify GenAI availability before starting long-running report generation
        if not globals().get('_GENAI_AVAILABLE', True):
            print("[Report] GenAI not currently available - attempting quick verify...")
            try:
                verify_genai_model()
            except Exception:
                pass

        if not globals().get('_GENAI_AVAILABLE', True):
            # Model still unavailable - create placeholder reports and skip full generation
            print("⚠ GenAI model unavailable. Skipping AI-driven report generation.")
            try:
                doc = docx.Document()
                doc.add_heading('Report Unavailable', level=1)
                doc.add_paragraph('The AI model required to generate the student analysis is not reachable at this time. Please check GENAI_API_KEY and network connectivity.')
                try:
                    os.makedirs(os.path.dirname(student_analysis_path), exist_ok=True)
                except Exception:
                    pass
                doc.save(student_analysis_path)
            except Exception as e:
                print(f"⚠ Could not write placeholder analysis file: {e}")
            try:
                doc2 = docx.Document()
                doc2.add_heading('Tailored Module Unavailable', level=1)
                doc2.add_paragraph('The AI model required to generate the tailored learning module is not reachable at this time.')
                try:
                    os.makedirs(os.path.dirname(student_tailored_module_path), exist_ok=True)
                except Exception:
                    pass
                doc2.save(student_tailored_module_path)
            except Exception as e:
                print(f"⚠ Could not write placeholder tailored module file: {e}")

            # Update progress to complete and avoid starting animation
            try:
                _report_progress = 100
            except Exception:
                pass
            skip_reports = True
        else:
            _REPORT_CREATION_ACTIVE = True
            try:
                # Reset completion-shown flag at start of a new report generation
                globals()['_REPORT_COMPLETION_SHOWN'] = False
            except Exception:
                pass
            try:
            
                print("[Report] _REPORT_CREATION_ACTIVE set to True")
            except Exception:
                pass
            # Notify local Flask proxy to enable guard against discovery clears
            try:
                try:
                    _http_session.get('http://127.0.0.1:5000/debug/set_report_guard', params={'active': 1}, timeout=1.0)
                    print('[Report] Notified Flask to set REPORT_CREATION_GUARD=1')
                except Exception:
                    # best-effort notification - ignore failures
                    pass
            except Exception:
                pass
            try:
                _start_report_creation_animation()
                print("[Quiz] Started report-creation OLED animation")
                try:
                    # initialize progress indicator
                    _report_progress = 5
                    print(f"[Report] _report_progress set to {_report_progress}")
                    # Quick health-check: attempt to ping the OLED proxy with the initial value
                    try:
                        # Dispatch the initial report progress probe asynchronously so
                        # the caller doesn't block on Flask forwarding to the device.
                        def _oled_probe_send(p):
                            try:
                                _http_session.get('http://127.0.0.1:5000/oled', params={'value': 'report', 'progress': p}, timeout=3)
                            except Exception:
                                # Silent fallback; this probe is best-effort only
                                pass

                        threading.Thread(target=_oled_probe_send, args=(_report_progress,), daemon=True).start()
                        print("[Quiz] Report animation probe dispatched (async)")
                    except Exception as _e:
                        print(f"[Quiz] Could not dispatch asynchronous OLED probe: {_e}")
                except Exception:
                    pass
            except Exception:
                pass
    except Exception:
        skip_reports = True

    print(f"Student: {CURRENT_STUDENT_NAME} (LRN: {CURRENT_STUDENT_ID})")
    print(f"Analysis file: {student_analysis_path}")
    print(f"Tailored module file: {student_tailored_module_path}")
    print(f"{'='*70}\n")
    # ========================================================
    
    try:
        _report_progress = 10
        print(f"[Report] _report_progress set to {_report_progress}")
    except Exception:
        pass
    if not skip_reports:
        quiz_analysis(conversation_file_path, student_analysis_path)
    else:
        print("[Report] Skipping quiz_analysis due to GenAI unavailability")
    try:
        _report_progress = 45
        print(f"[Report] _report_progress set to {_report_progress}")
    except Exception:
        pass
    print("Quiz was analyzed successfully.")
    try:
        _report_progress = 50
        print(f"[Report] _report_progress set to {_report_progress}")
    except Exception:
        pass
    if not skip_reports:
        tailored_module(conversation_file_path, student_analysis_path, student_tailored_module_path)
    else:
        print("[Report] Skipping tailored_module due to GenAI unavailability")
    try:
        _report_progress = 80
        print(f"[Report] _report_progress set to {_report_progress}")
    except Exception:
        pass
    print("Tailored module was created successfully.")
    # Create CSV File of the quiz results
    # Define the destination directory (canonical uploads folder)
    destination_directory = UPLOAD_DIR
    # Define the destination path
    destination_path = os.path.join(destination_directory, 'conversation_log.csv')          
    # Define the conversation log file path (now stored in the canonical uploads folder)
    log_file_path = conversation_file_path

    # Get the Unix timestamp for the file creation time
    created_timestamp = os.path.getctime(log_file_path)
    modified_timestamp = os.path.getmtime(log_file_path)

    # Convert the Unix timestamps to datetime objects
    created_datetime = datetime.fromtimestamp(created_timestamp)
    modified_datetime = datetime.fromtimestamp(modified_timestamp)

    # Define the desired datetime format (MySQL compatible: YYYY-MM-DD HH:MM:SS)
    datetime_format = "%Y-%m-%d %H:%M:%S"

    # Format the datetime objects
    formatted_created_datetime = created_datetime.strftime(datetime_format)
    formatted_modified_datetime = modified_datetime.strftime(datetime_format)

    # Define the CSV file path inside the canonical uploads folder
    csv_file_path = os.path.join(UPLOAD_DIR, 'conversation_log.csv')
    # Define the CSV fieldnames
    fieldnames = ["id", "student_quiz_id", "description", "image", "choices", "answer", "student_answer", "score", "status", "created", "modified"]
    # Initialize a list to store conversation log entries
    conversation_log = []

    # Read the conversation log file line by line
    from collections import deque
    with open(log_file_path, mode="r", encoding="utf-8") as logfile:
        current_entry = {}
        current_metadata = {}  # Store metadata from [METADATA] line
        recent_lines = deque(maxlen=12)  # keep last few raw lines for context/backfill
        
        for line in logfile:
            recent_lines.append(line.rstrip('\n'))
            # Check for metadata line (contains question_id, choices, image)
            if line.startswith("[METADATA]"):
                # Parse metadata: question_id=X, choices=Y, image=Z
                current_metadata = {}
                parts = line.replace("[METADATA]", "").strip().split(",")
                for part in parts:
                    if "=" in part:
                        key, val = part.split("=", 1)
                        current_metadata[key.strip()] = val.strip()
                continue
            
            # Check if the line contains the question number and description
            match = re.match(r"(\d+(?:st|nd|rd|th) question): (.+)", line.strip())
            if match:
                if current_entry:
                    # Avoid appending duplicate consecutive questions (e.g., when student asked to repeat)
                    prev_desc = current_entry.get('description', '').strip()
                    if not (conversation_log and conversation_log[-1].get('description', '').strip() == prev_desc):
                        conversation_log.append(current_entry)
                    else:
                        # Duplicate detected (likely a "repeat" request). Preserve any
                        # student-provided answer or resolved correct/incorrect info
                        # so it isn't lost when the repeated question is logged anew.
                        saved_student_answer = current_entry.get('student_answer', '')
                        saved_answer = current_entry.get('answer', '')
                        saved_score = current_entry.get('score', '')
                        saved_status = current_entry.get('status', '')
                        print(f"⚠ Skipping duplicate logged question (repeat detected): '{prev_desc[:40]}...' - preserving student_answer={saved_student_answer}")
                # Start a NEW entry with the actual student_quiz_id and metadata
                current_entry = {
                    "student_quiz_id": student_quiz_id,
                    "description": match.group(2),
                    "image": current_metadata.get('image', ''),
                    "choices": current_metadata.get('choices', ''),
                    "status": "1"  # Status is ALWAYS 1 for used questions (not 0!)
                }
                # If we preserved a student answer from the skipped duplicate, carry it
                # into the new entry so the final logged question contains the answer.
                try:
                    if 'saved_student_answer' in locals() and saved_student_answer:
                        current_entry["student_answer"] = saved_student_answer
                    if 'saved_answer' in locals() and saved_answer:
                        current_entry["answer"] = saved_answer
                    if 'saved_score' in locals() and saved_score:
                        current_entry["score"] = saved_score
                    if 'saved_status' in locals() and saved_status:
                        current_entry["status"] = saved_status
                except Exception:
                    pass
            elif "Student Answer:" in line:
                # Extract student answer
                student_answer = line.split(":")[1].strip()
                current_entry["student_answer"] = student_answer
            # Some recorder logs may record the student's spoken answer as a bare
            # line (e.g. "15." or "15") without the "Student Answer:" label.
            # If we haven't yet captured a student_answer for the current entry,
            # try to recognize such bare answers by comparing against the
            # question's `choices` metadata or by accepting a short numeric/text
            # token as a fallback.
            elif current_entry and not current_entry.get("student_answer"):
                stripped = line.strip()
                # Skip lines that look like labeled fields
                if stripped and ":" not in stripped:
                    # Normalize candidate (remove trailing dot)
                    candidate = stripped.rstrip('.')
                    try:
                        raw_choices = current_entry.get('choices', '')
                        choices_list = ast.literal_eval(raw_choices) if raw_choices else []
                    except Exception:
                        choices_list = []
                    choices_norm = [str(c).strip() for c in choices_list]
                    # If the candidate matches one of the known choices, accept it
                    if choices_norm and candidate in choices_norm:
                        current_entry["student_answer"] = candidate
                    else:
                        # Fallback: accept short alphanumeric tokens as answers
                        if len(candidate) <= 8 and re.match(r'^[A-Za-z0-9]+$', candidate):
                            current_entry["student_answer"] = candidate
            elif "Tama!" in line or "Tama" in line:
                # Correct answer - answer field should be same as student answer
                # Status is already set to "1" when entry was created
                if "answer" not in current_entry:
                    current_entry["answer"] = current_entry.get("student_answer", "")
                # If student_answer missing, try to backfill from recent raw lines
                if not current_entry.get('student_answer'):
                    # search backward in recent_lines for a bare answer candidate
                    for raw in reversed(list(recent_lines)):
                        if raw.startswith('[METADATA]') or raw.startswith(str((len(recent_lines)))):
                            continue
                        cand = raw.strip()
                        if cand and ":" not in cand:
                            cand = cand.rstrip('.')
                            try:
                                choices_list = ast.literal_eval(current_entry.get('choices', '')) if current_entry.get('choices') else []
                            except Exception:
                                choices_list = []
                            choices_norm = [str(c).strip() for c in choices_list]
                            if choices_norm and cand in choices_norm:
                                current_entry['student_answer'] = cand
                                break
                            if len(cand) <= 8 and re.match(r'^[A-Za-z0-9]+$', cand):
                                current_entry['student_answer'] = cand
                                break

                # Deduplicate before appending
                this_desc = current_entry.get('description', '').strip()
                if not (conversation_log and conversation_log[-1].get('description', '').strip() == this_desc):
                    conversation_log.append(current_entry)
                else:
                    print(f"⚠ Skipping duplicate logged question on correct answer: '{this_desc[:40]}...'")
                current_entry = {}  # Reset for next question
                current_metadata = {}  # Reset metadata
            elif "Oops! Ang tamang sagot ay" in line:
                # Extract correct answer
                correct_answer = line.split("ay")[1].strip()
                current_entry["answer"] = correct_answer
                current_entry["status"] = "1"  # Keep status as 1 (question was used)
                # If student_answer missing, try to backfill from recent lines
                if not current_entry.get('student_answer'):
                    for raw in reversed(list(recent_lines)):
                        cand = raw.strip()
                        if cand and ":" not in cand:
                            cand = cand.rstrip('.')
                            try:
                                choices_list = ast.literal_eval(current_entry.get('choices', '')) if current_entry.get('choices') else []
                            except Exception:
                                choices_list = []
                            choices_norm = [str(c).strip() for c in choices_list]
                            if choices_norm and cand in choices_norm:
                                current_entry['student_answer'] = cand
                                break
                            if len(cand) <= 8 and re.match(r'^[A-Za-z0-9]+$', cand):
                                current_entry['student_answer'] = cand
                                break

                # Deduplicate before appending
                this_desc = current_entry.get('description', '').strip()
                if not (conversation_log and conversation_log[-1].get('description', '').strip() == this_desc):
                    conversation_log.append(current_entry)
                else:
                    print(f"⚠ Skipping duplicate logged question on incorrect answer: '{this_desc[:40]}...'")
                current_entry = {}  # Reset for next question
                current_metadata = {}  # Reset metadata

    # Add the last entry to conversation log
    if current_entry:
        last_desc = current_entry.get('description', '').strip()
        if not (conversation_log and conversation_log[-1].get('description', '').strip() == last_desc):
            conversation_log.append(current_entry)
        else:
            print(f"⚠ Skipping duplicate trailing question entry: '{last_desc[:40]}...'")
    
    # DEBUG: Show what we collected from the log file
    print(f"\n{'='*70}")
    print(f"📝 CONVERSATION LOG PARSING RESULTS")
    try:
        if _report_creation_stop_event and not _report_creation_stop_event.is_set():
            try:
                _report_progress = 100
            except Exception:
                pass
            # Give the ESP a chance to play the completion animation and confirm it
            # was shown. Poll the local Flask proxy (/oled_status) until the
            # recorder reports the completion frame was shown, or until timeout.
            try:
                # Wait longer for the recorder to confirm the final animation.
                # Configurable via config.REPORT_CONFIRM_TIMEOUT (seconds).
                timeout = float(getattr(config, 'REPORT_CONFIRM_TIMEOUT', 60.0))
                poll_interval = 0.35
                start = time.time()
                completed = False
                status_url = 'http://127.0.0.1:5000/oled_status'
                print(f"[Quiz] Waiting up to {timeout}s for recorder to confirm report completion...")
                # While waiting, repeatedly send an explicit 100% progress
                # update to the Flask proxy so the ESP receives the target=100
                # frequently and its internal displayedProgress can reach 100.
                attempt_num = 0
                while time.time() - start < timeout:
                    attempt_num += 1
                    # Send a 100% progress update to the OLED proxy each loop (best-effort)
                    try:
                        oled_url = 'http://127.0.0.1:5000/oled'
                        params = {'value': 'report', 'progress': 100}
                        resp_put = http_get_with_retries(oled_url, params=params, timeout=4.0, retries=2, backoff=0.4)
                        if resp_put is None:
                            print(f"[ReportAnim] Could not send 100% progress to proxy (attempt {attempt_num})")
                    except Exception as e:
                        print(f"[ReportAnim] Exception when sending progress to proxy: {e}")

                    # Poll the proxy for recorder status (completion_played)
                    try:
                        r = http_get_with_retries(status_url, timeout=4.0, retries=3, backoff=0.5)
                        if r is not None and getattr(r, 'status_code', None) == 200:
                            try:
                                j = r.json()
                                if j.get('completion_played') in (True, 'true', 'True', '1'):
                                    completed = True
                                    try:
                                        globals()['_REPORT_COMPLETION_SHOWN'] = True
                                    except Exception:
                                        pass
                                    print("[Quiz] Recorder confirmed report completion animation shown")
                                    break
                            except Exception as je:
                                print(f"[ReportAnim] Could not parse oled_status JSON: {je}")
                        else:
                            code = getattr(r, 'status_code', None) if r is not None else None
                            print(f"[ReportAnim] oled_status not ready (status={code})")
                            # Try a direct device probe as a fallback when local proxy seems unresponsive
                            try:
                                direct = try_direct_oled_status(esp_record_host, timeout=4.0)
                                if direct and direct.get('completion_played') in (True, 'true', 'True', 1, '1'):
                                    completed = True
                                    print("[ReportAnim] Recorder (direct) confirmed report completion animation shown")
                                    break
                                else:
                                    print("[ReportAnim] Direct device oled_status did not confirm completion")
                            except Exception as de:
                                print(f"[ReportAnim] Direct device probe failed: {de}")
                    except Exception as e:
                        print(f"[ReportAnim] Exception while polling oled_status: {e}")

                    # Gentle backoff between polls to avoid hammering the proxy when it's down
                    time.sleep(max(poll_interval, 0.35))

                if not completed:
                    # Try an extended retry period before giving up
                    extra_wait = float(getattr(config, 'REPORT_CONFIRM_EXTRA_WAIT', 30.0))
                    print(f"[Quiz] Recorder did not confirm within {timeout}s; attempting an extra {extra_wait}s retry before proceeding...")
                    start2 = time.time()
                    while time.time() - start2 < extra_wait:
                        try:
                            try:
                                _http_session.get('http://127.0.0.1:5000/oled', params={'value': 'report', 'progress': 100}, timeout=4.0)
                            except Exception:
                                pass
                            r2 = _http_session.get(status_url, timeout=4.0)
                            if r2 is not None and getattr(r2, 'status_code', None) == 200:
                                try:
                                    j2 = r2.json()
                                    if j2.get('completion_played') in (True, 'true', 'True', '1'):
                                            completed = True
                                            try:
                                                globals()['_REPORT_COMPLETION_SHOWN'] = True
                                            except Exception:
                                                pass
                                            print("[Quiz] Recorder confirmed report completion animation shown (during extra retry)")
                                            break
                                except Exception:
                                    pass
                        except Exception:
                            pass
                        time.sleep(poll_interval)
                    if not completed:
                        print("[Quiz] Warning: recorder did not confirm completion after extra retry; proceeding anyway")
            except Exception:
                pass

            # Now stop the report animation loop
            _report_creation_stop_event.set()
            print("[Quiz] Stopped report-creation OLED animation")
    except Exception:
        pass
    # Clear active flag so state-change prompts and confirmations can resume.
    try:
        _REPORT_CREATION_ACTIVE = False
    except Exception:
        pass
    try:
        print("[Report] _REPORT_CREATION_ACTIVE cleared (False)")
    except Exception:
        pass
    # Notify Flask to clear the guard so registry clears can proceed
    try:
        try:
            _http_session.get('http://127.0.0.1:5000/debug/set_report_guard', params={'active': 0}, timeout=1.0)
            print('[Report] Notified Flask to clear REPORT_CREATION_GUARD')
        except Exception:
            pass
    except Exception:
        pass
    # Build updated conversation rows from parsed conversation_log
    updated_conversation_log = []
    for e in conversation_log:
        row = {
            "id": "",
            "student_quiz_id": e.get("student_quiz_id", student_quiz_id),
            "description": e.get("description", ""),
            "image": e.get("image", ""),
            "choices": e.get("choices", ""),
            "answer": e.get("answer", e.get("student_answer", "")),
            "student_answer": e.get("student_answer", ""),
            "score": e.get("score", "0"),
            "status": e.get("status", "1"),
            "created": formatted_created_datetime,
            "modified": formatted_modified_datetime
        }
        updated_conversation_log.append(row)

    # Write CSV directly into the uploads folder
    with open(csv_file_path, mode="w", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        # Write data without header
        for entry in updated_conversation_log:
            writer.writerow(entry)
    print("CSV file exported successfully.")
    try:
        _report_progress = 90
    except Exception:
        pass
    
    # DEBUG: Check CSV content before database insertion
    print(f"\n{'='*70}")
    print(f"📄 CSV FILE CONTENT CHECK")
    print(f"{'='*70}")
    print(f"CSV file location: {destination_path}")
    try:
        with open(destination_path, mode="r", encoding="utf-8") as check_file:
            csv_content = check_file.read()
            line_count = len(csv_content.split('\n'))
            print(f"Total lines in CSV: {line_count}")
            print(f"First 500 characters:\n{csv_content[:500]}")
    except Exception as e:
        print(f"⚠ Could not read CSV: {e}")
    print(f"{'='*70}\n")

    try:
        con = mysql.connector.connect(
            host=getattr(config, 'MYSQL_HOST', 'localhost'),
            port=getattr(config, 'MYSQL_PORT', 3306),
            database=getattr(config, 'MYSQL_DB', ''),
            user=getattr(config, 'MYSQL_USER', ''),
            password=getattr(config, 'MYSQL_PASS', '')
        )
        cur = con.cursor()

        print(f"\n{'='*70}")
        print(f"📊 INSERTING QUIZ RESULTS INTO DATABASE")
        print(f"{'='*70}")
        print(f"Student Quiz ID: {student_quiz_id}")
        print(f"Reading from CSV: {destination_path}")
        
        # If student_quiz_id is None, we need to create the student_quiz record now
        if student_quiz_id is None:
            print("⚠ student_quiz_id is None - attempting to create student_quiz record now...")
            try:
                # Get actual student_id from LRN
                cur.execute("SELECT id FROM students WHERE lrn = %s", (CURRENT_STUDENT_ID,))
                student_row = cur.fetchone()
                if student_row:
                    actual_student_id = student_row[0]
                    print(f"  ✓ Found student_id: {actual_student_id} for LRN: {CURRENT_STUDENT_ID}")
                    
                    # Get subject_id
                    cur.execute("SELECT subject_id FROM questions WHERE teacher_id = %s AND status = 1 LIMIT 1", (CURRENT_TEACHER_ID,))
                    subj_row = cur.fetchone()
                    subject_id = subj_row[0] if subj_row else 1
                    print(f"  ✓ Using subject_id: {subject_id}")
                    
                    # Create student_quiz record
                    cur.execute(
                        """INSERT INTO student_quiz (student_id, subject_id, total_questions, created, modified)
                           VALUES (%s, %s, 0, NOW(), NOW())""",
                        (actual_student_id, subject_id)
                    )
                    student_quiz_id = cur.lastrowid
                    con.commit()
                    print(f"  ✓ Created student_quiz record with ID: {student_quiz_id}")
                else:
                    print(f"  ❌ Could not find student with LRN: {CURRENT_STUDENT_ID}")
            except Exception as create_err:
                print(f"  ❌ Failed to create student_quiz record: {create_err}")
        
        inserted_count = 0
        # Read the CSV file and iterate over its rows
        with open(destination_path, mode="r", newline="", encoding="utf-8") as csvfile:
            reader = csv.reader(csvfile)
            # Iterate over the rows and construct queries
            for row_num, row in enumerate(reader, 1):
                try:
                    # Skip completely empty rows
                    if not row or all(not cell.strip() for cell in row):
                        print(f"⚠ Skipping empty row {row_num}")
                        continue
                    
                    # Skip if we don't have a valid student_quiz_id
                    if student_quiz_id is None:
                        print(f"❌ Skipping row {row_num}: student_quiz_id is None")
                        continue
                    
                    # Use parameterized query to avoid SQL syntax errors and injection
                    # This properly handles empty strings, quotes, and special characters
                    query = """INSERT INTO `student_quiz_questions` 
                               (`student_quiz_id`, `description`, `image`, `choices`, `answer`, 
                                `student_answer`, `score`, `status`, `created`, `modified`) 
                               VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)"""
                    
                    # Use the student_quiz_id variable directly (not from CSV which may be empty)
                    # CSV columns: 0=id, 1=student_quiz_id, 2=description, 3=image, 4=choices, 
                    #              5=answer, 6=student_answer, 7=score, 8=status, 9=created, 10=modified
                    values = (
                        student_quiz_id,  # Use the variable, not row[1]
                        row[2] if len(row) > 2 and row[2] else '',    # description
                        row[3] if len(row) > 3 and row[3] else '',    # image
                        row[4] if len(row) > 4 and row[4] else '',    # choices
                        row[5] if len(row) > 5 and row[5] else '',    # answer
                        row[6] if len(row) > 6 and row[6] else '',    # student_answer
                        row[7] if len(row) > 7 and row[7] else '0',   # score
                        row[8] if len(row) > 8 and row[8] else '1',   # status (default 1)
                        row[9] if len(row) > 9 and row[9] else None,  # created
                        row[10] if len(row) > 10 and row[10] else None # modified
                    )
                    
                    # Execute the query with parameterized values
                    cur.execute(query, values)
                    inserted_count += 1
                    desc_preview = values[1][:50] if values[1] else 'N/A'
                    print(f"✓ Row {row_num}: Inserted question - student_quiz_id={student_quiz_id}, description='{desc_preview}...'")
                except IndexError as idx_err:
                    print(f"⚠ Error: Row {row_num} does not contain expected number of fields: {idx_err}")
                    print(f"  Row content (len={len(row)}): {row[:5]}...")  # Show first 5 fields for debugging
                except Exception as e:
                    print(f"❌ Error on row {row_num}: {e}")
                    print(f"  Problematic row: {row[:5]}...")  # Show first 5 fields for debugging

        # Commit changes to the database
        con.commit()
        print(f"\n✓ Successfully inserted {inserted_count} quiz questions into database")
        print(f"{'='*70}\n")
        try:
            _report_progress = 98
        except Exception:
            pass
        
        # Update the parent student_quiz record with the total number of questions
        try:
            # total_questions is stored on the student_quiz parent record (id = student_quiz_id)
            cur.execute(
                """UPDATE student_quiz SET total_questions = %s WHERE id = %s""",
                (inserted_count, student_quiz_id)
            )
            con.commit()
            print(f"✓ Updated student_quiz (id={student_quiz_id}) with total_questions = {inserted_count}")
        except Exception as e:
            print(f"⚠ Could not update total_questions in student_quiz table: {e}")
        # Update the student's remarks field to record this assessment attempt.
        # Overwrite any existing remarks each time the student takes an assessment.
        try:
            if CURRENT_STUDENT_ID:
                try:
                    # Resolve actual student primary key (id) from LRN (CURRENT_STUDENT_ID stores LRN)
                    student_pk = None
                    try:
                        lookup_cur = con.cursor()
                        lookup_cur.execute("SELECT id FROM students WHERE lrn = %s", (CURRENT_STUDENT_ID,))
                        row = lookup_cur.fetchone()
                        if row:
                            student_pk = row[0]
                        else:
                            # If lookup failed, assume CURRENT_STUDENT_ID might already be the PK
                            student_pk = CURRENT_STUDENT_ID
                        lookup_cur.close()
                    except Exception:
                        student_pk = CURRENT_STUDENT_ID

                    # Compose a short teacher remark including score and percentage
                    # Determine score by summing inserted question scores for this student_quiz_id
                    try:
                        score = None
                        total_questions = inserted_count if inserted_count and inserted_count > 0 else None
                        if student_quiz_id and total_questions:
                            score_cur = con.cursor()
                            try:
                                score_cur.execute("SELECT SUM(CAST(score AS SIGNED)) FROM student_quiz_questions WHERE student_quiz_id = %s", (student_quiz_id,))
                                srow = score_cur.fetchone()
                                score = srow[0] if srow and srow[0] is not None else 0
                            except Exception:
                                score = None
                            finally:
                                score_cur.close()
                        else:
                            score = None

                        pct = (score / total_questions) * 100 if (score is not None and total_questions and total_questions > 0) else None
                    except Exception:
                        pct = None

                    if pct is not None:
                        # Build a teacher-facing short analysis (NO numeric score visible)
                        # Determine weak topics by inspecting incorrect question descriptions
                        try:
                            weak_descriptions = []
                            detail_cur = con.cursor()
                            try:
                                detail_cur.execute(
                                    "SELECT description FROM student_quiz_questions WHERE student_quiz_id = %s AND (score = 0 OR score IS NULL) LIMIT 6",
                                    (student_quiz_id,)
                                )
                                rows = detail_cur.fetchall()
                                if rows:
                                    weak_descriptions = [r[0] for r in rows if r and r[0]]
                            except Exception:
                                weak_descriptions = []
                            finally:
                                detail_cur.close()
                        except Exception:
                            weak_descriptions = []

                        # Simple keyword extraction from descriptions
                        def extract_keywords(descs, max_k=3):
                            if not descs:
                                return []
                            text = ' '.join(descs).lower()
                            tokens = re.findall(r"\b[a-z]{3,}\b", text)
                            stop = set(['the','and','for','with','when','what','which','that','this','are','from','use','about','problem','problems'])
                            counts = {}
                            for t in tokens:
                                if t in stop:
                                    continue
                                counts[t] = counts.get(t, 0) + 1
                            if not counts:
                                return []
                            # sort by freq
                            items = sorted(counts.items(), key=lambda x: x[1], reverse=True)
                            return [it[0] for it in items[:max_k]]

                        weak_topics = extract_keywords(weak_descriptions, max_k=3)

                        # Craft concise teacher message (no numbers). Keep it short.
                        if pct >= 90:
                            remarks_val = "Shows strong understanding of the assessed topics; recommend enrichment and more challenging tasks."
                        elif pct >= 70:
                            if weak_topics:
                                remarks_val = f"Good understanding overall; focus targeted practice on: {', '.join(weak_topics)}."
                            else:
                                remarks_val = "Good understanding overall; continue with targeted practice on weaker items."
                        elif pct >= 40:
                            if weak_topics:
                                remarks_val = f"Struggles with several items. Focus on: {', '.join(weak_topics)} and review foundational steps."
                            else:
                                remarks_val = "Shows gaps in understanding. Focus on foundational concepts and guided practice."
                        else:
                            if weak_topics:
                                remarks_val = f"Needs additional support; concentrate on: {', '.join(weak_topics)} with step-by-step instruction."
                            else:
                                remarks_val = "Needs additional support; begin with reviewing basic skills and guided instruction."

                        # If percentage couldn't be computed, use a generic recorded message
                        if pct is None:
                            remarks_val = f"Assessment recorded: {inserted_count} questions logged on {formatted_modified_datetime}"

                    # Perform the DB update using a fresh cursor and ensure it's closed
                    upd_cur = con.cursor()
                    try:
                        upd_cur.execute("""UPDATE students SET remarks = %s WHERE id = %s""", (remarks_val, student_pk))
                        con.commit()
                        print(f"✓ Updated students.remarks for student id={student_pk}: {remarks_val}")
                    finally:
                        try:
                            upd_cur.close()
                        except Exception:
                            pass
                except Exception as ue:
                    print(f"⚠ Could not update students.remarks for id={CURRENT_STUDENT_ID} (lookup as {student_pk}): {ue}")
        except Exception:
            pass
        
        # Note: Removed the problematic UPDATE query that was updating last 10 entries
        # Quiz results should not be modified after insertion
        
    except Error as error:
        print(f"❌ Insert data failed due to {error}")
    finally:
        if con.is_connected():
            # Close cursor and database connection
            cur.close()
            con.close()
            print("MySQL connection is closed")
    
    # AUTOMATIC RETURN TO ASSISTING MODE after quiz completes
    # Stop report-creation animation (if running) before finishing
    try:
        if _report_creation_stop_event and not _report_creation_stop_event.is_set():
            try:
                _report_progress = 100
            except Exception:
                pass
            # Give the ESP a chance to play the completion animation and confirm it
            # was shown. Poll the local Flask proxy (/oled_status) until the
            # recorder reports the completion frame was shown, or until timeout.
            try:
                timeout = float(getattr(config, 'REPORT_CONFIRM_TIMEOUT', 60.0))
                poll_interval = 0.35
                start = time.time()
                completed = False
                status_url = 'http://127.0.0.1:5000/oled_status'
                print(f"[Quiz] Waiting up to {timeout}s for recorder to confirm report completion...")
                # While waiting, repeatedly send an explicit 100% progress
                # update to the Flask proxy so the ESP receives the target=100
                # frequently and its internal displayedProgress can reach 100.
                while time.time() - start < timeout:
                    try:
                        # Send a 100% progress update to the OLED proxy each loop
                        try:
                            _http_session.get('http://127.0.0.1:5000/oled', params={'value': 'report', 'progress': 100}, timeout=0.8)
                        except Exception:
                            pass

                        r = _http_session.get(status_url, timeout=1.0)
                        if r is not None and getattr(r, 'status_code', None) == 200:
                            try:
                                j = r.json()
                                if j.get('completion_played') in (True, 'true', 'True', '1'):
                                    completed = True
                                    try:
                                        globals()['_REPORT_COMPLETION_SHOWN'] = True
                                    except Exception:
                                        pass
                                    print("[Quiz] Recorder confirmed report completion animation shown")
                                    break
                            except Exception:
                                pass
                        else:
                            # If proxy did not respond successfully, attempt direct device check
                            try:
                                direct = try_direct_oled_status(esp_record_host, timeout=4.0)
                                if direct and direct.get('completion_played') in (True, 'true', 'True', 1, '1'):
                                    completed = True
                                    print("[Quiz] Recorder (direct) confirmed report completion animation shown")
                                    break
                                else:
                                    print("[ReportAnim] Direct probe did not confirm completion; continuing")
                            except Exception as _e:
                                print(f"[ReportAnim] Direct probe error: {_e}")
                    except Exception:
                        pass
                    time.sleep(poll_interval)

                if not completed:
                    extra_wait = float(getattr(config, 'REPORT_CONFIRM_EXTRA_WAIT', 30.0))
                    print(f"[Quiz] Recorder did not confirm within {timeout}s; attempting an extra {extra_wait}s retry before proceeding...")
                    start2 = time.time()
                    while time.time() - start2 < extra_wait:
                        try:
                            try:
                                _http_session.get('http://127.0.0.1:5000/oled', params={'value': 'report', 'progress': 100}, timeout=0.8)
                            except Exception:
                                pass
                            r2 = _http_session.get(status_url, timeout=1.0)
                            if r2 is not None and getattr(r2, 'status_code', None) == 200:
                                try:
                                    j2 = r2.json()
                                    if j2.get('completion_played') in (True, 'true', 'True', '1'):
                                            completed = True
                                            try:
                                                globals()['_REPORT_COMPLETION_SHOWN'] = True
                                            except Exception:
                                                pass
                                            print("[Quiz] Recorder confirmed report completion animation shown (during extra retry)")
                                            break
                                except Exception:
                                    pass
                        except Exception:
                            pass
                        time.sleep(poll_interval)
                    if not completed:
                        print("[Quiz] Warning: recorder did not confirm completion after extra retry; proceeding anyway")
            except Exception:
                pass

            # Now stop the report animation loop
            _report_creation_stop_event.set()
            print("[Quiz] Stopped report-creation OLED animation")
    except Exception:
        pass

    print("\n" + "="*70)
    print("✓ Quiz completed! Returning to Assisting Mode...")
    print("="*70)
    
    # CRITICAL: Re-enable state button after quiz ends
    print("[Quiz] Re-enabling state button - student can now change modes...")
    try:
        enable_button_resp = requests.get('http://127.0.0.1:5000/enable_state_button', timeout=5)
        if enable_button_resp.status_code == 200:
            print("✓ State button RE-ENABLED - student can change modes now")
        else:
            print(f"⚠ Could not enable button: status {enable_button_resp.status_code}")
    except Exception as e:
        print(f"⚠ Could not enable state button: {e}")
    
    # Show a final completion text on the OLED, then return to assisting mode.
    try:
        # Only send the final text if it hasn't already been shown during report creation
        if not globals().get('_REPORT_COMPLETION_SHOWN', False):
            # Send custom two-line text to the recorder OLED via Flask proxy.
            final_line1 = "Reports Complete"
            final_line2 = "Back to Assisting Mode"
            print(f"[Quiz] Sending final completion message to OLED: {final_line1} | {final_line2}")
            try:
                _http_session.get('http://127.0.0.1:5000/oled', params={
                    'value': 'text',
                    'line1': final_line1,
                    'line2': final_line2,
                    'hold': 1600
                }, timeout=4)
            except Exception:
                # Quiet fallback: set neutral expression if text display fails
                set_oled_expression('idle')
            # Small pause to allow state to settle and message to be perceived
            time.sleep(1.6)
        else:
            print('[Quiz] Final completion already shown earlier; skipping OLED text')
        # After the (possible) message, switch back to assisting mode
        set_state_to_assisting_mode()  # Switch back to state 0
    except Exception:
        try:
            set_oled_expression('idle')
            set_state_to_assisting_mode()
            time.sleep(2)
        except Exception:
            pass

    # Short welcome back in Assisting Mode: say a friendly short line and wait for student
    try:
        first_name = None
        try:
            if CURRENT_STUDENT_NAME and isinstance(CURRENT_STUDENT_NAME, str) and CURRENT_STUDENT_NAME.strip():
                first_name = CURRENT_STUDENT_NAME.strip().split()[0]
        except Exception:
            first_name = None

        welcome_text = f"Welcome back {first_name}. I'm ready to help you — what is your question?" if first_name else "Welcome back. Handa na akong tumulong — ano ang iyong tanong?"
        # Show a brief happy expression while welcoming
        set_oled_expression('happy')
        synthesize_speech(welcome_text)
        play_audio('response.wav')
        set_oled_expression('idle')
        # Reset prompt timer so the long-silence prompt won't trigger immediately
        last_prompt_time = time.time()
        print(f"[Assist] Played welcome back message to student: {first_name if first_name else '<unknown>'}")
    except Exception as _e:
        print(f"[Assist] Could not play welcome back message: {_e}")

    # Reset assisting-mode session state so the system behaves like a fresh start
    try:
        # Instead of immediately resuming listening (which can trigger the
        # "Sorry, I didn't hear that" retry loop), set a short grace period
        # after returning from quiz. During this window GENTA will not start
        # a new recording session, preventing aggressive immediate re-prompts.
        global _assist_defer_until
        _assist_defer_until = time.time() + 4.0  # seconds to defer listening
        first_iteration = True
        # Reset prompt timer so follow-up prompts still respect cooldown
        last_prompt_time = time.time()
        print(f"✓ Reset assisting-mode session flags: first_iteration=True; deferring listen for {_assist_defer_until - time.time():.1f}s")
    except Exception:
        pass

def state_monitor_thread():
    """Background thread that monitors state changes from ESP32."""
    global _STATE_CHANGE_REQUESTED, _STATE_MONITOR_ACTIVE, _CURRENT_STATE
    
    print("[State Monitor] Thread started")
    _STATE_MONITOR_ACTIVE = True
    # Faster checks to reduce perceived latency when student presses the button
    check_interval = 0.6  # Check state every 600ms
    
    while _STATE_MONITOR_ACTIVE:
        try:
            # If a QUIZ flow is active (student confirmed and QUIZZER running),
            # or if report creation is in progress, avoid fetching state to
            # prevent interference with these longer-running tasks.
            try:
                if globals().get('_QUIZ_FLOW_ACTIVE') or globals().get('_REPORT_CREATION_ACTIVE'):
                    # Sleep a bit and re-check later. This avoids hammering the
                    # state endpoint while reports are being generated and
                    # ensures report creation isn't interrupted by state reads.
                    time.sleep(check_interval)
                    continue
            except Exception:
                pass
            # Read current state from ESP32
            new_state = read_state_text_file(url_state)
            
            if new_state and new_state.strip():
                new_state = new_state.strip()
                
                # Check if state has changed. Use a short debounce to avoid reacting
                # to transient or noisy readings (network glitches or GPIO bounce).
                if new_state != _CURRENT_STATE:
                    # Short confirmation read to ensure the new state persists
                    confirmed = False
                    try:
                        time.sleep(0.18)
                        confirm_state = read_state_text_file(url_state)
                        if confirm_state and confirm_state.strip() == new_state:
                            confirmed = True
                        else:
                            print(f"[State Monitor] Transient state '{new_state}' ignored (not confirmed)")
                    except Exception:
                        # If confirmation read fails, be conservative and ignore
                        print("[State Monitor] Confirmation read failed; ignoring transient state change")

                    if not confirmed:
                        # Skip processing this transient change
                        time.sleep(check_interval)
                        continue

                    print(f"\n[State Monitor] State change detected: {_CURRENT_STATE} → {new_state}")
                    print(f"[State Monitor] {('ASSISTING MODE' if new_state == '0' else 'QUIZ MODE')}")

                    # If switching to Quiz Mode, give immediate feedback to the student
                    # by updating the OLED and playing a short prompt. This is intentionally
                    # lightweight and non-blocking so the long-running Assisting loop can
                    # still wind down gracefully while the student hears the confirmation.
                    try:
                        # We've detected a request to switch to Quiz Mode. Do not play
                        # a TTS confirmation here — let the main loop handle the full
                        # confirmation flow to avoid duplicate prompts and race conditions.
                        if new_state == '1':
                            try:
                                # Update OLED text only (no TTS) to give visual feedback.
                                _http_session.get('http://127.0.0.1:5000/oled', params={
                                    'value': 'text',
                                    'line1': 'Quiz requested',
                                    'line2': 'Handa ka na ba?',
                                    'hold': 3000
                                }, timeout=1.5)
                            except Exception:
                                try:
                                    set_oled_expression('quiz')
                                except Exception:
                                    pass
                            # Cancel any in-progress recording waits so the main loop
                            # can proceed immediately to the confirmation dialog.
                            try:
                                _http_session.get('http://127.0.0.1:5000/cancel_recording_wait', timeout=1.0)
                                print('[State Monitor] Sent cancel notification to Flask to abort recording waits')
                            except Exception:
                                pass
                    except Exception:
                        pass

                    # Signal the main loop to handle the full confirmation flow and restart
                    _STATE_CHANGE_REQUESTED = True
                    _STATE_MONITOR_ACTIVE = False  # Stop monitoring to allow restart
                    break
            
            time.sleep(check_interval)
            
        except Exception as e:
            # Silent fail - don't spam console with connection errors
            time.sleep(check_interval)
    
    print("[State Monitor] Thread stopped")

def main():
    global GENTA_State, _STATE_CHANGE_REQUESTED, _STATE_MONITOR_ACTIVE, _CURRENT_STATE
    
    print("\n" + "="*70)
    print("GENTA SYSTEM STARTING - State Monitoring Enabled")
    print("="*70)
    
    # CRITICAL: Clear old recordings on startup to ensure fresh session
    print("\n[Startup Cleanup] Clearing old recordings from ESP32...")
    try:
        clear_resp = _http_session.get('https://nonbasic-bob-inimical.ngrok-free.dev/clear', timeout=5)
        if clear_resp.status_code == 200:
            print("✓ Old recordings cleared successfully")
        else:
            print(f"⚠ Clear returned status: {clear_resp.status_code}")
    except Exception as e:
        print(f"⚠ Could not clear old recordings: {e}")
    
    # Delete local recording files too
    try:
        if os.path.exists(audio_raw_path):
            os.remove(audio_raw_path)
            print("✓ Removed local recording.wav")
        if os.path.exists(audio_converted_path):
            os.remove(audio_converted_path)
            print("✓ Removed local Recording.wav")
    except Exception as e:
        print(f"⚠ Local cleanup warning: {e}")
    
    # FIRST THING: Set state to 0 (Assisting Mode) on startup
    set_state_to_assisting_mode()
    time.sleep(1)  # Give ESP32 time to update state file

    # Optional: exercise OLED expressions for diagnostics if requested via env var
    try:
        if os.environ.get('GENTA_OLED_EXERCISE', '0').strip() == '1':
            print("[OLED TEST] GENTA_OLED_EXERCISE=1 -> exercising expressions now")
            try:
                threading.Thread(target=exercise_all_oled_expressions, args=(0.6,), daemon=True).start()
            except Exception:
                # best-effort only
                try:
                    exercise_all_oled_expressions(0.6)
                except Exception:
                    pass
    except Exception:
        pass
    
    print("Monitoring state from: " + url_state)
    
    # Test if ESP32 is serving state.txt properly
    print("\n[System Check] Testing ESP32 state.txt endpoint...")
    try:
        test_state = read_state_text_file(url_state)
        if test_state is not None:
            print(f"✓ ESP32 state.txt accessible: Current state = {test_state}")
        else:
            print("⚠ WARNING: Cannot reach state.txt endpoint")
            print("  Make sure:")
            print("  1. ESP32 is powered on and connected")
            print("  2. Flask proxy is running")
            print("  3. ngrok tunnel is active")
    except Exception as e:
        print(f"⚠ State endpoint test failed: {e}")
    
    print("\nPress GPIO 22 button on ESP32 to toggle between modes:")
    print("  State 0 = Assisting Mode")
    print("  State 1 = Quiz Mode")
    print("  Note: Button is DISABLED until LRN is entered")
    print("="*70 + "\n")
    
    while True:
        try:
            # Read current state from ESP32
            GENTA_State = read_state_text_file(url_state)
            
            if GENTA_State is None or not GENTA_State.strip():
                GENTA_State = "0"
                print("⚠ Could not read state from server. Using default state (0)...")
            
            # Clean the state value (remove whitespace)
            GENTA_State = GENTA_State.strip()
            _CURRENT_STATE = GENTA_State
            _STATE_CHANGE_REQUESTED = False
            
            # Start background state monitoring thread
            monitor_thread = threading.Thread(target=state_monitor_thread, daemon=True)
            monitor_thread.start()
            
            if GENTA_State == "0":
                print("\n" + "="*70)
                print("🎓 GENTA ASSISTING MODE (State 0) - ACTIVE")
                print("="*70 + "\n")
                
                try:
                    GENTA()
                    # If GENTA() exits normally, check if it was due to state change
                    if not _STATE_CHANGE_REQUESTED:
                        print("\n✓ GENTA Assisting Mode session completed normally")
                        break
                except KeyboardInterrupt:
                    print("\n\n⚠ Keyboard interrupt detected. Exiting GENTA system...")
                    _STATE_MONITOR_ACTIVE = False
                    break
                except Exception as e:
                    print(f"\n⚠ Error in GENTA Assisting Mode: {e}")
                    _STATE_MONITOR_ACTIVE = False
                    if not _STATE_CHANGE_REQUESTED:
                        print("Restarting in 3 seconds...")
                        time.sleep(3)
                    
            elif GENTA_State == "1":
                print("\n" + "="*70)
                print("📝 GENTA QUIZ MODE (State 1) - ACTIVE")
                print("="*70 + "\n")
                
                try:
                    # Pause state polling while we run the confirmation + quiz flow
                    try:
                        globals()['_QUIZ_FLOW_ACTIVE'] = True
                        print("[State] QUIZ flow PAUSING state monitor polling (confirmation + quiz)")
                    except Exception:
                        pass
                    # If report creation is currently active, defer asking for quiz confirmation
                    # to avoid interrupting report generation.
                    if globals().get('_REPORT_CREATION_ACTIVE'):
                        print("[State] Report creation in progress — deferring quiz confirmation until complete")
                        # Wait a short while for the report task to finish (non-blocking expectation).
                        wait_start = time.time()
                        while globals().get('_REPORT_CREATION_ACTIVE') and (time.time() - wait_start) < 30:
                            time.sleep(1)
                        if globals().get('_REPORT_CREATION_ACTIVE'):
                            print("[State] Report creation still active after wait — skipping confirmation this cycle")
                            # Skip confirmation for now; re-check next loop iteration
                            continue
                    # Show a state-change text animation on the OLED asking for confirmation
                    try:
                        sc_line1 = "State change"
                        sc_line2 = "confirmation: Quiz Mode?"

                        print(f"[State] Sending state-change prompt to OLED: {sc_line1} | {sc_line2}")
                        _http_session.get('http://127.0.0.1:5000/oled', params={
                            'value': 'text',
                            'line1': sc_line1,
                            'line2': sc_line2,
                            'hold': 1800
                        }, timeout=4)
                    except Exception:
                        # Fallback: show the default quiz expression if custom text fails
                        try:
                            set_oled_expression('quiz')
                        except Exception:
                            pass

                    # IMPORTANT: Ask for confirmation before starting quiz
                    quiz_confirmed = confirm_quiz_readiness()
                    
                    if quiz_confirmed:
                        # Student confirmed - proceed with quiz
                        # Trigger OLED quiz transition animation (non-blocking)
                        try:
                            set_oled_expression('quiz')
                        except Exception:
                            pass
                        # Student confirmed: mark quiz flow active so state monitor pauses
                        try:
                            globals()['_QUIZ_FLOW_ACTIVE'] = True
                            print("[State] QUIZ flow active - state monitor will pause polling")
                        except Exception:
                            pass

                        # Start QUIZZER and show loading animation since student confirmed
                        QUIZZER(start_loading=True)

                        # QUIZZER returned - clear QUIZ flow active so state monitor resumes
                        try:
                            globals()['_QUIZ_FLOW_ACTIVE'] = False
                            print("[State] QUIZ flow ended - state monitor may resume polling")
                        except Exception:
                            pass
                        # After QUIZZER completes, it automatically switches to state 0
                        # Continue loop to restart in Assisting Mode
                        print("\n🔄 Switching to Assisting Mode...")
                        _STATE_CHANGE_REQUESTED = True  # Trigger loop restart
                    else:
                        # Student declined or unclear - return to Assisting Mode
                        print("\n🔄 Student declined quiz, returning to Assisting Mode...")
                        try:
                            globals()['_QUIZ_FLOW_ACTIVE'] = False
                            print("[State] QUIZ flow cleared (declined) - state monitor may resume")
                        except Exception:
                            pass
                        set_state_to_assisting_mode()  # Reset state to 0
                        time.sleep(1)
                        _STATE_CHANGE_REQUESTED = True  # Trigger loop restart
                        
                except KeyboardInterrupt:
                    print("\n\n⚠ Keyboard interrupt detected. Exiting GENTA system...")
                    try:
                        globals()['_QUIZ_FLOW_ACTIVE'] = False
                    except Exception:
                        pass
                    _STATE_MONITOR_ACTIVE = False
                    break
                except Exception as e:
                    print(f"\n⚠ Error in GENTA Quiz Mode: {e}")
                    _STATE_MONITOR_ACTIVE = False
                    try:
                        globals()['_QUIZ_FLOW_ACTIVE'] = False
                    except Exception:
                        pass
                    if not _STATE_CHANGE_REQUESTED:
                        print("Restarting in 3 seconds...")
                        time.sleep(3)
                    
            else:
                print(f"⚠ Invalid state '{GENTA_State}' received from server.")
                print("Valid states: 0 (Assisting) or 1 (Quiz)")
                print("Waiting 5 seconds before retry...")
                _STATE_MONITOR_ACTIVE = False
                time.sleep(5)
            
            # If state change was requested, loop will restart with new state
            if _STATE_CHANGE_REQUESTED:
                _STATE_MONITOR_ACTIVE = False
                print("\n" + "="*70)
                print("🔄 RESTARTING GENTA with new state...")
                print("="*70)
                time.sleep(1)
                continue
                
        except KeyboardInterrupt:
            print("\n\n⚠ Keyboard interrupt detected. Exiting GENTA system...")
            _STATE_MONITOR_ACTIVE = False
            break
        except Exception as e:
            print(f"\n⚠ Unexpected error in main loop: {e}")
            _STATE_MONITOR_ACTIVE = False
            print("Restarting in 5 seconds...")
            time.sleep(5)

if __name__ == "__main__":
    # Support a test hook to run the QUIZZER directly for end-to-end smoke tests
    # without going through the full state monitoring loop. Set env var
    # `RUN_QUIZ_ON_STARTUP=1` to execute a single QUIZZER(start_loading=True)
    # run and exit. This is useful for automated testing and reproducing the
    # report-creation OLED animation flow.
    if os.environ.get('RUN_QUIZ_ON_STARTUP', '').strip() == '1':
        try:
            print("RUN_QUIZ_ON_STARTUP=1 -> running QUIZZER directly for smoke test")
            # Ensure minimal student context exists so filenames are sane
            try:
                if not CURRENT_STUDENT_NAME:
                    CURRENT_STUDENT_NAME = 'TestStudent'
            except Exception:
                CURRENT_STUDENT_NAME = 'TestStudent'
            try:
                if not CURRENT_STUDENT_ID:
                    CURRENT_STUDENT_ID = '0000'
            except Exception:
                CURRENT_STUDENT_ID = '0000'
            # Run QUIZZER once (it will create reports and exercise the report animation)
            QUIZZER(start_loading=True)
        except Exception as e:
            print(f"Error running QUIZZER via RUN_QUIZ_ON_STARTUP: {e}")
    else:
        main()